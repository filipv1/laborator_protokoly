"""
Finální test odstranění prázdných PP tabulek
Vytvoří dokument podobný skutečnému PP protokolu s tabulkou OSTATNI
"""
from docx import Document
from generate_word_from_two_sources import remove_empty_pp_rows


def create_ostatni_only_table():
    """Vytvoří dokument s tabulkou která má jen OSTATNI (jak to vypadá v reálném PP)"""
    doc = Document()

    doc.add_heading("Test PP - OSTATNI tabulka", 0)

    # Přidej několik normálních tabulek
    doc.add_paragraph("Normální text před tabulkou")

    # Tabulka OSTATNI - jak vypadá v reálném dokumentu
    table = doc.add_table(rows=3, cols=6)

    # Řádek 0 - záhlaví
    table.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table.rows[0].cells[1].text = "Svalová práce"
    table.rows[0].cells[2].text = "MUŽ 1"
    table.rows[0].cells[3].text = "MUŽ 2"
    table.rows[0].cells[4].text = "Ø\nvýskyt\nza směnu\n[min]"
    table.rows[0].cells[5].text = "Typ pracovní polohy"

    # Řádek 1 - druhé záhlaví (obsahuje "Výskyt za směnu")
    table.rows[1].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table.rows[1].cells[1].text = "Svalová práce"
    table.rows[1].cells[2].text = "Výskyt\nza směnu\n[min]"
    table.rows[1].cells[3].text = "Výskyt\nza směnu\n[min]"
    table.rows[1].cells[4].text = "Ø\nvýskyt\nza směnu\n[min]"
    table.rows[1].cells[5].text = "Typ pracovní polohy"

    # Řádek 2 - OSTATNI
    table.rows[2].cells[0].text = "OSTATNI"
    table.rows[2].cells[1].text = ""
    table.rows[2].cells[2].text = ""
    table.rows[2].cells[3].text = ""
    table.rows[2].cells[4].text = ""
    table.rows[2].cells[5].text = ""

    doc.add_paragraph("Text po tabulce")

    doc.save("test_ostatni_table.docx")
    print("✓ Vytvořen dokument s OSTATNI tabulkou: test_ostatni_table.docx")
    return "test_ostatni_table.docx"


def test_ostatni_removal():
    """Test odstranění OSTATNI tabulky"""
    print("\n" + "="*60)
    print("FINÁLNÍ TEST - ODSTRANĚNÍ OSTATNI TABULKY")
    print("="*60)

    # Vytvoř testovací dokument
    test_path = create_ostatni_only_table()

    # Spočítej tabulky před
    doc = Document(test_path)
    tables_before = len(doc.tables)
    print(f"\nPočet tabulek PŘED: {tables_before}")

    # Spusť funkci
    print("\nSpouštím remove_empty_pp_rows()...")
    print("-"*60)
    remove_empty_pp_rows(test_path)
    print("-"*60)

    # Spočítej tabulky po
    doc = Document(test_path)
    tables_after = len(doc.tables)
    print(f"\nPočet tabulek PO: {tables_after}")

    # Ověř výsledek
    print("\n" + "="*60)
    if tables_before == 1 and tables_after == 0:
        print("✓✓✓ ÚSPĚCH! Prázdná OSTATNI tabulka byla odstraněna!")
        print("Problém vyřešen - žádné prázdné tabulky v PP protokolech!")
        return True
    else:
        print(f"✗✗✗ SELHÁNÍ! Tabulka nebyla odstraněna")
        print(f"  Očekáváno: {tables_before} → 0")
        print(f"  Skutečnost: {tables_before} → {tables_after}")
        return False


if __name__ == "__main__":
    result = test_ostatni_removal()
    print("="*60)