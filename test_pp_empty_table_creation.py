"""
Vytvoření testovacího PP dokumentu s prázdnými tabulkami
"""
from docx import Document
from docx.shared import Pt
from generate_word_from_two_sources import remove_empty_pp_rows


def create_test_pp_document():
    """Vytvoří testovací PP dokument s prázdnými tabulkami"""
    doc = Document()

    # Přidej nadpis
    doc.add_heading("Test PP dokument s prázdnými tabulkami", 0)

    # Vytvoř několik PP tabulek s různými stavy

    # 1. Tabulka TRUP - všechny hodnoty 0 (měla by být odstraněna)
    doc.add_paragraph("Hodnocení pracovních poloh - TRUP")
    table1 = doc.add_table(rows=5, cols=6)
    # Záhlaví
    table1.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table1.rows[0].cells[1].text = "Svalová práce"
    table1.rows[0].cells[2].text = "MUŽ 1"
    table1.rows[0].cells[3].text = "MUŽ 2"
    table1.rows[0].cells[4].text = "Ø výskyt za směnu [min]"
    table1.rows[0].cells[5].text = "Typ pracovní polohy"
    # Section header
    table1.rows[1].cells[0].text = "TRUP"
    # Data řádky - všechny 0
    table1.rows[2].cells[0].text = "Předklon trupu"
    table1.rows[2].cells[1].text = "Statická"
    table1.rows[2].cells[2].text = "0"
    table1.rows[2].cells[3].text = "0"
    table1.rows[2].cells[4].text = "0"
    table1.rows[2].cells[5].text = "N"

    table1.rows[3].cells[0].text = "Záklon trupu"
    table1.rows[3].cells[1].text = "Dynamická"
    table1.rows[3].cells[2].text = "0"
    table1.rows[3].cells[3].text = "0"
    table1.rows[3].cells[4].text = "0"
    table1.rows[3].cells[5].text = "N"

    # OSTATNI řádek
    table1.rows[4].cells[0].text = "OSTATNI"
    table1.rows[4].cells[1].text = "Ultrathink"

    # 2. Tabulka PHK - některé hodnoty nenulové (neměla by být odstraněna)
    doc.add_paragraph("")
    doc.add_paragraph("Hodnocení pracovních poloh - PHK")
    table2 = doc.add_table(rows=5, cols=6)
    # Záhlaví
    table2.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table2.rows[0].cells[1].text = "Svalová práce"
    table2.rows[0].cells[2].text = "MUŽ 1"
    table2.rows[0].cells[3].text = "MUŽ 2"
    table2.rows[0].cells[4].text = "Ø výskyt za směnu [min]"
    table2.rows[0].cells[5].text = "Typ pracovní polohy"
    # Section header
    table2.rows[1].cells[0].text = "PHK"
    # Data řádky - některé nenulové
    table2.rows[2].cells[0].text = "Elevace paže"
    table2.rows[2].cells[1].text = "Statická"
    table2.rows[2].cells[2].text = "5"  # Nenulové!
    table2.rows[2].cells[3].text = "0"
    table2.rows[2].cells[4].text = "2.5"  # Nenulové!
    table2.rows[2].cells[5].text = "PP"

    table2.rows[3].cells[0].text = "Flexe paže"
    table2.rows[3].cells[1].text = "Dynamická"
    table2.rows[3].cells[2].text = "0"
    table2.rows[3].cells[3].text = "0"
    table2.rows[3].cells[4].text = "0"
    table2.rows[3].cells[5].text = "N"

    # OSTATNI řádek
    table2.rows[4].cells[0].text = "OSTATNI"
    table2.rows[4].cells[1].text = "Ultrathink"

    # 3. Tabulka HLAVA A KRK - všechny hodnoty 0 (měla by být odstraněna)
    doc.add_paragraph("")
    doc.add_paragraph("Hodnocení pracovních poloh - HLAVA A KRK")
    table3 = doc.add_table(rows=4, cols=6)
    # Záhlaví
    table3.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table3.rows[0].cells[1].text = "Svalová práce"
    table3.rows[0].cells[2].text = "MUŽ 1"
    table3.rows[0].cells[3].text = "MUŽ 2"
    table3.rows[0].cells[4].text = "Ø výskyt za směnu [min]"
    table3.rows[0].cells[5].text = "Typ pracovní polohy"
    # Section header
    table3.rows[1].cells[0].text = "HLAVA_KRK"
    # Data řádek - všechny 0
    table3.rows[2].cells[0].text = "Předklon hlavy"
    table3.rows[2].cells[1].text = "Statická"
    table3.rows[2].cells[2].text = "0"
    table3.rows[2].cells[3].text = "0"
    table3.rows[2].cells[4].text = "0"
    table3.rows[2].cells[5].text = "N"
    # OSTATNI řádek
    table3.rows[3].cells[0].text = "OSTATNI"
    table3.rows[3].cells[1].text = "Ultrathink"

    # Ulož dokument
    doc.save("test_pp_with_empty_tables.docx")
    print("✓ Vytvořen testovací dokument: test_pp_with_empty_tables.docx")
    print("  - Tabulka TRUP: všechny hodnoty 0 (měla by být odstraněna)")
    print("  - Tabulka PHK: některé hodnoty nenulové (neměla by být odstraněna)")
    print("  - Tabulka HLAVA_KRK: všechny hodnoty 0 (měla by být odstraněna)")


def count_tables(docx_path):
    """Spočítá všechny tabulky v dokumentu"""
    doc = Document(docx_path)
    return len(doc.tables)


def test_empty_table_removal():
    """Test odstranění prázdných tabulek"""
    print("\n" + "="*60)
    print("TEST ODSTRANĚNÍ CELÝCH PRÁZDNÝCH PP TABULEK")
    print("="*60)

    # Vytvoř testovací dokument
    create_test_pp_document()

    test_path = "test_pp_with_empty_tables.docx"

    # Spočítej tabulky před
    tables_before = count_tables(test_path)
    print(f"\nPočet tabulek před úpravou: {tables_before}")

    # Spusť funkci
    print("\nSpouštím remove_empty_pp_rows()...")
    print("-"*60)
    remove_empty_pp_rows(test_path)
    print("-"*60)

    # Spočítej tabulky po
    tables_after = count_tables(test_path)
    print(f"\nPočet tabulek po úpravě: {tables_after}")

    # Vyhodnocení
    if tables_before == 3 and tables_after == 1:
        print("\n✓✓✓ ÚSPĚCH! Test prošel správně:")
        print("  - Odstraněny 2 prázdné tabulky (TRUP, HLAVA_KRK)")
        print("  - Zachována 1 tabulka s daty (PHK)")
    else:
        print(f"\n✗✗✗ CHYBA! Neočekávaný výsledek:")
        print(f"  - Očekáváno: 3 tabulky → 1 tabulka")
        print(f"  - Skutečnost: {tables_before} tabulek → {tables_after} tabulek")

    print(f"\nVýsledný dokument: {test_path}")
    print("Otevřete v MS Word pro vizuální kontrolu.")


if __name__ == "__main__":
    test_empty_table_removal()