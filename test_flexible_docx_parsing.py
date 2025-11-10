"""
Test flexibilního parsování Word tabulek
Testuje různé scénáře s chybějícími sloupci
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.shared import Pt
from core.docx_parser import DocxParser


def create_test_document(filename, headers, data_rows):
    """
    Vytvoří testovací Word dokument s tabulkou.

    Args:
        filename: Název souboru
        headers: List názvů sloupců
        data_rows: List řádků (každý řádek je list hodnot)
    """
    doc = Document()

    # Přidej info tabulku (první tabulka - dummy)
    doc.add_paragraph("Informační tabulka:")
    info_table = doc.add_table(rows=2, cols=2)
    info_table.cell(0, 0).text = "Místo měření"
    info_table.cell(0, 1).text = "Test"

    # Přidej časový snímek (druhá tabulka)
    doc.add_paragraph("\nČasový snímek:")

    num_cols = len(headers)
    num_rows = len(data_rows) + 1  # +1 pro header

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # Vyplň header
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = header
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Vyplň data
    for row_idx, row_data in enumerate(data_rows, start=1):
        for col_idx, value in enumerate(row_data):
            if col_idx < num_cols:  # Safety check
                table.cell(row_idx, col_idx).text = str(value) if value is not None else ""

    # Ulož
    doc.save(filename)
    print(f"✓ Vytvořen: {filename}")


def test_scenario(name, doc_path):
    """Otestuje jeden scénář"""
    print(f"\n{'='*60}")
    print(f"TEST: {name}")
    print(f"{'='*60}")

    result = DocxParser.parse_time_schedule_table(str(doc_path))

    print(f"\n📊 VÝSLEDEK:")
    print(f"   Celkem načteno řádků: {sum(1 for k, v in result.items() if k.startswith('line') and v['operation'])}")

    # Zobraz prvních 5 řádků
    print(f"\n   První 3 řádky:")
    for i in range(1, 4):
        line = result.get(f"line{i}", {})
        if line.get('operation'):
            print(f"   Line {i}:")
            print(f"      operation: {line.get('operation')}")
            print(f"      time_min: {line.get('time_min')}")
            print(f"      pieces_count: {line.get('pieces_count')}")

    print(f"\n   Total:")
    total = result.get("total", {})
    print(f"      time_min: {total.get('time_min')}")
    print(f"      pieces_count: {total.get('pieces_count')}")

    return result


def main():
    print("="*60)
    print("TESTOVÁNÍ FLEXIBILNÍHO PARSOVÁNÍ TABULEK")
    print("="*60)

    test_dir = Path("test_docx_flexible")
    test_dir.mkdir(exist_ok=True)

    # ===== SCÉNÁŘ 1: Plná tabulka (4 sloupce) =====
    doc1 = test_dir / "test_full_table.docx"
    create_test_document(
        doc1,
        headers=["Č.", "Činnost", "Čas (min)", "Počet kusů"],
        data_rows=[
            ["1", "Montáž dílu A", 45, 120],
            ["2", "Kontrola kvality", 30, 80],
            ["3", "Balení", 25, 100],
            ["Celkem", "", 100, 300]
        ]
    )
    result1 = test_scenario("Scénář 1: PLNÁ TABULKA (4 sloupce)", doc1)

    # ===== SCÉNÁŘ 2: Bez sloupce "Počet kusů" =====
    doc2 = test_dir / "test_no_pieces.docx"
    create_test_document(
        doc2,
        headers=["Č.", "Činnost", "Čas (min)"],
        data_rows=[
            ["1", "Montáž dílu A", 45],
            ["2", "Kontrola kvality", 30],
            ["3", "Balení", 25],
            ["Celkem", "", 100]
        ]
    )
    result2 = test_scenario("Scénář 2: BEZ SLOUPCE 'Počet kusů'", doc2)

    # ===== SCÉNÁŘ 3: Bez sloupce "Čas" =====
    doc3 = test_dir / "test_no_time.docx"
    create_test_document(
        doc3,
        headers=["Č.", "Činnost", "Počet kusů"],
        data_rows=[
            ["1", "Montáž dílu A", 120],
            ["2", "Kontrola kvality", 80],
            ["3", "Balení", 100],
            ["Celkem", "", 300]
        ]
    )
    result3 = test_scenario("Scénář 3: BEZ SLOUPCE 'Čas'", doc3)

    # ===== SCÉNÁŘ 4: Jen činnosti (2 sloupce) =====
    doc4 = test_dir / "test_operations_only.docx"
    create_test_document(
        doc4,
        headers=["Č.", "Činnost"],
        data_rows=[
            ["1", "Montáž dílu A"],
            ["2", "Kontrola kvality"],
            ["3", "Balení"]
        ]
    )
    result4 = test_scenario("Scénář 4: JEN ČINNOSTI (2 sloupce)", doc4)

    # ===== SCÉNÁŘ 5: Jiné pořadí sloupců =====
    doc5 = test_dir / "test_different_order.docx"
    create_test_document(
        doc5,
        headers=["Činnost", "Čas (min)", "Č.", "Počet kusů"],
        data_rows=[
            ["Montáž dílu A", 45, "1", 120],
            ["Kontrola kvality", 30, "2", 80],
            ["Balení", 25, "3", 100],
            ["", 100, "Celkem", 300]
        ]
    )
    result5 = test_scenario("Scénář 5: JINÉ POŘADÍ SLOUPCŮ", doc5)

    # ===== SCÉNÁŘ 6: Různé názvy sloupců =====
    doc6 = test_dir / "test_different_names.docx"
    create_test_document(
        doc6,
        headers=["Por.", "Popis práce", "Doba [min]", "Kusy"],
        data_rows=[
            ["1", "Montáž dílu A", 45, 120],
            ["2", "Kontrola kvality", 30, 80],
            ["3", "Balení", 25, 100],
            ["Celkem", "", 100, 300]
        ]
    )
    result6 = test_scenario("Scénář 6: RŮZNÉ NÁZVY SLOUPCŮ", doc6)

    # ===== SCÉNÁŘ 7: Čísla jako text (edge case) =====
    doc7 = test_dir / "test_numbers_as_text.docx"
    create_test_document(
        doc7,
        headers=["Č.", "Činnost", "Čas (min)", "Počet kusů"],
        data_rows=[
            ["1", "Montáž dílu A", "45", "120"],
            ["2", "Kontrola kvality", "třicet", "osmdesát"],  # Text místo čísel
            ["3", "Balení", "25", "100"],
        ]
    )
    result7 = test_scenario("Scénář 7: ČÍSLA JAKO TEXT (edge case)", doc7)

    # ===== SUMMARY =====
    print(f"\n{'='*60}")
    print(f"SHRNUTÍ TESTŮ")
    print(f"{'='*60}")
    print(f"✓ Všechny testy dokončeny")
    print(f"✓ Testovací dokumenty uloženy v: {test_dir.absolute()}")
    print(f"\nVÝSLEDKY:")
    print(f"  Scénář 1 (plná tabulka): {len([k for k,v in result1.items() if k.startswith('line') and v['operation']])} řádků")
    print(f"  Scénář 2 (bez kusů): {len([k for k,v in result2.items() if k.startswith('line') and v['operation']])} řádků")
    print(f"  Scénář 3 (bez času): {len([k for k,v in result3.items() if k.startswith('line') and v['operation']])} řádků")
    print(f"  Scénář 4 (jen činnosti): {len([k for k,v in result4.items() if k.startswith('line') and v['operation']])} řádků")
    print(f"  Scénář 5 (jiné pořadí): {len([k for k,v in result5.items() if k.startswith('line') and v['operation']])} řádků")
    print(f"  Scénář 6 (různé názvy): {len([k for k,v in result6.items() if k.startswith('line') and v['operation']])} řádků")
    print(f"  Scénář 7 (čísla jako text): {len([k for k,v in result7.items() if k.startswith('line') and v['operation']])} řádků")


if __name__ == "__main__":
    main()
