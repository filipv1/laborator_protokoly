"""
Test finální opravy PP table removal - záhlaví se už nepočítá jako data
"""
from docx import Document
from generate_word_from_two_sources import remove_empty_pp_rows


def create_test_document_with_header():
    """Vytvoří testovací dokument s PP tabulkou která má jen záhlaví a OSTATNI"""
    doc = Document()

    doc.add_heading("Test PP - pouze záhlaví a OSTATNI", 0)

    # Vytvoř tabulku jak vypadá v reálném PP dokumentu
    table = doc.add_table(rows=3, cols=6)

    # Řádek 0 - první část záhlaví
    table.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table.rows[0].cells[1].text = "Svalová práce"
    table.rows[0].cells[2].text = "MUŽ 1"
    table.rows[0].cells[3].text = "MUŽ 2"
    table.rows[0].cells[4].text = "Ø\nvýskyt\nza směnu\n[min]"
    table.rows[0].cells[5].text = "Typ pracovní polohy"

    # Řádek 1 - druhá část záhlaví (problematický řádek který se počítal jako data)
    table.rows[1].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table.rows[1].cells[1].text = "Svalová práce"
    table.rows[1].cells[2].text = "Výskyt\nza směnu\n[min]"
    table.rows[1].cells[3].text = "Výskyt\nza směnu\n[min]"
    table.rows[1].cells[4].text = "Ø\nvýskyt\nza směnu\n[min]"
    table.rows[1].cells[5].text = "Typ pracovní polohy"

    # Řádek 2 - OSTATNI
    table.rows[2].cells[0].text = "OSTATNI"
    table.rows[2].cells[1].text = "Ultrathink"
    table.rows[2].cells[2].text = ""
    table.rows[2].cells[3].text = ""
    table.rows[2].cells[4].text = ""
    table.rows[2].cells[5].text = ""

    doc.save("test_pp_header_only.docx")
    print("✓ Vytvořen testovací dokument: test_pp_header_only.docx")
    print("  Obsahuje tabulku pouze se záhlavím a OSTATNI (měla by být odstraněna)")
    return "test_pp_header_only.docx"


def test_header_only_removal():
    """Test odstranění tabulky která má jen záhlaví"""
    print("\n" + "="*60)
    print("TEST ODSTRANĚNÍ TABULKY S POUZE ZÁHLAVÍM")
    print("="*60)

    # Vytvoř testovací dokument
    test_path = create_test_document_with_header()

    # Spočítej tabulky před
    doc = Document(test_path)
    tables_before = len(doc.tables)
    print(f"\nPočet tabulek před úpravou: {tables_before}")

    # Spusť funkci
    print("\nSpouštím remove_empty_pp_rows()...")
    print("-"*60)
    remove_empty_pp_rows(test_path)
    print("-"*60)

    # Spočítej tabulky po
    doc = Document(test_path)
    tables_after = len(doc.tables)
    print(f"\nPočet tabulek po úpravě: {tables_after}")

    # Vyhodnocení
    if tables_before == 1 and tables_after == 0:
        print("\n✓✓✓ ÚSPĚCH! Tabulka pouze se záhlavím byla odstraněna!")
        return True
    else:
        print(f"\n✗✗✗ CHYBA! Tabulka nebyla odstraněna!")
        print(f"  Očekáváno: 1 → 0")
        print(f"  Skutečnost: {tables_before} → {tables_after}")
        return False


def create_mixed_test_document():
    """Vytvoří dokument s více PP tabulkami - některé prázdné, některé s daty"""
    doc = Document()

    doc.add_heading("Test PP - mix prázdných a plných tabulek", 0)

    # TABULKA 1: Pouze záhlaví a OSTATNI (měla by být odstraněna)
    doc.add_paragraph("Tabulka 1: Pouze záhlaví")
    table1 = doc.add_table(rows=3, cols=6)
    # Záhlaví
    table1.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table1.rows[0].cells[2].text = "MUŽ 1"
    table1.rows[0].cells[3].text = "MUŽ 2"
    table1.rows[0].cells[4].text = "Průměr"
    # Section
    table1.rows[1].cells[0].text = "TRUP"
    # OSTATNI
    table1.rows[2].cells[0].text = "OSTATNI"
    table1.rows[2].cells[1].text = "Ultrathink"

    # TABULKA 2: S datovými řádky (neměla by být odstraněna)
    doc.add_paragraph("\nTabulka 2: S daty")
    table2 = doc.add_table(rows=4, cols=6)
    # Záhlaví
    table2.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table2.rows[0].cells[2].text = "MUŽ 1"
    table2.rows[0].cells[3].text = "MUŽ 2"
    table2.rows[0].cells[4].text = "Průměr"
    # Section
    table2.rows[1].cells[0].text = "PHK"
    # Data řádek
    table2.rows[2].cells[0].text = "Vzpažení paže"
    table2.rows[2].cells[1].text = "Statická"
    table2.rows[2].cells[2].text = "10"  # Nenulové!
    table2.rows[2].cells[3].text = "5"
    table2.rows[2].cells[4].text = "7.5"
    # OSTATNI
    table2.rows[3].cells[0].text = "OSTATNI"

    doc.save("test_pp_mixed.docx")
    print("✓ Vytvořen smíšený testovací dokument: test_pp_mixed.docx")
    return "test_pp_mixed.docx"


def test_mixed_tables():
    """Test s více tabulkami"""
    print("\n" + "="*60)
    print("TEST SMÍŠENÝCH TABULEK")
    print("="*60)

    test_path = create_mixed_test_document()

    doc = Document(test_path)
    tables_before = len(doc.tables)
    print(f"Počet tabulek před: {tables_before}")

    print("\nSpouštím remove_empty_pp_rows()...")
    print("-"*60)
    remove_empty_pp_rows(test_path)
    print("-"*60)

    doc = Document(test_path)
    tables_after = len(doc.tables)
    print(f"Počet tabulek po: {tables_after}")

    if tables_before == 2 and tables_after == 1:
        print("\n✓✓✓ ÚSPĚCH! Prázdná tabulka odstraněna, tabulka s daty zachována!")
        return True
    else:
        print(f"\n✗✗✗ CHYBA! Nesprávný počet tabulek")
        return False


if __name__ == "__main__":
    print("FINÁLNÍ TEST PP TABLE REMOVAL")
    print("="*60)

    # Test 1
    result1 = test_header_only_removal()

    # Test 2
    result2 = test_mixed_tables()

    print("\n" + "="*60)
    print("VÝSLEDKY:")
    print(f"  Test 1 (pouze záhlaví): {'✓ PASS' if result1 else '✗ FAIL'}")
    print(f"  Test 2 (mix tabulek): {'✓ PASS' if result2 else '✗ FAIL'}")
    print("="*60)