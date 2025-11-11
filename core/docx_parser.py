"""
DOCX Parser - parsování Word dokumentů
"""
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
import unicodedata
import re


class DocxParser:
    """Parsování Word dokumentů pro získání tabulek"""

    # Vzory pro detekci sloupců (různé varianty názvů)
    COLUMN_PATTERNS = {
        "number": ["č.", "číslo", "number", "#", "pořadí", "por.", "c."],
        "operation": ["činnost", "operace", "operation", "popis", "práce", "aktivita", "cinnost"],
        "time_min": ["čas", "time", "min", "minut", "trvání", "doba", "cas"],
        "pieces_count": ["počet", "kusy", "pieces", "ks", "count", "množství", "pocet", "mnozstvi"]
    }

    @staticmethod
    def _normalize_column_name(name: str) -> str:
        """
        Normalizuje název sloupce pro matching.

        - Převede na lowercase
        - Odstraní diakritiku
        - Odstraní speciální znaky (závorky, čárky, atd.)
        - Redukuje mezery

        Args:
            name: Původní název sloupce

        Returns:
            Normalizovaný název
        """
        if not name:
            return ""

        # Lowercase
        name = name.lower()

        # Odstraň diakritiku (č→c, š→s, atd.)
        name = ''.join(
            c for c in unicodedata.normalize('NFD', name)
            if unicodedata.category(c) != 'Mn'
        )

        # Odstraň speciální znaky, nech jen písmena, čísla, mezery
        name = re.sub(r'[^a-z0-9\s]', '', name)

        # Redukuj vícenásobné mezery
        name = re.sub(r'\s+', ' ', name)

        return name.strip()

    @staticmethod
    def _detect_column_mapping(header_cells) -> Dict[str, Optional[int]]:
        """
        Detekuje mapování column_type → column_index podle názvů v headeru.

        Args:
            header_cells: Buňky z header řádku tabulky

        Returns:
            Dictionary: {"number": 0, "operation": 1, "time_min": 2, "pieces_count": None}
            None znamená, že sloupec nebyl nalezen
        """
        column_map = {
            "number": None,
            "operation": None,
            "time_min": None,
            "pieces_count": None
        }

        # Normalizuj názvy všech header sloupců
        header_names = [DocxParser._normalize_column_name(cell.text) for cell in header_cells]

        # Track které sloupce už byly přiřazeny (aby se neopakovaly)
        assigned_indices = set()

        # Pro každý typ sloupce zkus najít match v headeru
        # Prioritní pořadí: operation > time_min > pieces_count > number
        priority_order = ["operation", "time_min", "pieces_count", "number"]

        for col_type in priority_order:
            patterns = DocxParser.COLUMN_PATTERNS[col_type]

            for idx, header_name in enumerate(header_names):
                if idx in assigned_indices:
                    continue  # Tento sloupec už byl přiřazen

                # Zkontroluj jestli header obsahuje nějaký z patternů
                # WORD-BASED MATCHING: Rozdělím header na slova a testuji každé slovo
                header_words = header_name.split()
                match_found = False

                for pattern in patterns:
                    # NORMALIZUJ PATTERN (patterns mohou obsahovat diakritiku)
                    pattern_normalized = DocxParser._normalize_column_name(pattern)

                    # Přesný match celého headeru (rychlá cesta)
                    if pattern_normalized in header_name:
                        match_found = True
                        break

                    # Word-based matching: testuji každé slovo zvlášť
                    for word in header_words:
                        # Pro krátké zkratky (≤3 znaky): přesný match začátku
                        if len(pattern_normalized) <= 3 and word.startswith(pattern_normalized):
                            match_found = True
                            break

                        # Pro delší patterns (≥4 znaky): COMMON PREFIX matching (pokrývá české skloňování)
                        if len(pattern_normalized) >= 4 and len(word) >= 4:
                            # Pokud jsou slova podobné délky (rozdíl max 3 znaky)
                            if abs(len(word) - len(pattern_normalized)) <= 3:
                                # Porovnej common prefix (bez posledních 2 znaků kvůli koncovkám)
                                compare_len = min(len(word), len(pattern_normalized)) - 2
                                if compare_len > 0 and word[:compare_len] == pattern_normalized[:compare_len]:
                                    match_found = True
                                    break
                            # Pokud pattern je kratší, zkus standardní startswith
                            elif len(pattern_normalized) < len(word) and word.startswith(pattern_normalized):
                                match_found = True
                                break
                            # Pokud word je kratší, zkus opačně
                            elif len(word) < len(pattern_normalized) and pattern_normalized.startswith(word):
                                match_found = True
                                break

                    if match_found:
                        break

                if match_found:
                    column_map[col_type] = idx
                    assigned_indices.add(idx)
                    break

                if column_map[col_type] is not None:
                    break  # Našli jsme match pro tento typ, přejdi na další

        return column_map

    @staticmethod
    def _safe_get_cell_value(cells, column_index: Optional[int]) -> Optional[str]:
        """
        Bezpečně získá hodnotu buňky.

        Args:
            cells: Seznam buněk z řádku
            column_index: Index sloupce (může být None)

        Returns:
            Text z buňky nebo None pokud index neexistuje
        """
        if column_index is None or column_index >= len(cells):
            return None
        return cells[column_index].text.strip()

    @staticmethod
    def parse_time_schedule_table(docx_path: str) -> Dict[str, Any]:
        """
        Naparsuje tabulku "Časové rozložení pracovní směny" z Word dokumentu.

        FLEXIBILNÍ NAČÍTÁNÍ:
        - Detekuje sloupce podle názvů v headeru (fuzzy matching)
        - Funguje i když některé sloupce chybí (time_min, pieces_count)
        - Vyžaduje pouze sloupec "Činnost" (operation)
        - Fallback na default pořadí pokud header nelze detekovat

        Args:
            docx_path: Cesta k .docx souboru

        Returns:
            Dictionary s line1...line20 a total
        """
        if not docx_path:
            return DocxParser._get_empty_time_schedule()

        try:
            doc = Document(docx_path)

            # Hledáme druhou tabulku (index 1)
            # Tabulka 1 = informační (místo měření, datum, ...)
            # Tabulka 2 = časové rozložení
            if len(doc.tables) < 2:
                print(f"⚠ Varování: Dokument má méně než 2 tabulky")
                return DocxParser._get_empty_time_schedule()

            table = doc.tables[1]  # Druhá tabulka (index 1)

            if len(table.rows) == 0:
                print(f"⚠ Varování: Tabulka je prázdná")
                return DocxParser._get_empty_time_schedule()

            # KROK 1: Detekuj mapování sloupců z headeru
            header_cells = table.rows[0].cells
            column_map = DocxParser._detect_column_mapping(header_cells)

            # Fallback na default pořadí pokud žádný sloupec nebyl detekován
            if all(v is None for v in column_map.values()):
                print(f"⚠ Varování: Nepodařilo se detekovat názvy sloupců, používám default pořadí")
                print(f"   Předpokládám: [Č., Činnost, Čas, Počet kusů]")
                column_map = {
                    "number": 0,
                    "operation": 1,
                    "time_min": 2 if len(header_cells) > 2 else None,
                    "pieces_count": 3 if len(header_cells) > 3 else None
                }

            # Debug výpis detekovaných sloupců
            print(f"✓ Detekované sloupce v časovém snímku:")
            print(f"   • Činnost: sloupec {column_map['operation']}" if column_map['operation'] is not None else "   ⚠ Činnost: NENALEZENA")
            print(f"   • Čas: sloupec {column_map['time_min']}" if column_map['time_min'] is not None else "   ⚠ Čas: chybí (bude None)")
            print(f"   • Počet kusů: sloupec {column_map['pieces_count']}" if column_map['pieces_count'] is not None else "   ⚠ Počet kusů: chybí (bude None)")

            # Validace: operation MUSÍ být detekována
            if column_map['operation'] is None:
                print(f"❌ CHYBA: Sloupec 'Činnost' nebyl nalezen! Tabulka se nenačte.")
                return DocxParser._get_empty_time_schedule()

            # KROK 2: Načti data řádek po řádku
            result = {}
            total_row = None
            row_count = 0

            for i, row in enumerate(table.rows):
                if i == 0:  # Přeskoč header
                    continue

                cells = row.cells

                # Získej hodnoty podle detekovaného mapování
                number = DocxParser._safe_get_cell_value(cells, column_map['number'])
                operation = DocxParser._safe_get_cell_value(cells, column_map['operation'])
                time_text = DocxParser._safe_get_cell_value(cells, column_map['time_min'])
                pieces_text = DocxParser._safe_get_cell_value(cells, column_map['pieces_count'])

                # Detekuj řádek "Celkem"
                if (number and "celkem" in number.lower()) or (operation and "celkem" in operation.lower()):
                    total_row = {
                        "time_min": DocxParser._parse_number(time_text) if time_text else None,
                        "pieces_count": DocxParser._parse_number(pieces_text) if pieces_text else None
                    }
                    continue

                # VALIDACE: operation MUSÍ existovat
                if not operation:
                    continue  # Přeskoč prázdný řádek

                # Normální řádek
                row_count += 1
                if row_count > 20:  # Max 20 řádků
                    break

                result[f"line{row_count}"] = {
                    "number": number or "",
                    "operation": operation,
                    "time_min": DocxParser._parse_number(time_text) if time_text else None,
                    "pieces_count": DocxParser._parse_number(pieces_text) if pieces_text else None
                }

            # Doplň prázdné řádky do 20
            for i in range(row_count + 1, 21):
                result[f"line{i}"] = {
                    "number": "",
                    "operation": "",
                    "time_min": None,
                    "pieces_count": None
                }

            # Přidej total
            if total_row:
                result["total"] = total_row
            else:
                result["total"] = {"time_min": None, "pieces_count": None}

            print(f"✓ Načteno {row_count} řádků z časového snímku")

            return result

        except Exception as e:
            print(f"❌ Chyba při parsování Word dokumentu: {e}")
            import traceback
            traceback.print_exc()
            return DocxParser._get_empty_time_schedule()

    @staticmethod
    def _parse_number(text: str) -> Optional[int]:
        """
        Převede textový řetězec na číslo.

        Args:
            text: Textový řetězec (např. "415", "-", "")

        Returns:
            Číslo nebo None
        """
        text = text.strip()
        if not text or text == "-":
            return None

        try:
            return int(text)
        except ValueError:
            return None

    @staticmethod
    def _get_empty_time_schedule() -> Dict[str, Any]:
        """
        Vrátí prázdnou strukturu časového snímku.

        Returns:
            Dictionary s prázdnými line1...line20 a total
        """
        result = {}
        for i in range(1, 21):
            result[f"line{i}"] = {
                "number": "",
                "operation": "",
                "time_min": None,
                "pieces_count": None
            }
        result["total"] = {"time_min": None, "pieces_count": None}
        return result
