"""
Text Generator - generování podmínkových textů pro Word protokoly
"""
import math
from typing import Dict, Any, Optional


def _math_round(value: float) -> int:
    """
    Klasické matematické zaokrouhlování (ne Python banker's rounding).

    Args:
        value: Číslo k zaokrouhlení

    Returns:
        Zaokrouhlené celé číslo
    """
    return math.floor(value + 0.5)


def _find_in_table_W4_Y51(table: Dict[str, Dict], fmax_rounded: int) -> Optional[Dict[str, Any]]:
    """
    Najde řádek v table_W4_Y51, kde fmax odpovídá zaokrouhlené hodnotě.

    Args:
        table: table_W4_Y51 z results_data
        fmax_rounded: Zaokrouhlená hodnota Fmax

    Returns:
        Řádek z tabulky {"fmax": ..., "phk": ..., "lhk": ...} nebo None
    """
    # Projdi tabulku (skip header na indexu "1")
    for key, row in table.items():
        if key == "1":  # Skip header
            continue

        # Porovnej fmax
        if row.get("fmax") == fmax_rounded:
            return row

    return None


def _calculate_druhy_text_podminka_limit1(results_data: Dict[str, Any]) -> str:
    """
    Vypočítá text pro druhy_text_podminka_limit1 na základě překročení hygienických limitů.

    Args:
        results_data: Data z lsz_results.json

    Returns:
        Vygenerovaný text podle podmínek
    """
    # Definice textových variant
    TEXTS = {
        (1, 1): "překračují průměrné hygienické limity počtu pohybů pro naměřené vynakládané svalové síly extenzorů a flexorů předloktí pravé horní končetiny.",
        (1, 0): "překračují průměrné hygienické limity počtu pohybů pro naměřené vynakládané svalové síly extenzorů předloktí pravé horní končetiny. Pro flexory byl hygienický limit zachován.",
        (0, 1): "překračují průměrné hygienické limity počtu pohybů pro naměřené vynakládané svalové síly flexorů předloktí pravé horní končetiny. Pro extenzory byl hygienický limit zachován.",
        (0, 0): "nepřekračují průměrné hygienické limity počtu pohybů pro naměřené vynakládané svalové síly extenzorů a flexorů předloktí pravé horní končetiny."
    }

    # Načti data
    fmax_ext = results_data.get("Fmax_Phk_Extenzor")
    fmax_flex = results_data.get("Fmax_Phk_Flexor")
    phk_movements = results_data.get("phk_number_of_movements")
    table = results_data.get("table_W4_Y51", {})

    # Fallback pokud chybí data
    if fmax_ext is None or fmax_flex is None or phk_movements is None or not table:
        return TEXTS[(0, 0)]

    # VĚTEV 1: Extenzory
    fmax_ext_rounded = _math_round(fmax_ext)
    row_ext = _find_in_table_W4_Y51(table, fmax_ext_rounded)

    if row_ext is None:
        extenzor_flag = 0
    else:
        limit_ext = row_ext.get("phk", 0)
        extenzor_flag = 1 if phk_movements > limit_ext else 0

    # VĚTEV 2: Flexory
    fmax_flex_rounded = _math_round(fmax_flex)
    row_flex = _find_in_table_W4_Y51(table, fmax_flex_rounded)

    if row_flex is None:
        flexor_flag = 0
    else:
        limit_flex = row_flex.get("phk", 0)
        flexor_flag = 1 if phk_movements > limit_flex else 0

    # Vybrat text podle kombinace
    return TEXTS[(extenzor_flag, flexor_flag)]


def _calculate_treti_text_podminka_limit1(results_data: Dict[str, Any]) -> str:
    """
    Vypočítá text pro treti_text_podminka_limit1 na základě překročení hygienických limitů pro LHK.

    Args:
        results_data: Data z lsz_results.json

    Returns:
        Vygenerovaný text podle podmínek
    """
    # Definice textových variant
    TEXTS = {
        (1, 1): "překračují průměrné hygienické limity počtu pohybů pro naměřené vynakládané svalové síly extenzorů a flexorů předloktí levé horní končetiny.",
        (1, 0): "překračují průměrné hygienické limity počtu pohybů pro naměřené vynakládané svalové síly extenzorů předloktí levé horní končetiny. Pro flexory byl hygienický limit zachován.",
        (0, 1): "překračují průměrné hygienické limity počtu pohybů pro naměřené vynakládané svalové síly flexorů předloktí levé horní končetiny. Pro extenzory byl hygienický limit zachován.",
        (0, 0): "nepřekračují průměrné hygienické limity počtu pohybů pro naměřené vynakládané svalové síly extenzorů a flexorů předloktí levé horní končetiny."
    }

    # Načti data
    fmax_ext = results_data.get("Fmax_Lhk_Extenzor")
    fmax_flex = results_data.get("Fmax_Lhk_Flexor")
    lhk_movements = results_data.get("lhk_number_of_movements")
    table = results_data.get("table_W4_Y51", {})

    # Fallback pokud chybí data
    if fmax_ext is None or fmax_flex is None or lhk_movements is None or not table:
        return TEXTS[(0, 0)]

    # VĚTEV 1: Extenzory
    fmax_ext_rounded = _math_round(fmax_ext)
    row_ext = _find_in_table_W4_Y51(table, fmax_ext_rounded)

    if row_ext is None:
        extenzor_flag = 0
    else:
        limit_ext = row_ext.get("lhk", 0)
        extenzor_flag = 1 if lhk_movements > limit_ext else 0

    # VĚTEV 2: Flexory
    fmax_flex_rounded = _math_round(fmax_flex)
    row_flex = _find_in_table_W4_Y51(table, fmax_flex_rounded)

    if row_flex is None:
        flexor_flag = 0
    else:
        limit_flex = row_flex.get("lhk", 0)
        flexor_flag = 1 if lhk_movements > limit_flex else 0

    # Vybrat text podle kombinace
    return TEXTS[(extenzor_flag, flexor_flag)]


def _calculate_ctvrty_text_podminka(results_data: Dict[str, Any], worker_count: int = 2) -> str:
    """
    Vypočítá text pro ctvrty_text_podminka na základě rozložení svalových sil.

    Args:
        results_data: Data z lsz_results.json
        worker_count: Počet pracovníků (1 nebo 2)

    Returns:
        "nejsou" | "ojediněle" | "pravidelně"
    """
    # Načti tabulku force_distribution
    table = results_data.get("table_force_distribution", {})

    # Vyber správný řádek Celkem podle počtu pracovníků
    # Pro 1 pracovníka: řádek 7 (Celkem 1. měřené osoby)
    # Pro 2 pracovníky: řádek 21 (časově vážený průměr)
    celkem_row_key = "7" if worker_count == 1 else "21"
    row = table.get(celkem_row_key, {})

    # Pokud řádek neexistuje, fallback na "nejsou"
    if not row:
        return "nejsou"

    # Načti všechny 8 hodnot
    values = [
        row.get("force_55_70_phk_extenzory"),
        row.get("force_55_70_phk_flexory"),
        row.get("force_55_70_lhk_extenzory"),
        row.get("force_55_70_lhk_flexory"),
        row.get("force_over_70_phk_extenzory"),
        row.get("force_over_70_phk_flexory"),
        row.get("force_over_70_lhk_extenzory"),
        row.get("force_over_70_lhk_flexory")
    ]

    # Konvertuj None na 0 (fallback)
    values = [v if v is not None else 0 for v in values]

    # Podmínky (v tomto pořadí!)
    if all(v == 0 for v in values):
        return "nejsou"
    elif all(v >= 1 for v in values):
        return "pravidelně"
    else:
        return "ojediněle"


def _calculate_paty_text_podminka(results_data: Dict[str, Any], worker_count: int = 2) -> str:
    """
    Vypočítá text pro paty_text_podminka na základě nadlimitních svalových sil (nad 70% Fmax).

    Args:
        results_data: Data z lsz_results.json
        worker_count: Počet pracovníků (1 nebo 2)

    Returns:
        Text podle kombinace 4 hodnot force_over_70
    """
    # Načti tabulku force_distribution
    table = results_data.get("table_force_distribution", {})

    # Vyber správný řádek Celkem podle počtu pracovníků
    # Pro 1 pracovníka: řádek 7 (Celkem 1. měřené osoby)
    # Pro 2 pracovníky: řádek 21 (časově vážený průměr)
    celkem_row_key = "7" if worker_count == 1 else "21"
    row = table.get(celkem_row_key, {})

    # Pokud řádek neexistuje, fallback
    if not row:
        return "Při provádění práce nedochází k vynakládání nadlimitních svalových sil u všech měřených svalových skupin rukou a předloktí (nad 70 % Fmax)."

    # Načti 4 hodnoty (nadlimitní síly nad 70%)
    values = [
        row.get("force_over_70_phk_extenzory", 0),
        row.get("force_over_70_phk_flexory", 0),
        row.get("force_over_70_lhk_extenzory", 0),
        row.get("force_over_70_lhk_flexory", 0)
    ]

    # Konvertuj None na 0 (fallback)
    values = [v if v is not None else 0 for v in values]

    # Konvertuj na flags (>= 1 → 1, jinak 0)
    flags = tuple(1 if v >= 1 else 0 for v in values)

    # KOMPLETNÍ MAPPING VŠECH 16 KOMBINACÍ (2^4 = 16)
    PATTERNS = {
        (0, 0, 0, 0): "Při provádění práce nedochází k vynakládání nadlimitních svalových sil u všech měřených svalových skupin rukou a předloktí (nad 70 % Fmax).",

        (0, 0, 0, 1): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u flexorů LHK (nad 70 % Fmax).",
        (0, 0, 1, 0): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u extenzorů LHK (nad 70 % Fmax).",
        (0, 0, 1, 1): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u měřených svalových skupin levé ruky a předloktí (nad 70 % Fmax).",

        (0, 1, 0, 0): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u flexorů PHK (nad 70 % Fmax).",
        (0, 1, 0, 1): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u měřených flexorových svalových skupin rukou a předloktí (nad 70 % Fmax).",
        (0, 1, 1, 0): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u flexorů PHK a extenzorů LHK (nad 70 % Fmax).",
        (0, 1, 1, 1): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u flexorů PHK, extenzorů LHK a flexorů LHK (nad 70 % Fmax).",

        (1, 0, 0, 0): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u extenzorů PHK (nad 70 % Fmax).",
        (1, 0, 0, 1): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u extenzorů PHK a flexorů LHK (nad 70 % Fmax).",
        (1, 0, 1, 0): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u měřených extenzorových svalových skupin rukou a předloktí (nad 70 % Fmax).",
        (1, 0, 1, 1): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u extenzorů PHK, extenzorů LHK a flexorů LHK (nad 70 % Fmax).",

        (1, 1, 0, 0): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u měřených svalových skupin pravé ruky a předloktí (nad 70 % Fmax).",
        (1, 1, 0, 1): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u extenzorů PHK, flexorů PHK a flexorů LHK (nad 70 % Fmax).",
        (1, 1, 1, 0): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u extenzorů PHK, flexorů PHK a extenzorů LHK (nad 70 % Fmax).",
        (1, 1, 1, 1): "Při provádění práce dochází k vynakládání nadlimitních svalových sil u extenzorů PHK, flexorů PHK, extenzorů LHK a flexorů LHK (nad 70 % Fmax)."
    }

    return PATTERNS.get(flags, "Při provádění práce nedochází k vynakládání nadlimitních svalových sil u všech měřených svalových skupin rukou a předloktí (nad 70 % Fmax).")  # fallback pokud pattern neexistuje


def _calculate_sesty_text_podminka(results_data: Dict[str, Any], worker_count: int = 2) -> str:
    """
    Vypočítá text pro sesty_text_podminka na základě hodnot > 100.

    POUZE pro nadlimitní síly nad 70% Fmax (force_over_70_*).

    Args:
        results_data: Data z lsz_results.json
        worker_count: Počet pracovníků (1 nebo 2)

    Returns:
        Celá věta o pravidelnosti vynakládání nadlimitních svalových sil
    """
    # Načti tabulku force_distribution
    table = results_data.get("table_force_distribution", {})

    # Vyber správný řádek Celkem podle počtu pracovníků
    # Pro 1 pracovníka: řádek 7 (Celkem 1. měřené osoby)
    # Pro 2 pracovníky: řádek 21 (časově vážený průměr)
    celkem_row_key = "7" if worker_count == 1 else "21"
    row = table.get(celkem_row_key, {})

    # Pokud řádek neexistuje, fallback
    if not row:
        return "Vynakládání nadlimitních svalových sil není pravidelnou součástí výkonu prováděné práce."

    # Načti POUZE 4 hodnoty force_over_70 (nadlimitní síly nad 70%)
    values = [
        row.get("force_over_70_phk_extenzory", 0),
        row.get("force_over_70_phk_flexory", 0),
        row.get("force_over_70_lhk_extenzory", 0),
        row.get("force_over_70_lhk_flexory", 0)
    ]

    # Konvertuj None na 0 (fallback)
    values = [v if v is not None else 0 for v in values]

    # Pokud jakákoliv hodnota > 100
    if any(v > 100 for v in values):
        return "Vynakládání nadlimitních svalových sil je pravidelnou součástí výkonu prováděné práce."
    else:
        return "Vynakládání nadlimitních svalových sil není pravidelnou součástí výkonu prováděné práce."


def _calculate_sedmy_text_podminka(measurement_data: Dict[str, Any], results_data: Dict[str, Any], worker_count: int = 2) -> str:
    """
    Vypočítá text pro sedmy_text_podminka na základě překročení limitu pro velké svalové síly (55-70% Fmax).

    Logika:
    - Načti délku směny (work_duration v minutách)
    - Vypočti limit = (work_duration / 2) + 360
    - Sečti 4 hodnoty force_55_70_* z řádku Celkem
    - Pokud součet > limit → "překračuje u měřených svalových skupin..."
    - Jinak → "nepřekračuje u žádné z měřených svalových skupin..."

    Args:
        measurement_data: Data z measurement_data.json
        results_data: Data z lsz_results.json
        worker_count: Počet pracovníků (1 nebo 2)

    Returns:
        Celá věta o překročení/nepřekročení hygienického limitu
    """
    # Načti délku směny v minutách
    work_duration = measurement_data.get("section4_worker_a", {}).get("work_duration")

    # Pokud work_duration chybí, fallback
    if work_duration is None:
        return "Celosměnový počet těchto sil nepřekračuje u žádné z měřených svalových skupin rukou a předloktí daný hygienický limit."

    # Konvertuj na číslo (může být string z JSON)
    try:
        work_duration = float(work_duration)
    except (ValueError, TypeError):
        # Pokud konverze selže, fallback
        return "Celosměnový počet těchto sil nepřekračuje u žádné z měřených svalových skupin rukou a předloktí daný hygienický limit."

    # Vypočti limit podle vzorce: limit = (work_duration / 2) + 360
    limit = (work_duration / 2) + 360

    # Načti tabulku force_distribution
    table = results_data.get("table_force_distribution", {})

    # Vyber správný řádek Celkem podle počtu pracovníků
    # Pro 1 pracovníka: řádek 7 (Celkem 1. měřené osoby)
    # Pro 2 pracovníky: řádek 21 (časově vážený průměr)
    celkem_row_key = "7" if worker_count == 1 else "21"
    row = table.get(celkem_row_key, {})

    # Pokud řádek neexistuje, fallback
    if not row:
        return "Celosměnový počet těchto sil nepřekračuje u žádné z měřených svalových skupin rukou a předloktí daný hygienický limit."

    # Načti 4 hodnoty force_55_70_* (velké svalové síly 55-70% Fmax)
    suma_55_70 = (
        row.get("force_55_70_phk_extenzory", 0) +
        row.get("force_55_70_phk_flexory", 0) +
        row.get("force_55_70_lhk_extenzory", 0) +
        row.get("force_55_70_lhk_flexory", 0)
    )

    # Porovnej s limitem
    if suma_55_70 > limit:
        return "Celosměnový počet těchto sil překračuje u měřených svalových skupin rukou a předloktí daný hygienický limit."
    else:
        return "Celosměnový počet těchto sil nepřekračuje u žádné z měřených svalových skupin rukou a předloktí daný hygienický limit."


def _calculate_osmy_text_podminka(results_data: Dict[str, Any], worker_count: int = 2) -> str:
    """
    Vypočítá text pro osmy_text_podminka - seznam činností s force_over_70 > 100.

    Logika:
    - Pokud sesty_text_podminka = "není" → prázdný string
    - Pokud sesty_text_podminka = "je":
      - Projdi všechny řádky table_force_distribution (kromě "Celkem")
      - Pro každý řádek zkontroluj 4 hodnoty force_over_70_*
      - Pokud jakákoliv > 100 → přidej activity do seznamu
      - Vrať seznam oddělený čárkami

    Args:
        results_data: Data z lsz_results.json
        worker_count: Počet pracovníků (1 nebo 2)

    Returns:
        Seznam činností oddělený čárkami, nebo prázdný string
    """
    # Nejdřív zkontroluj sesty_text_podminka
    sesty = _calculate_sesty_text_podminka(results_data, worker_count)

    # Pokud sesty = "není", vrať prázdný string
    if sesty == "není":
        return ""

    # Sesty = "je", najdi konkrétní činnosti
    table = results_data.get("table_force_distribution", {})

    if not table:
        return ""

    activities = []

    # Projdi všechny řádky kromě "Celkem" řádků
    for key, row in table.items():
        # Načti activity
        activity = row.get("activity")
        if not activity:
            continue

        # Skip všechny řádky s "Celkem" (pro 1 pracovníka = 1 řádek, pro 2 = 3 řádky)
        if activity == "Celkem":
            continue

        # Zkontroluj 4 hodnoty force_over_70
        force_values = [
            row.get("force_over_70_phk_extenzory", 0),
            row.get("force_over_70_phk_flexory", 0),
            row.get("force_over_70_lhk_extenzory", 0),
            row.get("force_over_70_lhk_flexory", 0)
        ]

        # Konvertuj None na 0 (fallback pro null hodnoty z JSON)
        force_values = [v if v is not None else 0 for v in force_values]

        # Pokud jakákoliv > 100, přidej activity
        if any(v > 100 for v in force_values):
            activities.append(activity)

    # Vrať seznam oddělený čárkami (nebo prázdný string)
    if activities:
        return ", ".join(activities)
    else:
        return ""


def _calculate_devata_text_podminka(results_data: Dict[str, Any]) -> Dict[str, str]:
    """
    Vypočítá hodnoty pro tabulku hygienických limitů (devátá podmínka).

    Používá stejnou logiku jako druhá a třetí podmínka, ale vrací strukturovaná data
    pro všechny 4 svalové skupiny: PHK extenzory, PHK flexory, LHK extenzory, LHK flexory.

    Args:
        results_data: Data z lsz_results.json

    Returns:
        Dictionary s hodnotami pro tabulku:
        {
            "phk_extenzory": "Nad limitem" | "Pod limitem",
            "phk_flexory": "Nad limitem" | "Pod limitem",
            "lhk_extenzory": "Nad limitem" | "Pod limitem",
            "lhk_flexory": "Nad limitem" | "Pod limitem"
        }
    """
    # Default fallback hodnoty (všechno pod limitem)
    result = {
        "phk_extenzory": "Pod limitem",
        "phk_flexory": "Pod limitem",
        "lhk_extenzory": "Pod limitem",
        "lhk_flexory": "Pod limitem"
    }

    # Načti tabulku W4_Y51 a počty pohybů
    table = results_data.get("table_W4_Y51", {})
    if not table:
        return result

    phk_movements = results_data.get("phk_number_of_movements")
    lhk_movements = results_data.get("lhk_number_of_movements")

    # === PHK EXTENZORY ===
    fmax_phk_ext = results_data.get("Fmax_Phk_Extenzor")

    if fmax_phk_ext is not None and phk_movements is not None:
        fmax_rounded = _math_round(fmax_phk_ext)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            limit = row.get("phk", 0)
            if phk_movements > limit:
                result["phk_extenzory"] = "Nad limitem"

    # === PHK FLEXORY ===
    fmax_phk_flex = results_data.get("Fmax_Phk_Flexor")

    if fmax_phk_flex is not None and phk_movements is not None:
        fmax_rounded = _math_round(fmax_phk_flex)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            limit = row.get("phk", 0)
            if phk_movements > limit:
                result["phk_flexory"] = "Nad limitem"

    # === LHK EXTENZORY ===
    fmax_lhk_ext = results_data.get("Fmax_Lhk_Extenzor")

    if fmax_lhk_ext is not None and lhk_movements is not None:
        fmax_rounded = _math_round(fmax_lhk_ext)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            limit = row.get("lhk", 0)
            if lhk_movements > limit:
                result["lhk_extenzory"] = "Nad limitem"

    # === LHK FLEXORY ===
    fmax_lhk_flex = results_data.get("Fmax_Lhk_Flexor")

    if fmax_lhk_flex is not None and lhk_movements is not None:
        fmax_rounded = _math_round(fmax_lhk_flex)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            limit = row.get("lhk", 0)
            if lhk_movements > limit:
                result["lhk_flexory"] = "Nad limitem"

    return result


def _get_shift_duration_text(work_duration_minutes: float) -> str:
    """
    Převede délku směny v minutách na slovní vyjádření v češtině.

    Args:
        work_duration_minutes: Délka směny v minutách (240-840)

    Returns:
        Slovní vyjádření ve tvaru "(osmihodinovou)" nebo "(sedmi a půl hodinovou)"
    """
    # Mapování minut → text (4h až 14h, po půlhodinových incrementech)
    SHIFT_DURATION_MAP = {
        240: "(čtyřhodinovou)",
        270: "(čtyři a půl hodinovou)",
        300: "(pětihodinovou)",
        330: "(pět a půl hodinovou)",
        360: "(šestihodinovou)",
        390: "(šest a půl hodinovou)",
        420: "(sedmihodinovou)",
        450: "(sedmi a půl hodinovou)",
        480: "(osmihodinovou)",
        510: "(osmi a půl hodinovou)",
        540: "(devítihodinovou)",
        570: "(devíti a půl hodinovou)",
        600: "(desetihodinovou)",
        630: "(deseti a půl hodinovou)",
        660: "(jedenáctihodinovou)",
        690: "(jedenácti a půl hodinovou)",
        720: "(dvanáctihodinovou)",
        750: "(dvanácti a půl hodinovou)",
        780: "(třináctihodinovou)",
        810: "(třinácti a půl hodinovou)",
        840: "(čtrnáctihodinovou)"
    }

    # Zaokrouhli na nejbližších 30 minut
    rounded_minutes = round(work_duration_minutes / 30) * 30

    # Vrať text nebo fallback
    return SHIFT_DURATION_MAP.get(int(rounded_minutes), "(osmihodinovou)")


def _calculate_hygiene_limits(results_data: Dict[str, Any]) -> Dict[str, int]:
    """
    Vypočítá hygienické limity pro všechny 4 Fmax hodnoty z table_W4_Y51.

    Args:
        results_data: Data z lsz_results.json

    Returns:
        Dictionary s limity: {"phk_extenzor": 12500, "phk_flexor": 13000, ...}
    """
    table = results_data.get("table_W4_Y51", {})
    result = {}

    # PHK Extenzory
    fmax_phk_ext = results_data.get("Fmax_Phk_Extenzor")
    if fmax_phk_ext is not None:
        fmax_rounded = _math_round(fmax_phk_ext)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            result["phk_extenzor"] = row.get("phk", 0)

    # PHK Flexory
    fmax_phk_flex = results_data.get("Fmax_Phk_Flexor")
    if fmax_phk_flex is not None:
        fmax_rounded = _math_round(fmax_phk_flex)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            result["phk_flexor"] = row.get("phk", 0)

    # LHK Extenzory
    fmax_lhk_ext = results_data.get("Fmax_Lhk_Extenzor")
    if fmax_lhk_ext is not None:
        fmax_rounded = _math_round(fmax_lhk_ext)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            result["lhk_extenzor"] = row.get("lhk", 0)

    # LHK Flexory
    fmax_lhk_flex = results_data.get("Fmax_Lhk_Flexor")
    if fmax_lhk_flex is not None:
        fmax_rounded = _math_round(fmax_lhk_flex)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            result["lhk_flexor"] = row.get("lhk", 0)

    return result


def _calculate_desata_text_podminka(devata_data: Dict[str, str]) -> str:
    """
    Vypočítá text pro desatou podmínku - celkové překročení limitu.

    Zkontroluje všechny 4 hodnoty z devátého textu (phk_extenzory, phk_flexory,
    lhk_extenzory, lhk_flexory). Pokud je ALESPOŇ JEDNA "Nad limitem",
    vrátí "Nad limitem", jinak "Pod limitem".

    Args:
        devata_data: Výstup z _calculate_devata_text_podminka
                    {"phk_extenzory": "Nad limitem" | "Pod limitem", ...}

    Returns:
        "Nad limitem" | "Pod limitem"
    """
    # Zkontroluj všechny 4 hodnoty
    values = [
        devata_data.get("phk_extenzory", "Pod limitem"),
        devata_data.get("phk_flexory", "Pod limitem"),
        devata_data.get("lhk_extenzory", "Pod limitem"),
        devata_data.get("lhk_flexory", "Pod limitem")
    ]

    # Pokud je alespoň jedna hodnota "Nad limitem", vrať "Nad limitem"
    if any(v == "Nad limitem" for v in values):
        return "Nad limitem"
    else:
        return "Pod limitem"


def _calculate_jedenacta_text_podminka(results_data: Dict[str, Any]) -> str:
    """
    Vypočítá text pro jedenáctou podmínku - hierarchické vyhodnocení zatížení všech svalových skupin.

    Vyhodnotí všechny 4 svalové skupiny (PHK extenzory, PHK flexory, LHK extenzory, LHK flexory)
    vzhledem k jejich hygienickým limitům.

    HIERARCHIE (priorita od nejvyšší):
    3. Alespoň jeden sval překračuje hygienický limit (počet pohybů > limit)
    2. Alespoň jeden sval je nad 1/3 hygienického limitu, ale žádný nepřekračuje celý limit
    1. Žádný sval není nad 1/3 hygienického limitu

    Args:
        results_data: Data z lsz_results.json

    Returns:
        "1" - žádný sval není nad 1/3 hygienického limitu
        "2" - alespoň jeden sval je nad 1/3 limitu, ale žádný nepřekračuje limit
        "3" - alespoň jeden sval překračuje hygienický limit
    """
    table = results_data.get("table_W4_Y51", {})
    phk_movements = results_data.get("phk_number_of_movements")
    lhk_movements = results_data.get("lhk_number_of_movements")

    # Fallback pokud chybí data
    if not table or phk_movements is None or lhk_movements is None:
        return "1"

    # Flagy pro vyhodnocení
    over_limit = False  # Alespoň jeden sval > limit
    over_one_third = False  # Alespoň jeden sval > 1/3 limitu

    # === PHK EXTENZORY ===
    fmax_phk_ext = results_data.get("Fmax_Phk_Extenzor")
    if fmax_phk_ext is not None:
        fmax_rounded = _math_round(fmax_phk_ext)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            limit = row.get("phk", 0)
            one_third_limit = limit / 3

            if phk_movements > limit:
                over_limit = True
            elif phk_movements > one_third_limit:
                over_one_third = True

    # === PHK FLEXORY ===
    fmax_phk_flex = results_data.get("Fmax_Phk_Flexor")
    if fmax_phk_flex is not None:
        fmax_rounded = _math_round(fmax_phk_flex)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            limit = row.get("phk", 0)
            one_third_limit = limit / 3

            if phk_movements > limit:
                over_limit = True
            elif phk_movements > one_third_limit:
                over_one_third = True

    # === LHK EXTENZORY ===
    fmax_lhk_ext = results_data.get("Fmax_Lhk_Extenzor")
    if fmax_lhk_ext is not None:
        fmax_rounded = _math_round(fmax_lhk_ext)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            limit = row.get("lhk", 0)
            one_third_limit = limit / 3

            if lhk_movements > limit:
                over_limit = True
            elif lhk_movements > one_third_limit:
                over_one_third = True

    # === LHK FLEXORY ===
    fmax_lhk_flex = results_data.get("Fmax_Lhk_Flexor")
    if fmax_lhk_flex is not None:
        fmax_rounded = _math_round(fmax_lhk_flex)
        row = _find_in_table_W4_Y51(table, fmax_rounded)
        if row is not None:
            limit = row.get("lhk", 0)
            one_third_limit = limit / 3

            if lhk_movements > limit:
                over_limit = True
            elif lhk_movements > one_third_limit:
                over_one_third = True

    # Hierarchické vyhodnocení (priorita: 3 > 2 > 1)
    if over_limit:
        return "3"
    elif over_one_third:
        return "2"
    else:
        return "1"


def _calculate_pp_work_shift_coefficient(work_duration_min: int) -> float:
    """
    Vypočítá koeficient přepočtu pro nestandardní délky směn podle NV 361/2007, Tabulka č.8.

    Standardní směna: 480 min → koeficient = 1.0
    Každých 30 min navíc → +0.025 (max 1.2 při >689 min)
    Každých 30 min méně → -0.025 (min 0.8 při <270 min)

    Args:
        work_duration_min: Délka směny v minutách

    Returns:
        Koeficient přepočtu (0.8 - 1.2)

    Příklady:
        >>> _calculate_pp_work_shift_coefficient(480)
        1.0
        >>> _calculate_pp_work_shift_coefficient(510)  # +30 min
        1.05
        >>> _calculate_pp_work_shift_coefficient(450)  # -30 min
        0.975
    """
    if work_duration_min == 480:
        return 1.0

    difference = work_duration_min - 480

    # Vypočítej interval podle tabulky:
    # První rozsah (481-509 nebo 450-479) = 1-29 min rozdílu
    # Další rozsahy po 30 minutách (510-539, 420-449, atd.)
    if difference > 0:  # Delší směna
        if difference < 30:
            intervals = 1  # 481-509
        else:
            intervals = ((difference - 30) // 30) + 2  # 510-539, 540-569, ...
        coef = 1.0 + (intervals * 0.025)
        return min(coef, 1.2)  # Max cap na 1.2
    else:  # Kratší směna
        abs_diff = abs(difference)
        # Pro kratší směny jsou všechny rozsahy 30 minut (symetrické)
        intervals = ((abs_diff - 1) // 30) + 1  # 1-30→1, 31-60→2, 61-90→3, ...
        coef = 1.0 - (intervals * 0.025)
        return max(coef, 0.8)  # Min floor na 0.8


def _calculate_pp_category_limits(coefficient: float) -> Dict[str, Any]:
    """
    Vypočítá limitní hodnoty pro kategorie pracovních poloh (PP a N) s aplikovaným koeficientem.

    Base limity pro směnu 480 min:
    - PP: kategorie 1 (<100 min), kategorie 2 (101-160 min), kategorie 3 (>160 min)
    - N:  kategorie 1 (<20 min),  kategorie 2 (21-30 min),   kategorie 3 (>30 min)

    Args:
        coefficient: Koeficient přepočtu (z _calculate_pp_work_shift_coefficient)

    Returns:
        Dictionary s limitními hodnotami (celá čísla minut):
        {
            "pp_kat1_min": 0,
            "pp_kat1_max": 100,  # × coefficient
            "pp_kat2_min": 101,
            "pp_kat2_max": 160,  # × coefficient
            "pp_kat3_min": 161,
            "n_kat1_min": 0,
            "n_kat1_max": 20,    # × coefficient
            "n_kat2_min": 21,
            "n_kat2_max": 30,    # × coefficient
            "n_kat3_min": 31
        }

    Příklad:
        >>> _calculate_pp_category_limits(1.05)  # Směna 510 min
        {"pp_kat1_max": 105, "pp_kat2_max": 168, ...}
    """
    # Base limity pro standardní směnu (480 min)
    pp_base = [100, 160]
    n_base = [20, 30]

    # Aplikovat koeficient a zaokrouhlit matematicky
    pp_limit1 = _math_round(pp_base[0] * coefficient)
    pp_limit2 = _math_round(pp_base[1] * coefficient)
    n_limit1 = _math_round(n_base[0] * coefficient)
    n_limit2 = _math_round(n_base[1] * coefficient)

    return {
        # PP kategorie
        "pp_kat1_min": 0,
        "pp_kat1_max": pp_limit1,
        "pp_kat2_min": pp_limit1 + 1,
        "pp_kat2_max": pp_limit2,
        "pp_kat3_min": pp_limit2 + 1,
        # Kategorie 3 nemá max (> pp_kat3_min)

        # N kategorie
        "n_kat1_min": 0,
        "n_kat1_max": n_limit1,
        "n_kat2_min": n_limit1 + 1,
        "n_kat2_max": n_limit2,
        "n_kat3_min": n_limit2 + 1,
        # Kategorie 3 nemá max (> n_kat3_min)
    }


def _detect_czech_noun_gender(nazev_polohy: str) -> str:
    """
    Detekuje gramatický rod prvního podstatného jména v názvu pracovní polohy.

    Kombinuje slovník známých slov + pattern matching koncovek.

    Args:
        nazev_polohy: Název pracovní polohy (např. "Předklon trupu větší než 60°")

    Returns:
        "M" (mužský), "F" (ženský), nebo "N" (střední rod)
    """
    # Tokenizace - rozdělení podle mezer
    tokens = nazev_polohy.strip().split()
    if not tokens:
        return "N"  # Fallback střední rod

    # Adjektiva která přeskakujeme (hledáme první podstatné jméno)
    skip_adjectives = {
        "výrazný", "výrazná", "výrazné",
        "nevhodná", "nevhodný", "nevhodné",
        "extrémní", "extrémná", "extrémné",
        "vnitřní", "vnitřná", "vnitřné",
        "zevní", "zevná", "zevné"
    }

    # Najdi první podstatné jméno (ne adjektivum)
    first_noun = None
    for token in tokens:
        token_lower = token.lower().strip(",.;:")
        if token_lower not in skip_adjectives:
            first_noun = token_lower
            break

    if not first_noun:
        return "N"  # Fallback

    # Slovník známých slov (priorita před pattern matching)
    GENDER_DICT = {
        # Mužský rod
        "předklon": "M",
        "úklon": "M",
        "záklon": "M",
        "dřep": "M",
        "pohyb": "M",

        # Ženský rod
        "rotace": "F",
        "flexe": "F",
        "exkurze": "F",
        "poloha": "F",
        "polohy": "F",
        "práce": "F",
        "hlava": "F",

        # Střední rod
        "pootočení": "N",
        "ohnutí": "N",
        "zapažení": "N",
        "vzpažení": "N",
        "vleže": "N",
        "kleče": "N"
    }

    # Zkus slovník
    if first_noun in GENDER_DICT:
        return GENDER_DICT[first_noun]

    # Pattern matching podle koncovek
    if first_noun.endswith(("klon", "on")):
        return "M"  # předklon, úklon, záklon
    elif first_noun.endswith(("ace", "e", "a")):
        return "F"  # rotace, flexe, poloha, práce
    elif first_noun.endswith(("ení", "í")):
        return "N"  # pootočení, ohnutí, zapažení

    # Fallback - střední rod (nejbezpečnější)
    return "N"


def _decline_czech_adjective(base_adjective: str, gender: str) -> str:
    """
    Skloňuje české adjektivum podle gramatického rodu.

    Args:
        base_adjective: "Dynamická" nebo "Statická"
        gender: "M" (mužský), "F" (ženský), "N" (střední)

    Returns:
        Skloňované adjektivum v lowercase

    Examples:
        >>> _decline_czech_adjective("Dynamická", "M")
        "dynamický"
        >>> _decline_czech_adjective("Statická", "F")
        "statická"
        >>> _decline_czech_adjective("Dynamická", "N")
        "dynamické"
    """
    adj_lower = base_adjective.lower().strip()

    # Detekuj typ adjektiva
    is_dynamic = "dynamick" in adj_lower
    is_static = "statick" in adj_lower

    if is_dynamic:
        if gender == "M":
            return "dynamický"
        elif gender == "F":
            return "dynamická"
        else:  # N
            return "dynamické"
    elif is_static:
        if gender == "M":
            return "statický"
        elif gender == "F":
            return "statická"
        else:  # N
            return "statické"
    else:
        # Neznámé adjektivum - vrať jak je
        return adj_lower


def _format_pp_position_name(typ_svalove_prace: str, nazev_polohy: str) -> str:
    """
    Formátuje název polohy pro výstupní text s GRAMATICKY SPRÁVNÝM skloňováním.

    Detekuje gramatický rod prvního podstatného jména a správně skloňuje
    adjektivum "dynamická/statická" podle něj.

    Args:
        typ_svalove_prace: "Dynamická" nebo "Statická"
        nazev_polohy: "Předklon trupu..." atd.

    Returns:
        Formátovaný string s GRAMATICKY SPRÁVNÝM skloňováním:
        - "dynamický předklon trupu" (mužský rod)
        - "dynamická rotace hlavy" (ženský rod)
        - "dynamické pootočení" (střední rod)

    Examples:
        >>> _format_pp_position_name("Dynamická", "Předklon trupu větší než 60°")
        "dynamický předklon trupu větší než 60°"
        >>> _format_pp_position_name("Statická", "Rotace hlavy větší než 15°")
        "statická rotace hlavy větší než 15°"
    """
    # Detekuj gramatický rod prvního podstatného jména
    gender = _detect_czech_noun_gender(nazev_polohy)

    # Skloňuj adjektivum podle rodu
    if typ_svalove_prace and isinstance(typ_svalove_prace, str):
        typ_formatted = _decline_czech_adjective(typ_svalove_prace, gender)
    else:
        typ_formatted = ""

    # Převést název polohy na lowercase (první písmeno)
    if nazev_polohy and isinstance(nazev_polohy, str):
        nazev_lower = nazev_polohy[0].lower() + nazev_polohy[1:] if len(nazev_polohy) > 0 else ""
    else:
        nazev_lower = ""

    # Spojit
    if typ_formatted and nazev_lower:
        return f"{typ_formatted} {nazev_lower}"
    elif nazev_lower:
        return nazev_lower
    else:
        return "neznámá poloha"


def _is_single_worker_protocol(measurement_data: Dict[str, Any]) -> bool:
    """
    Detekuje, jestli jde o protokol pro jednoho pracovníka.

    Kontroluje, jestli section5_worker_b má vyplněné full_name.
    Pokud je full_name prázdný string nebo whitespace, jde o single-worker protokol.

    Args:
        measurement_data: Data z measurement_data.json

    Returns:
        True pokud je single-worker protokol (worker_b nemá jméno), jinak False

    Examples:
        >>> # Single worker (prázdné jméno)
        >>> data = {"section5_worker_b": {"full_name": ""}}
        >>> _is_single_worker_protocol(data)
        True

        >>> # Two workers (vyplněné jméno)
        >>> data = {"section5_worker_b": {"full_name": "Jan Novák"}}
        >>> _is_single_worker_protocol(data)
        False
    """
    worker_b = measurement_data.get("section5_worker_b", {})
    full_name = worker_b.get("full_name", "")

    # Pokud je full_name prázdný string, None nebo whitespace, je to single-worker
    is_single = not bool(full_name.strip() if full_name else False)

    return is_single


def _get_pp_max_category(
    results_data: Dict[str, Any],
    category_limits: Dict[str, Any],
    use_worker_a_only: bool = False
) -> int:
    """
    Určí nejvyšší kategorii pracovních poloh (1, 2, nebo 3).

    Projde všechny sekce a polohy, vypočítá jejich kategorie podle limitů,
    a vrátí nejvyšší nalezenou kategorii.

    Pro single-worker protokoly (use_worker_a_only=True) používá vyskyt_min_a,
    pro two-worker protokoly používá prumer_min.

    Args:
        results_data: pp_results.json s daty o polohách
        category_limits: Dictionary s limity kategorií (pp_kat1_max, n_kat1_max, atd.)
        use_worker_a_only: Pokud True, používá vyskyt_min_a místo prumer_min

    Returns:
        1, 2, nebo 3 (nejvyšší kategorie)
    """
    sections = ["trup", "hlava_krk", "phk", "lhk", "dk", "ostatni"]
    max_category = 1

    for section_name in sections:
        section = results_data.get(section_name, [])

        if isinstance(section, list):
            items = section
        elif isinstance(section, dict):
            items = section.values()
        else:
            continue

        for row_data in items:
            if not isinstance(row_data, dict):
                continue

            # Vybrat správnou hodnotu podle módu
            if use_worker_a_only:
                comparison_value = row_data.get("vyskyt_min_a", 0)
            else:
                comparison_value = row_data.get("prumer_min", 0)

            typ = row_data.get("typ_pracovni_polohy", "N")

            # Skip pokud je hodnota 0 nebo None
            if comparison_value == 0 or comparison_value is None:
                continue

            # Určit kategorii podle typu a limitů
            if typ == "PP":
                if comparison_value <= category_limits["pp_kat1_max"]:
                    category = 1
                elif comparison_value <= category_limits["pp_kat2_max"]:
                    category = 2
                else:
                    category = 3
            else:  # typ == "N"
                if comparison_value <= category_limits["n_kat1_max"]:
                    category = 1
                elif comparison_value <= category_limits["n_kat2_max"]:
                    category = 2
                else:
                    category = 3

            # Update max kategorie
            if category > max_category:
                max_category = category

    return max_category


def _get_pp_problematic_positions_list(
    results_data: Dict[str, Any],
    category_limits: Dict[str, Any],
    use_worker_a_only: bool = False
) -> str:
    """
    Vrací POUZE seznam problematických poloh (nejvyšší kategorie), oddělených čárkou.
    Použití v šabloně jako samostatný placeholder.

    Pro single-worker protokoly (use_worker_a_only=True) používá vyskyt_min_a,
    pro two-worker protokoly používá prumer_min.

    Args:
        results_data: pp_results.json s daty o polohách
        category_limits: Dictionary s limity kategorií (pp_kat1_max, n_kat1_max, atd.)
        use_worker_a_only: Pokud True, používá vyskyt_min_a místo prumer_min

    Returns:
        Čárkou oddělený seznam poloh (např. "dynamické předklon trupu, statické otočení hlavy")
        Prázdný string pokud není žádná problematická poloha
    """
    # Nejdřív zjisti max kategorii (předat parametr!)
    max_category = _get_pp_max_category(results_data, category_limits, use_worker_a_only)

    if max_category == 1:
        return ""  # Žádné problémy

    sections = ["trup", "hlava_krk", "phk", "lhk", "dk", "ostatni"]
    category_2_problems = []  # Kategorie 2+
    category_3_problems = []  # Kategorie 3

    # Projít všechny sekce a sbírat problematické polohy
    for section_name in sections:
        section = results_data.get(section_name, [])

        if isinstance(section, list):
            items = section
        elif isinstance(section, dict):
            items = section.values()
        else:
            continue

        for row_data in items:
            if not isinstance(row_data, dict):
                continue

            # Vybrat správnou hodnotu podle módu
            if use_worker_a_only:
                comparison_value = row_data.get("vyskyt_min_a", 0)
            else:
                comparison_value = row_data.get("prumer_min", 0)

            typ = row_data.get("typ_pracovni_polohy", "N")
            typ_svalove_prace = row_data.get("typ_svalove_prace", "")
            nazev_polohy = row_data.get("nazev_polohy", "")

            if comparison_value == 0 or comparison_value is None:
                continue

            # Určit kategorii podle typu a limitů
            if typ == "PP":
                if comparison_value <= category_limits["pp_kat1_max"]:
                    category = 1
                elif comparison_value <= category_limits["pp_kat2_max"]:
                    category = 2
                else:
                    category = 3
            else:  # typ == "N"
                if comparison_value <= category_limits["n_kat1_max"]:
                    category = 1
                elif comparison_value <= category_limits["n_kat2_max"]:
                    category = 2
                else:
                    category = 3

            # Sbírat problematické polohy
            formatted_name = _format_pp_position_name(typ_svalove_prace, nazev_polohy)

            if category >= 2:
                category_2_problems.append(formatted_name)

            if category == 3:
                category_3_problems.append(formatted_name)

    # Vrátit seznam podle max kategorie
    if max_category == 3:
        return ", ".join(category_3_problems)
    else:  # max_category == 2
        return ", ".join(category_2_problems)


def _calculate_prvni_pp_podminka_kategorie(
    results_data: Dict[str, Any],
    category_limits: Dict[str, Any],
    use_worker_a_only: bool = False
) -> str:
    """
    První PP podmínka: Kategorizace pracovních poloh podle času.

    Projde všechny sekce (trup, hlava_krk, phk, lhk, dk, ostatni) a pro každý řádek
    určí kategorii (1, 2, nebo 3) podle času a typ_pracovni_polohy.

    Pro single-worker protokoly (use_worker_a_only=True) používá vyskyt_min_a,
    pro two-worker protokoly používá prumer_min.

    Logika:
    - Pokud VŠECHNY polohy jsou kategorie 1 → text kategorie 1
    - Pokud ALESPOŇ JEDNA je kategorie 2 (ale žádná 3) → text kategorie 2 + výčet
    - Pokud ALESPOŇ JEDNA je kategorie 3 → text kategorie 3 + výčet

    Args:
        results_data: pp_results.json s daty o polohách
        category_limits: Dictionary s limity kategorií (pp_kat1_max, n_kat1_max, atd.)
        use_worker_a_only: Pokud True, používá vyskyt_min_a místo prumer_min

    Returns:
        Text podle nejvyšší nalezené kategorie
    """
    # Sekce které obsahují data o polohách
    sections = ["trup", "hlava_krk", "phk", "lhk", "dk", "ostatni"]

    max_category = 1  # Začínáme s kategorií 1 (optimisticky)

    # Sbírat problematické polohy
    category_2_problems = {"PP": [], "N": []}  # Překročily kat1 (jsou kat2+)
    category_3_problems = {"PP": [], "N": []}  # Překročily kat2 (jsou kat3)

    # Projít všechny sekce
    for section_name in sections:
        section = results_data.get(section_name, [])

        # Zkontrolovat, jestli je sekce list nebo dict
        if isinstance(section, list):
            # pp_results.json má sekce jako array
            items = section
        elif isinstance(section, dict):
            # Fallback pro dict strukturu (numeric string keys)
            items = section.values()
        else:
            continue

        # Projít všechny řádky v sekci
        for row_data in items:
            if not isinstance(row_data, dict):
                continue

            # Vybrat správnou hodnotu podle módu
            if use_worker_a_only:
                comparison_value = row_data.get("vyskyt_min_a", 0)
            else:
                comparison_value = row_data.get("prumer_min", 0)

            typ = row_data.get("typ_pracovni_polohy", "N")
            typ_svalove_prace = row_data.get("typ_svalove_prace", "")
            nazev_polohy = row_data.get("nazev_polohy", "")

            # Skip pokud je hodnota 0, null nebo prázdná (nenaměřená poloha)
            if comparison_value == 0 or comparison_value is None:
                continue

            # Určit kategorii podle typu a limitů
            if typ == "PP":
                if comparison_value <= category_limits["pp_kat1_max"]:
                    category = 1
                elif comparison_value <= category_limits["pp_kat2_max"]:
                    category = 2
                else:
                    category = 3
            else:  # typ == "N"
                if comparison_value <= category_limits["n_kat1_max"]:
                    category = 1
                elif comparison_value <= category_limits["n_kat2_max"]:
                    category = 2
                else:
                    category = 3

            # Update max kategorie (priorita: 3 > 2 > 1)
            if category > max_category:
                max_category = category

            # Sbírat problematické polohy
            if category >= 2:  # Překročily kat1
                formatted_name = _format_pp_position_name(typ_svalove_prace, nazev_polohy)
                category_2_problems[typ].append(formatted_name)

            if category == 3:  # Překročily kat2
                formatted_name = _format_pp_position_name(typ_svalove_prace, nazev_polohy)
                category_3_problems[typ].append(formatted_name)

    # Vygenerovat text podle max kategorie
    if max_category == 1:
        return "Celková doba zaujímání nepřijatelných pracovních poloh a podmíněně přijatelných poloh nepřekročila v průměru přípustný limit kategorie 1."

    elif max_category == 2:
        # Použít category_2_problems
        has_pp = len(category_2_problems["PP"]) > 0
        has_n = len(category_2_problems["N"]) > 0

        # Spojit všechny problematické polohy
        all_problems = category_2_problems["PP"] + category_2_problems["N"]
        problems_text = ", ".join(all_problems)

        # Vybrat šablonu
        if has_pp and has_n:
            # Obě
            return f"Celková doba zaujímání nepřijatelných pracovních poloh a podmíněně přijatelných poloh překročila v průměru přípustný limit kategorie 1 pro {problems_text}."
        elif has_pp:
            # Jen PP
            return f"Celková doba zaujímání podmíněně přijatelných poloh překročila v průměru přípustný limit kategorie 1 pro {problems_text}."
        else:  # has_n
            # Jen N
            return f"Celková doba zaujímání nepřijatelných pracovních poloh překročila v průměru přípustný limit kategorie 1 pro {problems_text}."

    else:  # max_category == 3
        # Použít category_3_problems
        has_pp = len(category_3_problems["PP"]) > 0
        has_n = len(category_3_problems["N"]) > 0

        # Spojit všechny problematické polohy
        all_problems = category_3_problems["PP"] + category_3_problems["N"]
        problems_text = ", ".join(all_problems)

        # Vybrat šablonu (stejná jako kat2, jen "kategorie 2" místo "kategorie 1")
        if has_pp and has_n:
            # Obě
            return f"Celková doba zaujímání nepřijatelných pracovních poloh a podmíněně přijatelných poloh překročila v průměru přípustný limit kategorie 2 pro {problems_text}."
        elif has_pp:
            # Jen PP
            return f"Celková doba zaujímání podmíněně přijatelných poloh překročila v průměru přípustný limit kategorie 2 pro {problems_text}."
        else:  # has_n
            # Jen N
            return f"Celková doba zaujímání nepřijatelných pracovních poloh překročila v průměru přípustný limit kategorie 2 pro {problems_text}."


def generate_conditional_texts(
    measurement_data: Dict[str, Any],
    results_data: Optional[Dict[str, Any]] = None,
    protocol_type: str = "LSZ"
) -> Dict[str, str]:
    """
    Vygeneruje podmínkové texty na základě dat z measurement_data.json a results_data.
    Podporuje LSZ, PP a CFZ protokoly.

    Args:
        measurement_data: Data z measurement_data.json (celý JSON)
        results_data: Data z results JSON (lsz_results.json nebo pp_results.json), volitelné
        protocol_type: Typ protokolu ("LSZ", "PP_CAS", "PP_KUSY", "CFZ")

    Returns:
        Dictionary s vygenerovanými texty:
        {
            "prvni_text_podminka_pocetdni": "...",
            "druhy_text_podminka_limit1": "...",  (LSZ-specific)
            ... další podmínkové texty podle protokolu
        }
    """
    texts = {}

    # Pro PP protokoly: použij stub implementaci (zatím jen prvni_text_podminka)
    if protocol_type in ("PP_CAS", "PP_KUSY"):
        print(f"Generuji PP podminkove texty")

        # Detekovat single-worker mód
        is_single_worker = _is_single_worker_protocol(measurement_data)

        if is_single_worker:
            print("  [INFO] Single-worker protocol detected - using worker A data (vyskyt_min_a)")
        else:
            print("  [INFO] Two-worker protocol - using average data (prumer_min)")

        # PP má jen generickou prvni_text_podminka
        section0 = measurement_data.get("section0_file_selection", {})
        section2 = measurement_data.get("section2_firma", {})
        measurement_days = section2.get("measurement_days", 1)
        gender = section0.get("workers_gender", "muži")
        worker_count = section0.get("worker_count", 2)

        prvni_text_varianty = {
            (1, 1, "muži"): "Měření probíhalo v jednom dni, v jedné průměrné směně. Měřen byl 1 pracovník – muž.",
            (1, 1, "ženy"): "Měření probíhalo v jednom dni, v jedné průměrné směně. Měřena byla 1 pracovnice – žena.",
            (2, 1, "muži"): "Měření probíhalo ve dvou dnech, ve dvou průměrných směnách. Měřen byl 1 pracovník – muž.",
            (2, 1, "ženy"): "Měření probíhalo ve dvou dnech, ve dvou průměrných směnách. Měřena byla 1 pracovnice – žena.",
            (1, 2, "muži"): "Měření probíhalo v jednom dni, v jedné průměrné směně. Měřeni byli 2 pracovníci – muži.",
            (1, 2, "ženy"): "Měření probíhalo v jednom dni, v jedné průměrné směně. Měřeny byly 2 pracovnice – ženy.",
            (2, 2, "muži"): "Měření probíhalo ve dvou dnech, ve dvou průměrných směnách. Měřeni byli 2 pracovníci – muži.",
            (2, 2, "ženy"): "Měření probíhalo ve dvou dnech, ve dvou průměrných směnách. Měřeny byly 2 pracovnice – ženy.",
        }

        texts["prvni_text_podminka_pocetdni"] = prvni_text_varianty.get(
            (measurement_days, worker_count, gender),
            prvni_text_varianty[(1, 2, "muži")]
        )

        # PP PODMÍNKA: Limitní hodnoty kategorií podle délky směny
        # Načíst z section4_worker_a.work_duration (STRING -> int)
        section4 = measurement_data.get("section4_worker_a", {})
        work_duration_str = section4.get("work_duration")

        if work_duration_str:
            work_duration_min = int(work_duration_str)
        else:
            work_duration_min = 480  # Fallback na standard

        # Vypočítat koeficient a kategoriální limity
        coefficient = _calculate_pp_work_shift_coefficient(work_duration_min)
        category_limits = _calculate_pp_category_limits(coefficient)

        # Přidat limity do texts (pro zobrazení v Word a pro další výpočty)
        texts.update(category_limits)

        # Debug vypis
        print(f"  PP smena: {work_duration_min} min -> koeficient: {coefficient:.3f}")
        print(f"  PP limity: kat1 0-{category_limits['pp_kat1_max']}, kat2 {category_limits['pp_kat2_min']}-{category_limits['pp_kat2_max']}, kat3 >{category_limits['pp_kat3_min']}")
        print(f"  N limity: kat1 0-{category_limits['n_kat1_max']}, kat2 {category_limits['n_kat2_min']}-{category_limits['n_kat2_max']}, kat3 >{category_limits['n_kat3_min']}")

        # PP PODMÍNKA: Kategorizace pracovních poloh
        if results_data is not None:
            texts["prvni_pp_podminka_kategorie"] = _calculate_prvni_pp_podminka_kategorie(
                results_data,
                category_limits,
                use_worker_a_only=is_single_worker
            )
            print(f"  PP kategorizace: {texts['prvni_pp_podminka_kategorie'][:50].encode('ascii', 'replace').decode('ascii')}...")

            # PP SEZNAM: Pouze seznam problematických poloh (pro samostatné použití v šabloně)
            texts["pp_problematicke_polohy_seznam"] = _get_pp_problematic_positions_list(
                results_data,
                category_limits,
                use_worker_a_only=is_single_worker
            )
            if texts["pp_problematicke_polohy_seznam"]:
                print(f"  PP seznam problémových poloh: {texts['pp_problematicke_polohy_seznam'][:80].encode('ascii', 'replace').decode('ascii')}...")
            else:
                print("  PP seznam problémových poloh: (prázdný - kategorie 1)")

            # PP KATEGORIE ČÍSLO: Pouze číslo nejvyšší kategorie (1, 2, nebo 3)
            max_category_number = _get_pp_max_category(
                results_data,
                category_limits,
                use_worker_a_only=is_single_worker
            )
            texts["pp_kategorie_cislo"] = str(max_category_number)
            print(f"  PP kategorie cislo: {texts['pp_kategorie_cislo']}")

            # PP PŘEKROČENÁ KATEGORIE: Číslo kategorie, která byla překročena (0, 1, nebo 2)
            # Kategorie 1 → překročena 0 (žádná)
            # Kategorie 2 → překročena 1
            # Kategorie 3 → překročena 2
            if max_category_number == 1:
                texts["pp_prekocrena_kategorie"] = ""
            else:
                texts["pp_prekocrena_kategorie"] = str(max_category_number - 1)
            print(f"  PP prekrocena kategorie: '{texts['pp_prekocrena_kategorie']}'")
        else:
            texts["prvni_pp_podminka_kategorie"] = "PP results data not available"
            texts["pp_problematicke_polohy_seznam"] = ""
            texts["pp_kategorie_cislo"] = "1"  # Default kategorie 1
            texts["pp_prekocrena_kategorie"] = ""  # Default žádná překročená
            print("  [WARNING] PP kategorizace preskocena - results_data is None")

        # PP-specific stub texty (další podmínky budou implementovány později)
        texts["pp_placeholder_text"] = "PP conditional texts - TODO"

        return texts

    # Pro LSZ: použij existující implementaci (níže)
    texts = {}

    # PODMÍNKA 1: Počet dnů měření + pohlaví + počet pracovníků
    section0 = measurement_data.get("section0_file_selection", {})
    section2 = measurement_data.get("section2_firma", {})
    measurement_days = section2.get("measurement_days", 1)  # Z section2_firma
    gender = section0.get("workers_gender", "muži")          # Z section0_file_selection
    worker_count = section0.get("worker_count", 2)           # Z section0_file_selection

    # Matice všech kombinací (2 dny × 2 worker_count × 2 pohlaví = 8 variant)
    prvni_text_varianty = {
        # 1 pracovník, 1 den
        (1, 1, "muži"): "Měření probíhalo v jednom dni, v jedné průměrné směně. Měřen byl 1 pracovník – muž.",
        (1, 1, "ženy"): "Měření probíhalo v jednom dni, v jedné průměrné směně. Měřena byla 1 pracovnice – žena.",

        # 1 pracovník, 2 dny
        (2, 1, "muži"): "Měření probíhalo ve dvou dnech, ve dvou průměrných směnách. Měřen byl 1 pracovník – muž.",
        (2, 1, "ženy"): "Měření probíhalo ve dvou dnech, ve dvou průměrných směnách. Měřena byla 1 pracovnice – žena.",

        # 2 pracovníci, 1 den
        (1, 2, "muži"): "Měření probíhalo v jednom dni, v jedné průměrné směně. Měřeni byli 2 pracovníci – muži.",
        (1, 2, "ženy"): "Měření probíhalo v jednom dni, v jedné průměrné směně. Měřeny byly 2 pracovnice – ženy.",

        # 2 pracovníci, 2 dny
        (2, 2, "muži"): "Měření probíhalo ve dvou dnech, ve dvou průměrných směnách. Měřeni byli 2 pracovníci – muži.",
        (2, 2, "ženy"): "Měření probíhalo ve dvou dnech, ve dvou průměrných směnách. Měřeny byly 2 pracovnice – ženy.",
    }

    texts["prvni_text_podminka_pocetdni"] = prvni_text_varianty.get((measurement_days, worker_count, gender), prvni_text_varianty[(1, 2, "muži")])

    # PODMÍNKA 2: Hygienické limity PHK (pouze pokud jsou k dispozici results_data)
    if results_data is not None:
        texts["druhy_text_podminka_limit1"] = _calculate_druhy_text_podminka_limit1(results_data)

    # PODMÍNKA 3: Hygienické limity LHK (pouze pokud jsou k dispozici results_data)
    if results_data is not None:
        texts["treti_text_podminka_limit1"] = _calculate_treti_text_podminka_limit1(results_data)

    # PODMÍNKA 4: Rozložení svalových sil (pouze pokud jsou k dispozici results_data)
    if results_data is not None:
        texts["ctvrty_text_podminka"] = _calculate_ctvrty_text_podminka(results_data, worker_count)

    # PODMÍNKA 5: Nadlimitní svalové síly nad 70% Fmax (pouze pokud jsou k dispozici results_data)
    if results_data is not None:
        texts["paty_text_podminka"] = _calculate_paty_text_podminka(results_data, worker_count)

    # PODMÍNKA 6: Hodnoty > 100 (pouze pokud jsou k dispozici results_data)
    if results_data is not None:
        texts["sesty_text_podminka"] = _calculate_sesty_text_podminka(results_data, worker_count)

    # PODMÍNKA 7: Překročení limitu pro velké svalové síly 55-70% Fmax (vyžaduje measurement_data i results_data)
    if results_data is not None:
        texts["sedmy_text_podminka"] = _calculate_sedmy_text_podminka(measurement_data, results_data, worker_count)

    # PODMÍNKA 8: Seznam činností s force_over_70 > 100 (pouze pokud jsou k dispozici results_data)
    if results_data is not None:
        texts["osmy_text_podminka"] = _calculate_osmy_text_podminka(results_data, worker_count)

    # PODMÍNKA 9: Tabulka hygienických limitů (Nad limitem / Pod limitem pro všechny 4 svalové skupiny)
    if results_data is not None:
        texts["devata_text_podminka"] = _calculate_devata_text_podminka(results_data)

    # PODMÍNKA 10: Celkové překročení limitu (pokud je alespoň jedna svalová skupina nad limitem)
    if results_data is not None:
        devata_data = texts.get("devata_text_podminka", {})
        texts["desata_text_podminka"] = _calculate_desata_text_podminka(devata_data)

    # PODMÍNKA 11: Hierarchické vyhodnocení zatížení (1/3 limitu vs celý limit)
    if results_data is not None:
        texts["jedenacta_text_podminka"] = _calculate_jedenacta_text_podminka(results_data)

    # HYGIENICKÉ LIMITY: Číselné hodnoty limitů z table_W4_Y51
    if results_data is not None:
        texts["hygiene_limits"] = _calculate_hygiene_limits(results_data)

    # HYGIENICKÝ LIMIT PRO 55-70% FMAX: Vypočítané číslo podle délky směny
    work_duration = measurement_data.get("section4_worker_a", {}).get("work_duration")
    if work_duration is not None:
        try:
            work_duration = float(work_duration)
            texts["hygiene_limit_55_70"] = int((work_duration / 2) + 360)
            texts["shift_duration_text"] = _get_shift_duration_text(work_duration)
        except (ValueError, TypeError):
            texts["hygiene_limit_55_70"] = 600  # Fallback pro 8h směnu
            texts["shift_duration_text"] = "(osmihodinovou)"
    else:
        texts["hygiene_limit_55_70"] = 600  # Fallback pro 8h směnu
        texts["shift_duration_text"] = "(osmihodinovou)"

    # PODMÍNĚNÉ PLACEHOLDERY: Vypíší text pouze pokud je překročení
    # Sesty text - pokud DOCHÁZÍ k pravidelným nadlimitním silám (> 100)
    if results_data is not None:
        sesty = texts.get("sesty_text_podminka", "")
        if "je pravidelnou součástí" in sesty:
            texts["sesty_text_if_true"] = texts["sesty_text_podminka"]
        else:
            texts["sesty_text_if_true"] = ""

    # Sedmy text - pokud PŘEKRAČUJE limit pro 55-70% Fmax
    if results_data is not None:
        sedmy = texts.get("sedmy_text_podminka", "")
        if "překračuje" in sedmy and "nepřekračuje" not in sedmy:
            texts["sedmy_text_if_true"] = texts["sedmy_text_podminka"]
        else:
            texts["sedmy_text_if_true"] = ""

    return texts


# ============================================================================
# HOLTER FORMATTING - Zvýraznění vybraných holterů ve Word tabulce
# ============================================================================

# Mapování holter ID → číslo holteru v textu
HOLTER_MAPPING = {
    "A": "60/16",
    "B": "65/17",
    "C": "84/19",
    "D": "85/19",
    "E": "86/20",
    "F": "87/20"
}


def get_selected_holter_numbers(measurement_data: Dict[str, Any]) -> list:
    """
    Zjistí vybraná čísla holterů z measurement_data.

    Args:
        measurement_data: Data z measurement_data.json (celý JSON)

    Returns:
        List čísel holterů (např. ["65/17", "85/19"])
    """
    selected = []

    # Pracovník A
    holter_a = measurement_data.get("section4_worker_a", {}).get("emg_holter")
    if holter_a and holter_a in HOLTER_MAPPING:
        selected.append(HOLTER_MAPPING[holter_a])

    # Pracovník B
    holter_b = measurement_data.get("section5_worker_b", {}).get("emg_holter")
    if holter_b and holter_b in HOLTER_MAPPING:
        selected.append(HOLTER_MAPPING[holter_b])

    return selected


def highlight_selected_holters(docx_path: str, selected_holter_numbers: list) -> None:
    """
    Zvýrazní tučně řádky s vybranými holtry v tabulce měřících přístrojů.

    BEZPEČNÝ PŘÍSTUP: Post-processing po vyrendrování šablony.
    Nepoužívá RichText v tabulkách (riskantní), ale upravuje hotový dokument.

    Args:
        docx_path: Cesta k vygenerovanému Word dokumentu
        selected_holter_numbers: List čísel holterů (např. ["65/17", "85/19"])

    Returns:
        None (upravuje dokument in-place)
    """
    from docx import Document

    # Pokud nejsou vybrané holteru, není co zvýrazňovat
    if not selected_holter_numbers:
        return

    # Otevři dokument
    doc = Document(docx_path)

    # Najdi správnou tabulku (ta s nadpisem "Typ" v prvním sloupci)
    for table in doc.tables:
        if len(table.rows) == 0:
            continue

        # Zkontroluj první buňku prvního řádku (hlavička)
        first_cell_text = table.rows[0].cells[0].text.strip()

        # Pokud obsahuje "Typ", je to správná tabulka
        if "Typ" in first_cell_text:
            # Projdi všechny řádky kromě hlavičky
            for row_idx, row in enumerate(table.rows[1:], start=1):
                cell_text = row.cells[0].text

                # Pokud řádek obsahuje vybraný holter
                if any(holter_num in cell_text for holter_num in selected_holter_numbers):
                    # Nastav bold pro všechny runs v prvním sloupci
                    for paragraph in row.cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

            # Tabulku jsme našli a zpracovali, můžeme skončit
            break

    # Ulož změny (přepíše originální soubor)
    doc.save(docx_path)


# ============================================================================
# FORCE DISTRIBUTION HIGHLIGHTING - Červené zvýraznění nadlimitních hodnot
# ============================================================================

def highlight_force_distribution_values(docx_path: str, measurement_data: Dict[str, Any], results_data: Dict[str, Any]) -> None:
    """
    Červeně zvýrazní nadlimitní hodnoty v tabulce force_distribution.

    DVOJE KRITÉRIA:
    1. force_over_70_* > 100 → červeně
    2. force_55_70_* > hygienický limit → červeně
       - Limit = (work_duration / 2) + 360 (stejná logika jako sedmá podmínka)

    BEZPEČNÝ PŘÍSTUP: Post-processing po vyrendrování šablony.
    Nepoužívá RichText v tabulkách (riskantní), ale upravuje hotový dokument.

    Args:
        docx_path: Cesta k vygenerovanému Word dokumentu
        measurement_data: Data z measurement_data.json (pro work_duration)
        results_data: Data z lsz_results.json (pro table_force_distribution)

    Returns:
        None (upravuje dokument in-place)
    """
    from docx import Document
    from docx.shared import RGBColor

    # Vypočti hygienický limit pro force_55_70 (stejná logika jako sedmá podmínka)
    work_duration = measurement_data.get("section4_worker_a", {}).get("work_duration")

    if work_duration is None:
        # Fallback - pokud není work_duration, použij výchozí 8h směnu
        limit_55_70 = 600
    else:
        try:
            work_duration = float(work_duration)
            limit_55_70 = (work_duration / 2) + 360
        except (ValueError, TypeError):
            limit_55_70 = 600  # Fallback

    # Načti tabulku force_distribution
    table_force_dist = results_data.get("table_force_distribution", {})
    if not table_force_dist:
        return  # Není co zvýrazňovat

    # Otevři dokument
    doc = Document(docx_path)

    # Najdi správnou tabulku - tabulka force_distribution
    # Identifikace:
    # - 9 sloupců (activity + 8 hodnot)
    # - Obsahuje "Výskyt sil" nebo "Rozpis" v hlavičce
    target_table = None
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue

        # Zkontroluj hlavičku
        header_row = table.rows[0]

        # Musí mít 9 sloupců
        if len(header_row.cells) != 9:
            continue

        header_text = " ".join([cell.text for cell in header_row.cells]).lower()

        # Musí obsahovat klíčová slova
        if ("výskyt sil" in header_text or "vyskyt sil" in header_text) and \
           ("rozpis" in header_text or "innost" in header_text or "activity" in header_text):
            target_table = table
            break

    if target_table is None:
        return  # Tabulka nenalezena

    # Identifikuj indexy sloupců (force_over_70 a force_55_70)
    # Struktura tabulky (podle TABLES_ANALYSIS.md):
    # Sloupec 0: activity
    # Sloupec 1-4: force_55_70 (PHK ext, PHK flex, LHK ext, LHK flex)
    # Sloupec 5-8: force_over_70 (PHK ext, PHK flex, LHK ext, LHK flex)

    force_55_70_columns = [1, 2, 3, 4]  # PHK ext, PHK flex, LHK ext, LHK flex
    force_over_70_columns = [5, 6, 7, 8]  # PHK ext, PHK flex, LHK ext, LHK flex

    # Projdi všechny řádky
    for row_idx, row in enumerate(target_table.rows):
        # Zpracuj VŠECHNY řádky včetně "Celkem" a "Časově vážený průměr"
        # Skip pouze řádky hlavičky - pozná se tak, že sloupec 5 není číslo

        # Zkontroluj, jestli je to hlavička (sloupec 5 není číslo)
        try:
            test_value = row.cells[5].text.strip()
            float(test_value)  # Pokud to selže, je to hlavička
        except (ValueError, IndexError):
            continue  # Skip hlavičku

        # === FORCE_OVER_70: Zkontroluj hodnoty > 100 ===
        for col_idx in force_over_70_columns:
            if col_idx >= len(row.cells):
                continue

            cell = row.cells[col_idx]
            cell_text = cell.text.strip()

            try:
                value = float(cell_text)
                if value > 100:
                    # Obarvi červeně - MUSÍME PŘEPSAT OBSAH, NE JEN MĚNIT BARVU
                    # Docxtpl může vytvářet speciální run strukturu, která se neobarvuje správně

                    # Smaž všechny paragraphy
                    for paragraph in cell.paragraphs:
                        p = paragraph._element
                        p.getparent().remove(p)

                    # Vytvoř nový paragraph s červeným runem
                    new_para = cell.add_paragraph()
                    run = new_para.add_run(cell_text)
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True
            except (ValueError, TypeError):
                pass  # Není číslo, skip

        # === FORCE_55_70: Zkontroluj hodnoty > limit ===
        for col_idx in force_55_70_columns:
            if col_idx >= len(row.cells):
                continue

            cell = row.cells[col_idx]
            cell_text = cell.text.strip()

            try:
                value = float(cell_text)
                if value > limit_55_70:
                    # Obarvi červeně - MUSÍME PŘEPSAT OBSAH, NE JEN MĚNIT BARVU

                    # Smaž všechny paragraphy
                    for paragraph in cell.paragraphs:
                        p = paragraph._element
                        p.getparent().remove(p)

                    # Vytvoř nový paragraph s červeným runem
                    new_para = cell.add_paragraph()
                    run = new_para.add_run(cell_text)
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True
            except (ValueError, TypeError):
                pass  # Není číslo, skip

    # Ulož změny (přepíše originální soubor)
    doc.save(docx_path)


def highlight_nad_limitem_values(docx_path: str) -> None:
    """
    Červeně a tučně zvýrazní texty "nad limitem" nebo "Nad limitem" v tabulce s porovnáním limitů.

    Tabulka má strukturu:
    - 6 sloupců
    - Sloupec 0: Pozice (profession_name)
    - Sloupce 1-4: PHK Extenzory, PHK Flexory, LHK Extenzory, LHK Flexory (obsahují "Nad limitem" nebo "Pod limitem")
    - Sloupec 5: Konečné hodnocení

    Args:
        docx_path: Cesta k vygenerovanému Word dokumentu

    Returns:
        None (upravuje dokument in-place)
    """
    from docx import Document
    from docx.shared import RGBColor

    # Otevři dokument
    doc = Document(docx_path)

    # Najdi správnou tabulku - tabulka s "Porovnání naměřených" a "Konečné hodnocení"
    # Identifikace:
    # - 6 sloupců
    # - Obsahuje "Porovnání" nebo "Konečné hodnocení" v hlavičce
    target_table = None
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue

        # Musí mít 6 sloupců
        if len(table.columns) != 6:
            continue

        # Zkontroluj hlavičku (první dva řádky)
        header_text = ""
        for row_idx in range(min(2, len(table.rows))):
            header_text += " ".join([cell.text for cell in table.rows[row_idx].cells]).lower()

        # Musí obsahovat klíčová slova
        if ("porovnání" in header_text or "porovnani" in header_text) and \
           ("konečné hodnocení" in header_text or "konecne hodnoceni" in header_text):
            target_table = table
            break

    if target_table is None:
        return  # Tabulka nenalezena

    # Sloupce 1-4 obsahují texty "Nad limitem" / "Pod limitem"
    limit_columns = [1, 2, 3, 4]

    # Projdi všechny řádky
    for row_idx, row in enumerate(target_table.rows):
        # Přeskoč hlavičkové řádky (první 2 řádky)
        if row_idx < 2:
            continue

        # Projdi sloupce 1-4
        for col_idx in limit_columns:
            if col_idx >= len(row.cells):
                continue

            cell = row.cells[col_idx]
            cell_text = cell.text.strip()

            # Zkontroluj, jestli obsahuje "nad limitem" (case insensitive)
            if "nad limitem" in cell_text.lower():
                # Obarvi červeně a tučně

                # Smaž všechny paragraphy
                for paragraph in cell.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)

                # Vytvoř nový paragraph s červeným a tučným runem
                new_para = cell.add_paragraph()
                run = new_para.add_run(cell_text)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True

    # Ulož změny (přepíše originální soubor)
    doc.save(docx_path)


def highlight_pp_typ_svalove_prace(docx_path: str) -> None:
    """
    Obarví background buněk s typem svalové práce v PP tabulkách.

    - "Statická" → background #FDE9D9 (světle oranžová/béžová)
    - "Dynamická" → background #95B3D7 (světle modrá)

    Args:
        docx_path: Cesta k vygenerovanému Word dokumentu

    Returns:
        None (upravuje dokument in-place)
    """
    from docx import Document

    # Otevři dokument
    doc = Document(docx_path)

    colored_count = 0

    # Projdi všechny tabulky
    for table_idx, table in enumerate(doc.tables):
        # Zkontroluj, jestli je to PP tabulka
        if not is_pp_table(table):
            continue

        # Projdi všechny řádky (kromě hlavičky)
        for row_idx, row in enumerate(table.rows):
            # Přeskoč první řádek (hlavička)
            if row_idx == 0:
                continue

            # Sloupec 1 = Typ svalové práce
            if len(row.cells) < 2:
                continue

            cell = row.cells[1]
            cell_text = cell.text.strip().lower()

            # Přeskoč prázdné buňky a hlavičky
            if not cell_text or "svalová práce" in cell_text or "muž" in cell_text or "výskyt" in cell_text:
                continue

            # Obarvi podle typu
            if "statická" in cell_text or "staticka" in cell_text:
                set_cell_background(cell, "FDE9D9")  # Světle oranžová/béžová
                colored_count += 1
            elif "dynamická" in cell_text or "dynamicka" in cell_text:
                set_cell_background(cell, "95B3D7")  # Světle modrá
                colored_count += 1

    # Ulož změny
    doc.save(docx_path)
    print(f"  ✓ Obarveno {colored_count} buněk s typem svalové práce")


def set_cell_background(cell, hex_color: str) -> None:
    """
    Nastaví background (shading) buňky na zadanou hex barvu.

    Používá XML manipulaci pro nastavení w:shd elementu s w:fill atributem.

    Args:
        cell: python-docx Cell objekt
        hex_color: Hex kód barvy bez # (např. "90EE90" pro světle zelenou)

    Example:
        set_cell_background(cell, "90EE90")  # Zelená
        set_cell_background(cell, "FFFF00")  # Žlutá
        set_cell_background(cell, "FFC8C8")  # Světle červená
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Vytvoř shading element
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), hex_color)

    # Přidej do cell properties (tcPr)
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.append(shading_elm)


def is_pp_table(table) -> bool:
    """
    Zkontroluje, jestli je tabulka PP tabulka (pracovní polohy).

    PP tabulka má:
    - 6 sloupců (two-worker: nazev, typ_prace, vyskyt_a, vyskyt_b, prumer, typ_pp)
    - 4 sloupce (single-worker: nazev, typ_prace, vyskyt_a, typ_pp)
    - Obsahuje klíčová slova: TRUP, HLAVA_KRK, PHK, LHK, atd.

    Args:
        table: python-docx Table objekt

    Returns:
        True pokud je to PP tabulka, False jinak
    """
    # Kontrola počtu sloupců (4 nebo 6)
    if len(table.columns) not in [4, 6]:
        return False

    # Klíčová slova pro PP tabulky
    pp_keywords = [
        "TRUP", "Trup",
        "HLAVA_KRK", "HLAVA A KRK", "Hlava a krk",
        "PHK", "Pravá horní končetina",
        "LHK", "Levá horní končetina",
        "DOLNÍ KONČETINY", "Dolní končetiny", "DKK",
        "OSTATNÍ ČÁSTI TĚLA", "Ostatní části těla", "OST",
        "Statická", "Dynamická"  # Typy svalové práce
    ]

    # Zkontroluj prvních 5 řádků na přítomnost keywords
    for row in table.rows[:5]:
        row_text = " ".join(cell.text for cell in row.cells)
        for keyword in pp_keywords:
            if keyword in row_text:
                return True

    return False


def is_pp_section_header(text: str) -> bool:
    """
    Zkontroluje, jestli je text section header (TRUP, HLAVA_KRK, atd.).

    Args:
        text: Text z buňky

    Returns:
        True pokud je to section header, False jinak
    """
    section_keywords = [
        "TRUP", "HLAVA_KRK", "HLAVA A KRK", "PHK", "LHK",
        "DOLNÍ KONČETINY", "DKK", "OSTATNÍ ČÁSTI TĚLA",
        "OSTATNI", "OST", "ULTRATHINK"
    ]

    text_upper = text.upper().strip()
    return text_upper in section_keywords


def calculate_pp_category(
    prumer_min: float,
    typ_pracovni_polohy: str,
    category_limits: Dict[str, float]
) -> int:
    """
    Vypočítá kategorii (1, 2, nebo 3) pro PP pracovní polohu.

    Logika podle NV 361/2007 Table 8:
    - Typ PP: pp_kat1_max, pp_kat2_max
    - Typ N: n_kat1_max, n_kat2_max

    Args:
        prumer_min: Průměrný čas v minutách
        typ_pracovni_polohy: "PP" nebo "N"
        category_limits: Dict s limity (pp_kat1_max, n_kat1_max, atd.)

    Returns:
        Kategorie 1, 2, nebo 3
    """
    if typ_pracovni_polohy == "PP":
        if prumer_min <= category_limits["pp_kat1_max"]:
            return 1
        elif prumer_min <= category_limits["pp_kat2_max"]:
            return 2
        else:
            return 3
    else:  # typ == "N"
        if prumer_min <= category_limits["n_kat1_max"]:
            return 1
        elif prumer_min <= category_limits["n_kat2_max"]:
            return 2
        else:
            return 3


def highlight_pp_categories(
    docx_path: str,
    measurement_data: Dict[str, Any],
    results_data: Dict[str, Any]
) -> None:
    """
    Zvýrazní čísla minut v PP tabulkách barevným backgroundem podle kategorií.

    POST-PROCESSING funkce volaná po vyrendrování Jinja2 šablony.
    Používá stejný přístup jako highlight_force_distribution_values() pro LSZ.

    Kategorie → Barvy:
    - Kategorie 1: Zelená (90EE90) - OK, bezpečné
    - Kategorie 2: Žlutá (FFFF00) - Varování
    - Kategorie 3: Červená (FFC8C8) - Kritické

    Zvýrazňuje buňky: vyskyt_min_a (col 2), vyskyt_min_b (col 3), prumer_min (col 4)

    Args:
        docx_path: Cesta k vygenerovanému Word dokumentu
        measurement_data: Data z measurement_data.json
        results_data: Data z pp_results.json

    Returns:
        None (upravuje dokument in-place)
    """
    from docx import Document

    # Vypočítej category_limits z work_duration (celková doba práce v minutách)
    section4 = measurement_data.get("section4_worker_a", {})
    work_duration_str = section4.get("work_duration")

    if work_duration_str:
        work_duration_min = int(work_duration_str)
    else:
        # Fallback - standardní směna
        work_duration_min = 480
        print(f"  ⚠ work_duration nenalezen, používám výchozích 480 min")

    # Spočítej koeficient a limity
    coefficient = _calculate_pp_work_shift_coefficient(work_duration_min)
    category_limits = _calculate_pp_category_limits(coefficient)
    print(f"  → Koeficient směny: {coefficient:.3f}, PP kat1_max: {category_limits['pp_kat1_max']}, N kat1_max: {category_limits['n_kat1_max']}")

    # Barevné schéma (hex kódy bez #)
    COLORS = {
        1: "68EE32",  # Zelená
        2: "FFC000",  # Oranžová
        3: "C00000"   # Tmavě červená
    }

    # Otevři dokument
    doc = Document(docx_path)

    total_highlighted = 0
    tables_processed = 0

    # Projdi všechny tabulky
    for table_idx, table in enumerate(doc.tables):
        # Identifikuj PP tabulky
        if not is_pp_table(table):
            continue

        print(f"  → Zpracovávám PP tabulku (index {table_idx})")
        tables_processed += 1

        # Detekuj single-worker vs two-worker
        num_cols = len(table.columns)
        is_single_worker = (num_cols == 4)

        # Určit indexy sloupců podle typu
        if is_single_worker:
            # 4 sloupce: 0=název, 1=typ_práce, 2=vyskyt_a, 3=typ_polohy
            vyskyt_a_idx = 2
            typ_polohy_idx = 3
            value_columns = [2]  # Jen vyskyt_a
        else:
            # 6 sloupců: 0=název, 1=typ_práce, 2=vyskyt_a, 3=vyskyt_b, 4=prumer, 5=typ_polohy
            vyskyt_a_idx = 2
            prumer_idx = 4
            typ_polohy_idx = 5
            value_columns = [2, 3, 4]  # vyskyt_a, vyskyt_b, prumer

        # Projdi datové řádky (skip řádek 0 = hlavička)
        for row_idx in range(1, len(table.rows)):
            row = table.rows[row_idx]

            # Bezpečné čtení hodnot
            try:
                nazev_polohy = row.cells[0].text.strip()
                typ_pracovni_polohy = row.cells[typ_polohy_idx].text.strip()  # "N" nebo "PP"
            except (IndexError, AttributeError):
                continue

            # Skip section headery (TRUP, HLAVA_KRK, atd.)
            if is_pp_section_header(nazev_polohy):
                continue

            # Skip záhlaví (obsahuje "Výskyt za směnu", "MUŽ", atd.)
            nazev_upper = nazev_polohy.upper()
            if any(keyword in nazev_upper for keyword in ["NEPŘIJATELNÁ", "PODMÍNĚNĚ", "PRACOVNÍ POLOHA", "SVALOVÁ PRÁCE"]):
                continue

            # Načti hodnotu pro výpočet kategorie
            try:
                if is_single_worker:
                    # Pro single-worker používej vyskyt_a
                    value_text = row.cells[vyskyt_a_idx].text.strip()
                else:
                    # Pro two-worker používej průměr
                    value_text = row.cells[prumer_idx].text.strip()

                value_min = float(value_text.replace(",", "."))
            except (ValueError, IndexError, AttributeError):
                continue

            # Skip prázdné řádky (value_min == 0)
            if value_min == 0 or value_min is None:
                continue

            # Skip pokud typ_pracovni_polohy není "N" nebo "PP"
            if typ_pracovni_polohy not in ["N", "PP"]:
                continue

            # Vypočítej kategorii
            category = calculate_pp_category(value_min, typ_pracovni_polohy, category_limits)

            # Zvýrazni příslušné buňky
            for col_idx in value_columns:
                try:
                    cell = row.cells[col_idx]
                    cell_text = cell.text.strip()

                    # Skip prázdné buňky
                    if not cell_text or cell_text == "0":
                        continue

                    # Nastav background podle kategorie
                    set_cell_background(cell, COLORS[category])
                    total_highlighted += 1

                except (IndexError, AttributeError):
                    continue

    # Ulož změny
    if total_highlighted > 0:
        doc.save(docx_path)
        print(f"  ✓ Zvýrazněno {total_highlighted} buněk v {tables_processed} PP tabulkách")
    else:
        print(f"  ⚠ Žádné buňky ke zvýraznění")


def highlight_lsz_category(
    docx_path: str,
    measurement_data: Dict[str, Any],
    results_data: Dict[str, Any]
) -> None:
    """
    Zvýrazní číslo kategorie v LSZ tabulce barevným textem a backgroundem.

    POST-PROCESSING funkce volaná po vyrendrování Jinja2 šablony.
    Používá stejný přístup jako highlight_force_distribution_values() a highlight_pp_categories().

    Najde tabulku podle textu "Navrhovaná kategorie:" a obarví třetí sloupec podle hodnoty.

    Kategorie → Barvy (text/background):
    - 1: černá text/zelená background - OK, nízké zatížení
    - 2: černá text/oranžová background - Varování, střední zatížení
    - 3: bílá text/červená background - Kritické, vysoké zatížení

    Tabulka má strukturu:
    - Sloupec 0: Dlouhý popisný text ("Celosměnový počet pohybů...")
    - Sloupec 1: "Navrhovaná kategorie:"
    - Sloupec 2: Číslo kategorie (1, 2, nebo 3)

    Args:
        docx_path: Cesta k vygenerovanému Word dokumentu
        measurement_data: Data z measurement_data.json
        results_data: Data z lsz_results.json

    Returns:
        None (upravuje dokument in-place)
    """
    from docx import Document
    from docx.shared import RGBColor, Pt

    # Otevři dokument
    doc = Document(docx_path)

    # Najdi tabulku s "Navrhovaná kategorie:"
    target_table = None
    target_row_idx = None

    for table in doc.tables:
        # Kontrola 3 sloupců
        if len(table.columns) != 3:
            continue

        # Hledej řádek s "Navrhovaná kategorie:"
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) >= 2:
                cell_text = row.cells[1].text.strip()
                # Flexibilní matching - různé varianty textu
                if ("kategorie" in cell_text.lower() and
                    ("navrhovaná" in cell_text.lower() or "navrhovana" in cell_text.lower())):
                    target_table = table
                    target_row_idx = row_idx
                    break

        if target_table:
            break

    if not target_table or target_row_idx is None:
        print("  ⚠ Tabulka s 'Navrhovaná kategorie' nebyla nalezena")
        return

    # Přečti hodnotu z třetího sloupce
    cell = target_table.rows[target_row_idx].cells[2]
    cell_text = cell.text.strip()

    try:
        category = int(cell_text)
    except (ValueError, TypeError):
        print(f"  ⚠ Nelze převést hodnotu kategorie na číslo: '{cell_text}'")
        return

    # Definuj barvy podle kategorie
    if category == 1:
        text_color = RGBColor(0, 0, 0)  # Černá
        bg_color = "68EE32"  # Zelená
        desc = "černá/zelená (OK)"
    elif category == 2:
        text_color = RGBColor(0, 0, 0)  # Černá
        bg_color = "FFC000"  # Oranžová
        desc = "černá/oranžová (varování)"
    elif category == 3:
        text_color = RGBColor(255, 255, 255)  # Bílá
        bg_color = "FF0000"  # Červená
        desc = "bílá/červená (kritické)"
    else:
        print(f"  ⚠ Neplatná hodnota kategorie: {category} (očekáváno 1, 2, nebo 3)")
        return

    # Smaž všechny paragraphy v buňce (docxtpl může vytvářet speciální strukturu)
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Vytvoř nový paragraph s obarvením
    new_para = cell.add_paragraph()
    new_para.alignment = 1  # Center
    run = new_para.add_run(cell_text)
    run.font.color.rgb = text_color
    run.font.size = Pt(36)
    run.font.name = "Cambria"
    run.bold = True  # Tučně pro lepší viditelnost

    # Nastav background buňky
    set_cell_background(cell, bg_color)

    # Ulož dokument
    doc.save(docx_path)
    print(f"  ✓ LSZ kategorie {category} zvýrazněna ({desc})")


def highlight_pp_category_number(
    docx_path: str,
    measurement_data: Dict[str, Any],
    results_data: Dict[str, Any]
) -> None:
    """
    Zvýrazní číslo kategorie v PP tabulce barevným textem a backgroundem.

    POST-PROCESSING funkce volaná po vyrendrování Jinja2 šablony.
    Používá stejný přístup jako highlight_lsz_category().

    Najde finální tabulku s "Navrhovaná kategorie:" a obarví třetí sloupec podle hodnoty.

    Kategorie → Barvy (text/background):
    - 1: černá text/zelená background - OK, přijatelné pracovní polohy
    - 2: černá text/oranžová background - Varování, podmíněně přijatelné
    - 3: bílá text/červená background - Kritické, nepřijatelné pracovní polohy

    Tabulka má strukturu:
    - Sloupec 0: "Výskyt nepřijatelných a podmíněně přijatelných pracovních poloh"
    - Sloupec 1: "Navrhovaná kategorie:"
    - Sloupec 2: Číslo kategorie (1, 2, nebo 3)

    Args:
        docx_path: Cesta k vygenerovanému Word dokumentu
        measurement_data: Data z measurement_data.json
        results_data: Data z pp_results.json

    Returns:
        None (upravuje dokument in-place)
    """
    from docx import Document
    from docx.shared import RGBColor, Pt

    # Otevři dokument
    doc = Document(docx_path)

    # Najdi tabulku s "Navrhovaná kategorie:" NEBO "Výskyt nepřijatelných"
    target_table = None
    target_row_idx = None

    for table in doc.tables:
        # Kontrola 3 sloupců
        if len(table.columns) != 3:
            continue

        # Hledej řádek s "Navrhovaná kategorie:" nebo "Výskyt nepřijatelných"
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) >= 2:
                # Check sloupec 0 (může obsahovat "Výskyt nepřijatelných")
                cell0_text = row.cells[0].text.strip().lower()
                # Check sloupec 1 (měl by obsahovat "Navrhovaná kategorie:")
                cell1_text = row.cells[1].text.strip().lower()

                # Hledáme řádek kde:
                # - Sloupec 0 obsahuje "výskyt nepřijatelných" nebo
                # - Sloupec 1 obsahuje "navrhovaná kategorie"
                if (("výskyt" in cell0_text and "nepřijatelných" in cell0_text) or
                    ("vyskyt" in cell0_text and "neprijatelnych" in cell0_text) or
                    ("kategorie" in cell1_text and ("navrhovaná" in cell1_text or "navrhovana" in cell1_text))):
                    target_table = table
                    target_row_idx = row_idx
                    break

        if target_table:
            break

    if not target_table or target_row_idx is None:
        print("  ⚠ PP tabulka s 'Navrhovaná kategorie' nebyla nalezena")
        return

    # Přečti hodnotu z třetího sloupce
    cell = target_table.rows[target_row_idx].cells[2]
    cell_text = cell.text.strip()

    try:
        category = int(cell_text)
    except (ValueError, TypeError):
        print(f"  ⚠ Nelze převést PP kategorii na číslo: '{cell_text}'")
        return

    # Definuj barvy podle kategorie (stejné jako u LSZ)
    if category == 1:
        text_color = RGBColor(0, 0, 0)  # Černá
        bg_color = "68EE32"  # Zelená
        desc = "černá/zelená (OK)"
    elif category == 2:
        text_color = RGBColor(0, 0, 0)  # Černá
        bg_color = "FFC000"  # Oranžová
        desc = "černá/oranžová (varování)"
    elif category == 3:
        text_color = RGBColor(255, 255, 255)  # Bílá
        bg_color = "FF0000"  # Červená
        desc = "bílá/červená (kritické)"
    else:
        print(f"  ⚠ Neplatná hodnota PP kategorie: {category} (očekáváno 1, 2, nebo 3)")
        return

    # Smaž všechny paragraphy v buňce (docxtpl může vytvářet speciální strukturu)
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Vytvoř nový paragraph s obarvením
    new_para = cell.add_paragraph()
    new_para.alignment = 1  # Center
    run = new_para.add_run(cell_text)
    run.font.color.rgb = text_color
    run.font.size = Pt(36)
    run.font.name = "Cambria"
    run.bold = True  # Tučně pro lepší viditelnost

    # Nastav background buňky
    set_cell_background(cell, bg_color)

    # Ulož dokument
    doc.save(docx_path)
    print(f"  ✓ PP kategorie {category} zvýrazněna ({desc})")
