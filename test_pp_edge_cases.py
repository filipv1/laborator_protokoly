"""
Test edge cases pro odstranění prázdných PP tabulek
"""
from docx import Document
from generate_word_from_two_sources import remove_empty_pp_rows


def create_edge_case_document():
    """Vytvoří dokument s různými edge cases"""
    doc = Document()

    doc.add_heading("Test Edge Cases - Odstranění PP tabulek", 0)

    # EDGE CASE 1: Tabulka pouze se záhlavím a OSTATNI (žádné datové řádky)
    doc.add_paragraph("Edge Case 1: Pouze záhlaví + OSTATNI")
    table1 = doc.add_table(rows=3, cols=6)
    # Záhlaví
    table1.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table1.rows[0].cells[1].text = "Svalová práce"
    table1.rows[0].cells[2].text = "MUŽ 1"
    table1.rows[0].cells[3].text = "MUŽ 2"
    table1.rows[0].cells[4].text = "Ø"
    table1.rows[0].cells[5].text = "Typ"
    # Section
    table1.rows[1].cells[0].text = "LHK"
    # OSTATNI
    table1.rows[2].cells[0].text = "OSTATNI"
    table1.rows[2].cells[1].text = "Ultrathink"

    # EDGE CASE 2: Tabulka se směsicí nul (0, 0.0, prázdné)
    doc.add_paragraph("\nEdge Case 2: Směs různých typů nul")
    table2 = doc.add_table(rows=5, cols=6)
    # Záhlaví
    table2.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table2.rows[0].cells[1].text = "Svalová práce"
    table2.rows[0].cells[2].text = "MUŽ 1"
    table2.rows[0].cells[3].text = "MUŽ 2"
    table2.rows[0].cells[4].text = "Ø"
    table2.rows[0].cells[5].text = "Typ"
    # Section
    table2.rows[1].cells[0].text = "DOLNÍ KONČETINY"
    # Data řádky - různé typy nul
    table2.rows[2].cells[0].text = "Poloha 1"
    table2.rows[2].cells[1].text = "Statická"
    table2.rows[2].cells[2].text = "0"     # string nula
    table2.rows[2].cells[3].text = ""      # prázdné
    table2.rows[2].cells[4].text = "0.0"   # desetinná nula
    table2.rows[2].cells[5].text = "N"

    table2.rows[3].cells[0].text = "Poloha 2"
    table2.rows[3].cells[1].text = "Dynamická"
    table2.rows[3].cells[2].text = ""      # prázdné
    table2.rows[3].cells[3].text = "0"     # string nula
    table2.rows[3].cells[4].text = ""      # prázdné
    table2.rows[3].cells[5].text = "N"

    table2.rows[4].cells[0].text = "OSTATNI"
    table2.rows[4].cells[1].text = "Ultrathink"

    # EDGE CASE 3: Tabulka s jedním nenulovým řádkem
    doc.add_paragraph("\nEdge Case 3: Jeden nenulový řádek")
    table3 = doc.add_table(rows=6, cols=6)
    # Záhlaví
    table3.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table3.rows[0].cells[1].text = "Svalová práce"
    table3.rows[0].cells[2].text = "MUŽ 1"
    table3.rows[0].cells[3].text = "MUŽ 2"
    table3.rows[0].cells[4].text = "Ø"
    table3.rows[0].cells[5].text = "Typ"
    # Section
    table3.rows[1].cells[0].text = "OSTATNÍ ČÁSTI TĚLA"
    # Data řádky
    table3.rows[2].cells[0].text = "Nulová poloha"
    table3.rows[2].cells[1].text = "Statická"
    table3.rows[2].cells[2].text = "0"
    table3.rows[2].cells[3].text = "0"
    table3.rows[2].cells[4].text = "0"
    table3.rows[2].cells[5].text = "N"

    table3.rows[3].cells[0].text = "Aktivní poloha"
    table3.rows[3].cells[1].text = "Dynamická"
    table3.rows[3].cells[2].text = "0.1"   # Malá nenulová hodnota!
    table3.rows[3].cells[3].text = "0"
    table3.rows[3].cells[4].text = "0.05"
    table3.rows[3].cells[5].text = "PP"

    table3.rows[4].cells[0].text = "Další nulová"
    table3.rows[4].cells[1].text = "Statická"
    table3.rows[4].cells[2].text = "0"
    table3.rows[4].cells[3].text = "0"
    table3.rows[4].cells[4].text = "0"
    table3.rows[4].cells[5].text = "N"

    table3.rows[5].cells[0].text = "OSTATNI"
    table3.rows[5].cells[1].text = "Ultrathink"

    # EDGE CASE 4: Tabulka s více section headers
    doc.add_paragraph("\nEdge Case 4: Více section headers")
    table4 = doc.add_table(rows=7, cols=6)
    # Záhlaví
    table4.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table4.rows[0].cells[1].text = "Svalová práce"
    table4.rows[0].cells[2].text = "MUŽ 1"
    table4.rows[0].cells[3].text = "MUŽ 2"
    table4.rows[0].cells[4].text = "Ø"
    table4.rows[0].cells[5].text = "Typ"
    # První section
    table4.rows[1].cells[0].text = "TRUP"
    # Data
    table4.rows[2].cells[0].text = "Poloha trupu"
    table4.rows[2].cells[1].text = "Statická"
    table4.rows[2].cells[2].text = "0"
    table4.rows[2].cells[3].text = "0"
    table4.rows[2].cells[4].text = "0"
    table4.rows[2].cells[5].text = "N"
    # Druhá section
    table4.rows[3].cells[0].text = "PHK"
    # Data
    table4.rows[4].cells[0].text = "Poloha PHK"
    table4.rows[4].cells[1].text = "Dynamická"
    table4.rows[4].cells[2].text = "0"
    table4.rows[4].cells[3].text = "0"
    table4.rows[4].cells[4].text = "0"
    table4.rows[4].cells[5].text = "N"
    # Třetí section
    table4.rows[5].cells[0].text = "LHK"
    # OSTATNI
    table4.rows[6].cells[0].text = "OSTATNI"
    table4.rows[6].cells[1].text = "Ultrathink"

    doc.save("test_pp_edge_cases.docx")
    print("✓ Vytvořen dokument s edge cases: test_pp_edge_cases.docx")
    print("  - Edge Case 1: Pouze záhlaví + OSTATNI (měla by být odstraněna)")
    print("  - Edge Case 2: Směs různých typů nul (měla by být odstraněna)")
    print("  - Edge Case 3: Jeden nenulový řádek (neměla by být odstraněna)")
    print("  - Edge Case 4: Více section headers, všechny nulové (měla by být odstraněna)")


def analyze_tables(docx_path):
    """Analyzuje tabulky v dokumentu"""
    doc = Document(docx_path)
    pp_count = 0
    for i, table in enumerate(doc.tables):
        if len(table.columns) == 6:
            pp_count += 1
            # Spočítej datové řádky
            data_rows = 0
            for row in table.rows[1:]:  # Skip header
                text = row.cells[0].text.strip().upper()
                if text not in ["TRUP", "HLAVA_KRK", "HLAVA A KRK", "PHK", "LHK",
                               "DOLNÍ KONČETINY", "DKK", "OSTATNÍ ČÁSTI TĚLA",
                               "OSTATNI", "OST", "ULTRATHINK"]:
                    if "OSTATNI" not in text and "ULTRATHINK" not in text.upper():
                        data_rows += 1
            print(f"  Tabulka {i}: {data_rows} datových řádků")
    return pp_count


def test_edge_cases():
    """Test edge cases"""
    print("\n" + "="*60)
    print("TEST EDGE CASES - ODSTRANĚNÍ PP TABULEK")
    print("="*60)

    # Vytvoř dokument
    create_edge_case_document()

    test_path = "test_pp_edge_cases.docx"

    # Analýza před
    print(f"\nAnalýza dokumentu PŘED úpravou:")
    tables_before = analyze_tables(test_path)
    print(f"Celkem PP tabulek: {tables_before}")

    # Spusť funkci
    print("\nSpouštím remove_empty_pp_rows()...")
    print("-"*60)
    remove_empty_pp_rows(test_path)
    print("-"*60)

    # Analýza po
    print(f"\nAnalýza dokumentu PO úpravě:")
    tables_after = analyze_tables(test_path)
    print(f"Celkem PP tabulek: {tables_after}")

    # Vyhodnocení
    expected_removed = 3  # Edge cases 1, 2, 4
    expected_remaining = 1  # Edge case 3

    if tables_before == 4 and tables_after == expected_remaining:
        print("\n✓✓✓ ÚSPĚCH! Všechny edge cases prošly správně:")
        print(f"  - Odstraněny {expected_removed} prázdné tabulky")
        print(f"  - Zachována {expected_remaining} tabulka s daty")
    else:
        print(f"\n✗✗✗ CHYBA! Neočekávaný výsledek:")
        print(f"  - Očekáváno: 4 tabulky → {expected_remaining} tabulka")
        print(f"  - Skutečnost: {tables_before} tabulek → {tables_after} tabulek")

    print(f"\nVýsledný dokument: {test_path}")


if __name__ == "__main__":
    test_edge_cases()