"""Analyzuj placeholdery v Word šabloně"""
import re
from docx import Document

doc = Document(r'Vzorové protokoly/Autorizované protokoly pro MUŽE/lsz_placeholdery_v2.docx')

placeholders = set()

# Procházej všechny paragrafy
for para in doc.paragraphs:
    text = para.text
    if '{{' in text and '}}' in text:
        matches = re.findall(r'\{\{[^}]+\}\}', text)
        for match in matches:
            placeholders.add(match.strip())

# Procházej tabulky
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text = cell.text
            if '{{' in text and '}}' in text:
                matches = re.findall(r'\{\{[^}]+\}\}', text)
                for match in matches:
                    placeholders.add(match.strip())

# Vypiš do souboru
with open('word_placeholders_analysis.txt', 'w', encoding='utf-8') as f:
    f.write(f'Celkový počet unikátních placeholderů: {len(placeholders)}\n\n')
    f.write('Seznam všech placeholderů:\n')
    f.write('='*50 + '\n')
    for i, ph in enumerate(sorted(placeholders), 1):
        f.write(f'{i}. {ph}\n')

print('Analýza dokončena - výsledky v word_placeholders_analysis.txt')
