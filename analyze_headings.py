"""
Analyzuje všechny nadpisy v Word dokumentu
"""
import sys
from docx import Document

# Fix Windows encoding
sys.stdout.reconfigure(encoding='utf-8')

if len(sys.argv) < 2:
    print("Pouziti: python analyze_headings.py <cesta_k_docx>")
    sys.exit(1)

docx_path = sys.argv[1]
doc = Document(docx_path)

print("=" * 80)
print(f"ANALYZA NADPISU: {docx_path}")
print("=" * 80)
print()

print(f"Celkem odstavců: {len(doc.paragraphs)}")
print()

print("--- HLEDÁNÍ NADPISŮ S 'ČASOV' ---")
found = False
for i, p in enumerate(doc.paragraphs):
    if 'ČASOV' in p.text.upper() or 'CASOV' in p.text.upper():
        print(f"Řádek {i}: '{p.text}'")
        print(f"  Style: {p.style.name}")
        print(f"  Text length: {len(p.text)}")
        print(f"  Exact bytes: {repr(p.text)}")
        print()
        found = True

if not found:
    print("❌ Žádný nadpis s 'ČASOV' nenalezen!")
    print()
    print("--- VŠECHNY NADPISY (obsahující velká písmena) ---")
    for i, p in enumerate(doc.paragraphs):
        # Hledej odstavce co jsou převážně velká písmena
        if p.text and len(p.text) > 5 and sum(1 for c in p.text if c.isupper()) > len(p.text) * 0.5:
            print(f"Řádek {i}: '{p.text[:80]}'")

print()
print("=" * 80)
print("HOTOVO")
print("=" * 80)
