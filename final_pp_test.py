"""
Finální test PP removal funkce
"""
from docx import Document
from pathlib import Path
from generate_word_from_two_sources import remove_empty_pp_rows

# Vytvoř realistický PP protokol
doc = Document()

# TABULKA 1: Hlavní tabulka s PP daty
table = doc.add_table(rows=14, cols=6)

# Hlavička
table.rows[0].cells[0].text = 'Nepřijatelná/podmíněně přijatelná pracovní poloha'
table.rows[0].cells[1].text = 'Svalová práce'
table.rows[0].cells[2].text = 'MUŽ 1\nVýskyt za směnu\n[min]'
table.rows[0].cells[3].text = 'MUŽ 2\nVýskyt za směnu\n[min]'
table.rows[0].cells[4].text = 'Ø\nvýskyt za směnu\n[min]'
table.rows[0].cells[5].text = 'Typ pracovní polohy'

# TRUP sekce
table.rows[1].cells[0].text = 'TRUP'
table.rows[2].cells[0].text = 'Předklon trupu větší než 60°'
table.rows[2].cells[1].text = 'Statická'
table.rows[2].cells[2].text = '0'
table.rows[2].cells[3].text = '0'
table.rows[2].cells[4].text = '0'
table.rows[2].cells[5].text = 'N'

table.rows[3].cells[0].text = 'Záklon bez opory celého těla'
table.rows[3].cells[1].text = 'Statická'
table.rows[3].cells[2].text = '3'
table.rows[3].cells[3].text = '5'
table.rows[3].cells[4].text = '4'
table.rows[3].cells[5].text = 'N'

table.rows[4].cells[0].text = 'Výrazný úklon či pootočení trupu větší než 20°'
table.rows[4].cells[1].text = 'Statická'
table.rows[4].cells[2].text = '0'
table.rows[4].cells[3].text = '5'
table.rows[4].cells[4].text = '2.5'
table.rows[4].cells[5].text = 'N'

table.rows[5].cells[0].text = 'Předklon trupu větší než 60° při frekvenci pohybů větší nebo rovné 2/min'
table.rows[5].cells[1].text = 'Dynamická'
table.rows[5].cells[2].text = '0'
table.rows[5].cells[3].text = '5'
table.rows[5].cells[4].text = '2.5'
table.rows[5].cells[5].text = 'N'

table.rows[6].cells[0].text = 'Záklon trupu s oporou těla'
table.rows[6].cells[1].text = 'Statická'
table.rows[6].cells[2].text = '0'
table.rows[6].cells[3].text = '0'
table.rows[6].cells[4].text = '0'
table.rows[6].cells[5].text = 'PP'

# HLAVA_KRK sekce
table.rows[7].cells[0].text = 'HLAVA_KRK'
table.rows[8].cells[0].text = 'Předklon hlavy větší než 25° bez podpory trupu'
table.rows[8].cells[1].text = 'Statická'
table.rows[8].cells[2].text = '0'
table.rows[8].cells[3].text = '0'
table.rows[8].cells[4].text = '0'
table.rows[8].cells[5].text = 'N'

table.rows[9].cells[0].text = 'Záklon hlavy bez podpory celé hlavy'
table.rows[9].cells[1].text = 'Statická'
table.rows[9].cells[2].text = '150'
table.rows[9].cells[3].text = '40'
table.rows[9].cells[4].text = '95'
table.rows[9].cells[5].text = 'N'

table.rows[10].cells[0].text = 'Úklon a rotace hlavy větší než 15°'
table.rows[10].cells[1].text = 'Statická'
table.rows[10].cells[2].text = '0'
table.rows[10].cells[3].text = '0'
table.rows[10].cells[4].text = '0'
table.rows[10].cells[5].text = 'N'

# PHK sekce
table.rows[11].cells[0].text = 'PHK'
table.rows[12].cells[0].text = 'Vzpažení paže 40 až 60°, jestliže paže není podepřena'
table.rows[12].cells[1].text = 'Statická'
table.rows[12].cells[2].text = '80'
table.rows[12].cells[3].text = '60'
table.rows[12].cells[4].text = '70'
table.rows[12].cells[5].text = 'PP'

table.rows[13].cells[0].text = 'Vzpažení paže větší než 60°'
table.rows[13].cells[1].text = 'Statická'
table.rows[13].cells[2].text = '0'
table.rows[13].cells[3].text = '0'
table.rows[13].cells[4].text = '0'
table.rows[13].cells[5].text = 'N'

# Ulož test soubor
test_file = Path('final_pp_test_before.docx')
doc.save(str(test_file))
print(f'Test file created: {test_file}')
print(f'Rows before: {len(table.rows)}')

# Očekávané řádky k odstranění (všechny mají 0,0,0):
print('\nExpected to be removed (all have 0,0,0):')
print('  - Row 2: Předklon trupu větší než 60°')
print('  - Row 6: Záklon trupu s oporou těla')
print('  - Row 8: Předklon hlavy větší než 25° bez podpory trupu')
print('  - Row 10: Úklon a rotace hlavy větší než 15°')
print('  - Row 13: Vzpažení paže větší než 60°')

# Aplikuj odstranění
print('\n' + '='*60)
print('Applying remove_empty_pp_rows()...')
print('='*60)
remove_empty_pp_rows(str(test_file))
print('='*60)

# Ulož výsledek
output_file = Path('final_pp_test_after.docx')
test_file.rename(output_file)

# Zkontroluj výsledek
doc_after = Document(str(output_file))
table_after = doc_after.tables[0]
print(f'\nRows after: {len(table_after.rows)}')
print(f'Removed: {len(table.rows) - len(table_after.rows)} rows')

print('\nRemaining rows:')
for i, row in enumerate(table_after.rows):
    text = row.cells[0].text[:50].encode('ascii', 'replace').decode('ascii')
    print(f'  Row {i}: {text}')

print('\nSUCCESS: All empty rows removed!' if len(table_after.rows) == 9 else 'PROBLEM: Not all rows removed!')