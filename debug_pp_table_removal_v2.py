"""
Debug skript pro kontrolu PP tabulek v dokumentu - v2
"""
import sys
from docx import Document
from pathlib import Path

# Force UTF-8 encoding
sys.stdout.reconfigure(encoding='utf-8')


def safe_print(text):
    """Bezpečný print pro Windows konzoli"""
    try:
        print(text)
    except UnicodeEncodeError:
        # Fallback na ASCII
        safe_text = text.encode('ascii', 'replace').decode('ascii')
        print(safe_text)


def analyze_pp_document(docx_path):
    """Analyzuje PP dokument a vypíše detailní info o tabulkách"""
    doc = Document(docx_path)

    pp_keywords = [
        "TRUP", "Trup",
        "HLAVA_KRK", "HLAVA A KRK", "Hlava a krk",
        "PHK", "Pravá horní končetina",
        "LHK", "Levá horní končetina",
        "DOLNÍ KONČETINY", "Dolní končetiny", "DKK",
        "OSTATNÍ ČÁSTI TĚLA", "Ostatní části těla", "OST",
        "Statická", "Dynamická"
    ]

    section_keywords = ["TRUP", "HLAVA_KRK", "HLAVA A KRK", "PHK", "LHK",
                       "DOLNÍ KONČETINY", "DKK", "OSTATNÍ ČÁSTI TĚLA", "OSTATNI", "OST"]

    safe_print(f"Analyzuji dokument: {docx_path}")
    safe_print(f"Celkem tabulek: {len(doc.tables)}")
    safe_print("="*60)

    for table_idx, table in enumerate(doc.tables):
        safe_print(f"\nTABULKA {table_idx}:")
        safe_print(f"  Pocet sloupcu: {len(table.columns)}")
        safe_print(f"  Pocet radku: {len(table.rows)}")

        if len(table.columns) == 6:
            # Zkontroluj PP keywords
            is_pp_table = False
            for row in table.rows[:5]:
                row_text = " ".join(cell.text for cell in row.cells)
                for keyword in pp_keywords:
                    if keyword in row_text:
                        is_pp_table = True
                        break
                if is_pp_table:
                    break

            if is_pp_table:
                safe_print("  -> Je to PP tabulka!")

                # Spočítej datové řádky
                data_rows = 0
                empty_rows = 0
                section_rows = 0
                ostatni_rows = 0

                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0:  # Skip header
                        continue

                    try:
                        nazev = row.cells[0].text.strip()
                        nazev_upper = nazev.upper()

                        # Je to sekční řádek?
                        if nazev_upper in section_keywords:
                            section_rows += 1
                            safe_print(f"    Radek {row_idx}: [SECTION] {nazev}")
                            continue

                        # Je to OSTATNI?
                        if "OSTATNI" in nazev_upper or "ULTRATHINK" in nazev_upper:
                            ostatni_rows += 1
                            safe_print(f"    Radek {row_idx}: [OSTATNI] {nazev}")
                            continue

                        # Zkontroluj hodnoty
                        vyskyt_a = row.cells[2].text.strip() if len(row.cells) > 2 else ""
                        vyskyt_b = row.cells[3].text.strip() if len(row.cells) > 3 else ""
                        prumer = row.cells[4].text.strip() if len(row.cells) > 4 else ""

                        is_empty = (vyskyt_a in ["0", "0.0", "", "None"] and
                                   vyskyt_b in ["0", "0.0", "", "None"] and
                                   prumer in ["0", "0.0", "", "None"])

                        if is_empty:
                            empty_rows += 1
                            safe_print(f"    Radek {row_idx}: [EMPTY] {nazev[:40]}... [{vyskyt_a},{vyskyt_b},{prumer}]")
                        else:
                            data_rows += 1
                            safe_print(f"    Radek {row_idx}: [DATA] {nazev[:40]}... [{vyskyt_a},{vyskyt_b},{prumer}]")

                    except Exception as e:
                        safe_print(f"    Radek {row_idx}: [ERROR] {str(e)}")

                safe_print(f"  SUMARIZACE:")
                safe_print(f"    - Datove radky (s hodnotami): {data_rows}")
                safe_print(f"    - Prazdne radky (vse 0): {empty_rows}")
                safe_print(f"    - Sekcni radky: {section_rows}")
                safe_print(f"    - OSTATNI radky: {ostatni_rows}")

                if data_rows == 0 and empty_rows == 0:
                    safe_print("  !!! TABULKA NEMA ZADNE DATOVE RADKY - MELA BY BYT ODSTRANENA !!!")
                elif data_rows == 0 and empty_rows > 0:
                    safe_print("  !!! TABULKA MA JEN PRAZDNE RADKY - PO JEJICH ODSTRANENI BY MELA BYT CELA SMAZANA !!!")


def main():
    # Analyzuj aktuální PP dokument
    pp_path = Path("projects/325_Ascari_sro/PP_325_Ascari_sro_CAS_protokol.docx")
    if pp_path.exists():
        analyze_pp_document(pp_path)
    else:
        safe_print(f"Soubor neexistuje: {pp_path}")


if __name__ == "__main__":
    main()