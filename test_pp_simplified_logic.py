"""
Test zjednodušené logiky PP table removal
"""
from docx import Document
from generate_word_from_two_sources import remove_empty_pp_rows


def create_realistic_pp_table():
    """Vytvoří realistickou PP tabulku jak vypadá po vyplnění placeholderů"""
    doc = Document()

    doc.add_heading("Realistický PP protokol", 0)

    # Tabulka 1: Všechny hodnoty nulové (měla by být odstraněna)
    doc.add_paragraph("Tabulka 1: TRUP - všechny nuly")
    table1 = doc.add_table(rows=6, cols=6)

    # Záhlaví
    table1.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table1.rows[0].cells[1].text = "Svalová práce"
    table1.rows[0].cells[2].text = "MUŽ 1"
    table1.rows[0].cells[3].text = "MUŽ 2"
    table1.rows[0].cells[4].text = "Ø"
    table1.rows[0].cells[5].text = "Typ"

    # Sekce
    table1.rows[1].cells[0].text = "TRUP"

    # Datové řádky - všechny nulové
    table1.rows[2].cells[0].text = "Předklon trupu větší než 60°"
    table1.rows[2].cells[1].text = "Statická"
    table1.rows[2].cells[2].text = "0"
    table1.rows[2].cells[3].text = "0"
    table1.rows[2].cells[4].text = "0"
    table1.rows[2].cells[5].text = "N"

    table1.rows[3].cells[0].text = "Záklon trupu"
    table1.rows[3].cells[1].text = "Dynamická"
    table1.rows[3].cells[2].text = "0"
    table1.rows[3].cells[3].text = "0"
    table1.rows[3].cells[4].text = "0"
    table1.rows[3].cells[5].text = "N"

    table1.rows[4].cells[0].text = "Výrazný úklon trupu"
    table1.rows[4].cells[1].text = "Statická"
    table1.rows[4].cells[2].text = "0"
    table1.rows[4].cells[3].text = "0"
    table1.rows[4].cells[4].text = "0"
    table1.rows[4].cells[5].text = "N"

    # OSTATNI
    table1.rows[5].cells[0].text = "OSTATNI"
    table1.rows[5].cells[1].text = "Ultrathink"

    # Tabulka 2: Některé hodnoty nenulové (neměla by být odstraněna)
    doc.add_paragraph("\nTabulka 2: PHK - některé hodnoty nenulové")
    table2 = doc.add_table(rows=6, cols=6)

    # Záhlaví
    table2.rows[0].cells[0].text = "Nepřijatelná/podmíněně přijatelná pracovní poloha"
    table2.rows[0].cells[1].text = "Svalová práce"
    table2.rows[0].cells[2].text = "MUŽ 1"
    table2.rows[0].cells[3].text = "MUŽ 2"
    table2.rows[0].cells[4].text = "Ø"
    table2.rows[0].cells[5].text = "Typ"

    # Sekce
    table2.rows[1].cells[0].text = "PHK"

    # Datové řádky - mix
    table2.rows[2].cells[0].text = "Vzpažení paže větší než 60°"
    table2.rows[2].cells[1].text = "Statická"
    table2.rows[2].cells[2].text = "15"  # Nenulové!
    table2.rows[2].cells[3].text = "10"
    table2.rows[2].cells[4].text = "12.5"
    table2.rows[2].cells[5].text = "PP"

    table2.rows[3].cells[0].text = "Zapažení"
    table2.rows[3].cells[1].text = "Dynamická"
    table2.rows[3].cells[2].text = "0"
    table2.rows[3].cells[3].text = "0"
    table2.rows[3].cells[4].text = "0"
    table2.rows[3].cells[5].text = "N"

    table2.rows[4].cells[0].text = "Elevace paže"
    table2.rows[4].cells[1].text = "Statická"
    table2.rows[4].cells[2].text = "0"
    table2.rows[4].cells[3].text = "0"
    table2.rows[4].cells[4].text = "0"
    table2.rows[4].cells[5].text = "N"

    # OSTATNI
    table2.rows[5].cells[0].text = "OSTATNI"
    table2.rows[5].cells[1].text = "Ultrathink"

    doc.save("test_pp_realistic.docx")
    print("✓ Vytvořen realistický testovací dokument: test_pp_realistic.docx")
    return "test_pp_realistic.docx"


def test_simplified_logic():
    """Test zjednodušené logiky"""
    print("\n" + "="*60)
    print("TEST ZJEDNODUŠENÉ LOGIKY")
    print("="*60)

    test_path = create_realistic_pp_table()

    # Počet tabulek před
    doc = Document(test_path)
    tables_before = len(doc.tables)
    print(f"Počet tabulek před: {tables_before}")

    print("\nSpouštím remove_empty_pp_rows()...")
    print("-"*60)
    remove_empty_pp_rows(test_path)
    print("-"*60)

    # Počet tabulek po
    doc = Document(test_path)
    tables_after = len(doc.tables)
    print(f"Počet tabulek po: {tables_after}")

    if tables_before == 2 and tables_after == 1:
        print("\n✓✓✓ ÚSPĚCH!")
        print("  - Tabulka 1 (všechny nuly) byla odstraněna")
        print("  - Tabulka 2 (některé nenulové) zůstala")
        return True
    else:
        print(f"\n✗✗✗ CHYBA!")
        print(f"  Očekáváno: 2 → 1")
        print(f"  Skutečnost: {tables_before} → {tables_after}")
        return False


if __name__ == "__main__":
    result = test_simplified_logic()
    print("="*60)
    print(f"VÝSLEDEK: {'✓ PASS' if result else '✗ FAIL'}")
    print("="*60)