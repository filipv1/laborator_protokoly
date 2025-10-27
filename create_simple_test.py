"""Jednoduchý test Word generování se dvěma JSONy"""
import json
from docx import Document
from docxtpl import DocxTemplate

# Vytvoř jednoduchou testovací šablonu
doc = Document()
doc.add_heading('TEST - DVA ZDROJE DAT', 0)

doc.add_heading('VSTUPNI DATA (input)', 1)
doc.add_paragraph('Firma: {{ input.section1_firma.company }}')
doc.add_paragraph('ICO: {{ input.section1_firma.ico }}')
doc.add_paragraph('Datum: {{ input.section1_firma.measurement_date }}')

doc.add_heading('VYSLEDKOVA DATA (results)', 1)
doc.add_paragraph('Fmax PHK Extenzor: {{ results.Fmax_Phk_Extenzor }} N')
doc.add_paragraph('Fmax PHK Flexor: {{ results.Fmax_Phk_Flexor }} N')
doc.add_paragraph('Pocet pohybu PHK: {{ results.phk_number_of_movements }}')

doc.add_heading('TABULKA SOMATOMETRIE', 1)
doc.add_paragraph('{% for row in results.table_somatometrie %}')
doc.add_paragraph('Radek {{ loop.index }}: Datum={{ row.datum }}, Vek={{ row.vek_roky }} let, Vyska={{ row.vyska_cm }} cm')
doc.add_paragraph('{% endfor %}')

# Ulož šablonu
template_path = 'simple_test_template.docx'
doc.save(template_path)

# Načti data
with open('measurement_data_example.json', encoding='utf-8') as f:
    input_data = json.load(f)

with open('lsz_results.json', encoding='utf-8') as f:
    results_data = json.load(f)

# Vytvoř kontext s vnořenou strukturou
context = {
    'input': input_data,
    'results': results_data
}

# Vygeneruj Word
doc = DocxTemplate(template_path)
doc.render(context)
output_path = 'SIMPLE_TEST_OUTPUT.docx'
doc.save(output_path)

print(f'Generated: {output_path}')
print(f'Input keys: {list(input_data.keys())}')
print(f'Results keys: {list(results_data.keys())[:10]}...')
print(f'Somatometrie rows: {len(results_data["table_somatometrie"])}')
