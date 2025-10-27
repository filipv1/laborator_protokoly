"""
Overit cervene barvy v Word dokumentu
"""
import sys
from docx import Document
from docx.shared import RGBColor

if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')

def verify_red_colors(docx_path):
    """Overi cervene barvy v dokumentu"""

    print("=" * 80)
    print(f"OVERENI CERVENYCH BAREV: {docx_path}")
    print("=" * 80)
    print()

    doc = Document(docx_path)

    # Najdi tabulku s force distribution
    target_table = None
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue

        # Zkontroluj hlavicku
        header_row = table.rows[0]
        header_text = " ".join([cell.text for cell in header_row.cells]).lower()

        if "extenzor" in header_text or "flexor" in header_text:
            target_table = table
            print(f"Nalezena tabulka force_distribution (tabulka #{table_idx})")
            print(f"Pocet radku: {len(table.rows)}")
            print(f"Pocet sloupcu: {len(table.rows[0].cells)}")
            print()
            break

    if target_table is None:
        print("ERROR: Tabulka force_distribution nenalezena!")
        return

    # Projdi vsechny radky a zkontroluj cervene barvy
    print("ANALYZA BAREV:")
    print("-" * 80)

    red_found = False

    for row_idx, row in enumerate(target_table.rows):
        row_activity = row.cells[0].text.strip()

        # Zobraz jen radky s daty (ne prazdne)
        if not row_activity or row_idx == 0:  # Skip hlavicku
            continue

        print(f"Radek {row_idx}: {row_activity}")

        # Zkontroluj kazdy sloupec
        for col_idx, cell in enumerate(row.cells[1:], start=1):  # Skip activity sloupec
            cell_value = cell.text.strip()

            if not cell_value or cell_value == '0':
                continue

            # Zkontroluj barvy vsech runs
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.color.rgb:
                        rgb = run.font.color.rgb

                        # Je cervena? (RGB 255, 0, 0)
                        if rgb == RGBColor(255, 0, 0):
                            print(f"  Sloupec {col_idx}: {cell_value} -> CERVENA ✓")
                            red_found = True
                        else:
                            print(f"  Sloupec {col_idx}: {cell_value} -> RGB({rgb[0]}, {rgb[1]}, {rgb[2]})")
                    else:
                        # Pokud neni nastavena barva, je to automaticka (cerna)
                        if cell_value not in ['0', '']:
                            print(f"  Sloupec {col_idx}: {cell_value} -> zadna barva (automaticka - cerna)")

        print()

    print("-" * 80)
    if red_found:
        print("VYSLEDEK: CERVENE HODNOTY NALEZENY! ✓")
    else:
        print("VYSLEDEK: ZADNE CERVENE HODNOTY! ✗")
    print("=" * 80)


if __name__ == "__main__":
    verify_red_colors("test_force_output_CERVENE.docx")
