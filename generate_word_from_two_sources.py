"""
Word Protocol Generator - načítá data ze dvou oddělených JSONů
- measurement_data.json = vstupní data (GUI wizard)
- lsz_results.json = výsledková data (načtená z Excel)
"""
import json
import sys
import argparse
from datetime import datetime
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx import Document
from pathlib import Path
from PIL import Image
from jinja2 import Environment
from core.text_generator import generate_conditional_texts, get_selected_holter_numbers, highlight_selected_holters, highlight_force_distribution_values

# Fix Windows console encoding
sys.stdout.reconfigure(encoding='utf-8')


def normalize_png(png_path):
    """
    Konvertuje PNG do standardního formátu, který python-docx rozpozná.
    xlwings občas exportuje PNG v nestandardním formátu.
    """
    try:
        img = Image.open(png_path)
        # Konvertuj do RGB (pokud je RGBA nebo jiný režim)
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')
        # Ulož zpět jako standardní PNG
        img.save(png_path, 'PNG')
        return True
    except Exception as e:
        print(f"⚠ Chyba při normalizaci PNG: {e}")
        return False


def format_czech_number(value):
    """
    Univerzální formátování čísel v českém formátu.

    Automaticky detekuje typ čísla a aplikuje správné formátování:
    - Desetinná čísla: zaokrouhlí na 1 des. místo, čárka místo tečky, mezery pro tisíce
    - Celá čísla: pouze mezery pro tisíce

    Args:
        value: Číslo k formátování (float, int, string nebo None)

    Returns:
        str: Formátované číslo v českém formátu

    Pravidla:
        - 1-999: bez mezer (např. "150", "8,5")
        - 1000+: mezery každé 3 cifry zprava (např. "2 222", "12 345,7")
        - Desetinná: vždy 1 des. místo s čárkou (např. "8,6", "12 345,7")
        - Celá: bez desetinných míst (např. "2 222", "450")

    Příklady:
        >>> format_czech_number(8.55)
        "8,6"
        >>> format_czech_number(11.899999)
        "11,9"
        >>> format_czech_number(450)
        "450"
        >>> format_czech_number(2222)
        "2 222"
        >>> format_czech_number(12345.67)
        "12 345,7"
        >>> format_czech_number(123456789.123)
        "123 456 789,1"
    """
    if value is None:
        return ""

    try:
        # Konvertuj na float
        num = float(value)

        # Detekuj, jestli je to prakticky celé číslo
        # (tolerance pro float nepřesnosti: 0.0001)
        is_integer = abs(num - round(num)) < 0.0001

        if is_integer:
            # CELÉ ČÍSLO: žádná desetinná místa
            num_int = int(round(num))
            num_str = str(abs(num_int))  # abs pro handling záporných čísel
            decimal_part = None
        else:
            # DESETINNÉ ČÍSLO: zaokrouhli na 1 des. místo
            num_rounded = round(num, 1)
            # Rozděl na celou a desetinnou část
            integer_part_num = int(abs(num_rounded))
            decimal_part_num = abs(num_rounded) - integer_part_num

            # Převeď desetinnou část na string (1 des. místo)
            decimal_str = f"{decimal_part_num:.1f}".split('.')[1]

            num_str = str(integer_part_num)
            decimal_part = decimal_str

        # FORMÁTOVÁNÍ CELÉ ČÁSTI s mezerami (každé 3 cifry zprava)
        if len(num_str) <= 3:
            # Číslo < 1000: bez mezer
            formatted_int = num_str
        else:
            # Číslo >= 1000: přidej mezery
            formatted_int = ""
            for i, digit in enumerate(reversed(num_str)):
                if i > 0 and i % 3 == 0:
                    formatted_int = " " + formatted_int
                formatted_int = digit + formatted_int

        # Handling záporných čísel
        if num < 0:
            formatted_int = "-" + formatted_int

        # VÝSLEDEK
        if decimal_part:
            return f"{formatted_int},{decimal_part}"
        else:
            return formatted_int

    except (ValueError, TypeError):
        # Pokud se nepodaří konvertovat, vrať původní hodnotu
        return str(value) if value else ""


def insert_conditional_text_before_heading(docx_path, input_data, heading_text="ČASOVÉ ROZLOŽENÍ PRACOVNÍ SMĚNY"):
    """
    Vloží podmíněný text před zadaný nadpis v Word dokumentu.

    POŽADAVEK 2: Podle hodnoty 'what_is_evaluated' vloží buď text pro KUSY nebo ČAS
    před nadpis "ČASOVÉ ROZLOŽENÍ PRACOVNÍ SMĚNY".

    Args:
        docx_path: Cesta k vygenerovanému Word dokumentu
        input_data: measurement_data.json (obsahuje section3_additional_data.what_is_evaluated)
        heading_text: Text nadpisu, před který se má vložit podmíněný text

    Returns:
        bool: True pokud se podařilo vložit text, False jinak
    """
    try:
        # Načti hodnotu what_is_evaluated
        what_is_evaluated = input_data.get("section3_additional_data", {}).get("what_is_evaluated", "kusy").lower()
        workplace = input_data.get("section2_firma", {}).get("workplace", "[pracoviště]")

        # Definuj texty podle podmínky
        if what_is_evaluated == "kusy":
            conditional_heading = "Norma"
            conditional_text = (
                f"Průměrná směna odpovídala stanovené normě pro pracoviště {workplace}, "
                f"která byla v den měření schválena zaměstnavatelem. Dle této normy byly přepočteny "
                f"počty pohybů a svalové síly rukou a předloktí obou horních končetin. "
                f"V den měření byla norma na lince stanovena na 228 ks/směna."
            )
        else:  # "čas"
            conditional_heading = "ČAS"
            conditional_text = (
                f"Průměrná směna vychází z časového snímku (viz níže), "
                f"jenž byl schválen zaměstnavatelem. Dle tohoto časového snímku byly přepočteny "
                f"pohyby a svalové síly rukou a předloktí obou horních končetin."
            )

        # Otevři dokument pomocí python-docx
        doc = Document(docx_path)

        # KROK 1: Hledej v normálních paragrafech
        for i, paragraph in enumerate(doc.paragraphs):
            if heading_text in paragraph.text:
                # Vlož nadpis před hlavní nadpis (centrovaný)
                new_heading = doc.paragraphs[i].insert_paragraph_before(conditional_heading)
                new_heading.alignment = 1  # 1 = CENTER

                # Vlož text před hlavní nadpis (normální formát)
                new_para = doc.paragraphs[i + 1].insert_paragraph_before(conditional_text)
                new_para.alignment = 3  # 3 = JUSTIFY (oboustranné zarovnání)

                print(f"  ✓ Podmíněný text vložen před paragraph (typ: {what_is_evaluated.upper()})")

                # Ulož dokument
                doc.save(docx_path)
                return True

        # KROK 2: Pokud nenalezen v paragrafech, hledej v tabulkách
        print(f"  → Nadpis nenalezen v paragrafech, hledám v tabulkách...")

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    if heading_text in cell.text:
                        # Nadpis nalezen v tabulce! Vlož text PŘED tabulku
                        # Najdi pozici tabulky v dokumentu
                        table_element = table._element
                        parent = table_element.getparent()

                        # Vlož 2 nové paragrafy před tabulku
                        # 1. Nadpis (centrovaný)
                        new_heading_para = doc.add_paragraph(conditional_heading)
                        new_heading_para.alignment = 1  # CENTER
                        parent.insert(parent.index(table_element), new_heading_para._element)

                        # 2. Text (zarovnaný)
                        new_text_para = doc.add_paragraph(conditional_text)
                        new_text_para.alignment = 3  # JUSTIFY
                        parent.insert(parent.index(table_element), new_text_para._element)

                        print(f"  ✓ Podmíněný text vložen před tabulku #{table_idx} (typ: {what_is_evaluated.upper()})")

                        # Ulož dokument
                        doc.save(docx_path)
                        return True

        print(f"  ⚠ Varování: Nadpis '{heading_text}' nebyl nalezen ani v paragrafech ani v tabulkách")
        return False

    except Exception as e:
        print(f"  ⚠ Chyba při vkládání podmíněného textu: {e}")
        import traceback
        traceback.print_exc()
        return False


def generate_word_protocol_v1(measurement_json, results_json, template_path, output_path):
    """
    VARIANTA 1: Vnořená struktura
    V Word šabloně používáš: {{ input.section1_firma.company }} a {{ results.Fmax_Phk_Extenzor }}
    """
    # Získej projekt folder (parent measurement_json)
    project_folder = Path(measurement_json).parent

    # Načti oba JSONy
    with open(measurement_json, encoding='utf-8') as f:
        input_data = json.load(f)

    with open(results_json, encoding='utf-8') as f:
        results_data = json.load(f)

    # Vygeneruj podmínkové texty a přidej do input_data
    conditional_texts = generate_conditional_texts(input_data, results_data)
    input_data["section_generated_texts"] = conditional_texts

    # Konvertuj tabulky z dict na list (pro indexování v Jinja2)
    # Vytvoř 1-indexed list (s dummy elementem na indexu 0)
    # Padded až do indexu 50 pro Word template placeholders
    MAX_ROWS = 50
    for key, value in results_data.items():
        if isinstance(value, dict) and value and all(k.isdigit() for k in value.keys()):
            # Převeď slovník s číselnými klíči na list (1-indexed)
            # Padded s prázdnými daty pro template placeholders
            results_data[key] = [None] + [value.get(str(i), {}) for i in range(1, MAX_ROWS + 1)]

    # Vytvoř vnořenou strukturu
    context = {
        "input": input_data,
        "results": results_data
    }

    # Vygeneruj Word
    doc = DocxTemplate(template_path)

    # ZAKOMENTOVÁNO: Vkládání grafů zakázáno
    # # Přidat grafy jako obrázky (z project folder) - PNG formát
    # graf1_path = project_folder / "lsz_charts" / "graf1.png"
    # graf2_path = project_folder / "lsz_charts" / "graf2.png"
    #
    # if graf1_path.exists():
    #     context["graf1"] = InlineImage(doc, str(graf1_path), width=Mm(150))
    # else:
    #     context["graf1"] = ""
    #
    # if graf2_path.exists():
    #     context["graf2"] = InlineImage(doc, str(graf2_path), width=Mm(150))
    # else:
    #     context["graf2"] = ""

    doc.render(context)
    doc.save(output_path)

    # POST-PROCESSING: Zvýrazni vybrané holteru tučně
    selected_holters = get_selected_holter_numbers(input_data)
    if selected_holters:
        print(f"  → Zvýrazňuji holteru: {selected_holters}")
        highlight_selected_holters(output_path, selected_holters)
        print(f"  ✓ Zvýraznění dokončeno")
    else:
        print(f"  ⚠ Žádné holteru k zvýraznění")

    # POST-PROCESSING: Červeně zvýrazni nadlimitní hodnoty v force_distribution
    print(f"  → Zvýrazňuji nadlimitní hodnoty v tabulce force_distribution...")
    highlight_force_distribution_values(output_path, input_data, results_data)
    print(f"  ✓ Červené zvýraznění dokončeno")

    print(f"[OK] Word vygenerován (VNOŘENÁ STRUKTURA): {output_path}")
    print(f"  - Vstupní sekce: {list(input_data.keys())}")
    print(f"  - Výsledkové tabulky: {len([k for k in results_data.keys() if k.startswith('table_')])}")
    print()
    print("V Word šabloně používej:")
    print("  {{ input.section1_firma.company }}")
    print("  {{ results.Fmax_Phk_Extenzor }}")
    print("  {% for row in results.table_somatometrie %}")


def generate_word_protocol_v2(measurement_json, results_json, template_path, output_path):
    """
    VARIANTA 2: Plochá struktura (flat merge)
    V Word šabloně používáš: {{ section1_firma.company }} a {{ Fmax_Phk_Extenzor }}
    UPOZORNĚNÍ: Pokud mají oba JSONy stejné klíče, dojde ke kolizi!
    """
    # Získej projekt folder (parent measurement_json)
    project_folder = Path(measurement_json).parent

    # Načti oba JSONy
    with open(measurement_json, encoding='utf-8') as f:
        input_data = json.load(f)

    with open(results_json, encoding='utf-8') as f:
        results_data = json.load(f)

    # Vygeneruj podmínkové texty a přidej do input_data
    conditional_texts = generate_conditional_texts(input_data, results_data)
    input_data["section_generated_texts"] = conditional_texts

    # Přejmenuj klíče pro kompatibilitu s Word template
    key_mapping = {
        "section1_firma": "section2_firma",
        "section2_additional_data": "section3_additional_data",
        "section3_worker_a": "section4_worker_a",
        "section4_worker_b": "section5_worker_b",
        "section5_final": "section6_final"
    }

    for old_key, new_key in key_mapping.items():
        if old_key in input_data:
            input_data[new_key] = input_data.pop(old_key)

    # Konvertuj tabulky z dict na list (pro indexování v Jinja2)
    # Vytvoř 1-indexed list (s dummy elementem na indexu 0)
    # Padded až do indexu 50 pro Word template placeholders
    MAX_ROWS = 50
    for key, value in results_data.items():
        if isinstance(value, dict) and value and all(k.isdigit() for k in value.keys()):
            # Převeď slovník s číselnými klíči na list (1-indexed)
            # Padded s prázdnými daty pro template placeholders
            results_data[key] = [None] + [value.get(str(i), {}) for i in range(1, MAX_ROWS + 1)]

    # Slouč do jedné úrovně
    context = {**input_data, **results_data}

    # Načti cestu k Word souboru z JSON
    copied_docx_path = input_data.get("section1_uploaded_docx", {}).get("copied_file_path")

    # Vytvoř template pro přístup k new_subdoc metodě
    tpl = DocxTemplate(template_path)

    # POŽADAVEK 3,4,5: Vytvoř vlastní Jinja2 environment s custom filtry
    jinja_env = Environment()
    jinja_env.filters['czech'] = format_czech_number
    print("✓ Custom filtr zaregistrován: |czech (univerzální formátování)")

    # Přidat subdoc do contextu
    if copied_docx_path and Path(copied_docx_path).exists():
        # POŽADAVEK 2: NEJDŘÍV uprav subdokument (vlož podmíněný text)
        print(f"  → Upravuji subdokument před vložením...")
        insert_conditional_text_before_heading(copied_docx_path, input_data)

        # TEĎ TEPRVE načti upravený subdokument
        subdoc = tpl.new_subdoc(copied_docx_path)
        context["popisprace"] = subdoc
        print(f"✓ Subdokument načten: {Path(copied_docx_path).name}")
    else:
        print("⚠ Varování: Popis práce nebyl nalezen, placeholder zůstane prázdný")
        context["popisprace"] = ""

    # POŽADAVEK 7: Přidej dnešní datum do contextu
    context["today_date"] = datetime.now().strftime("%d.%m.%Y")
    print(f"✓ Dnešní datum přidáno: {context['today_date']}")

    # ZAKOMENTOVÁNO: Vkládání grafů zakázáno
    # # Přidat grafy jako obrázky (z project folder) - PNG formát
    # graf1_path = project_folder / "lsz_charts" / "graf1.png"
    # graf2_path = project_folder / "lsz_charts" / "graf2.png"
    #
    # if graf1_path.exists():
    #     context["graf1"] = InlineImage(tpl, str(graf1_path), width=Mm(150))
    #     print(f"✓ Graf 1 přidán do Word")
    # else:
    #     print(f"⚠ Graf 1 nenalezen: {graf1_path}")
    #     context["graf1"] = ""
    #
    # if graf2_path.exists():
    #     context["graf2"] = InlineImage(tpl, str(graf2_path), width=Mm(150))
    #     print(f"✓ Graf 2 přidán do Word")
    # else:
    #     print(f"⚠ Graf 2 nenalezen: {graf2_path}")
    #     context["graf2"] = ""

    # Vygeneruj Word s vlastním Jinja2 environment (obsahuje custom filtry)
    tpl.render(context, jinja_env=jinja_env)
    tpl.save(output_path)

    # POST-PROCESSING: Zvýrazni vybrané holteru tučně
    selected_holters = get_selected_holter_numbers(input_data)
    if selected_holters:
        print(f"  → Zvýrazňuji holteru: {selected_holters}")
        highlight_selected_holters(output_path, selected_holters)
        print(f"  ✓ Zvýraznění dokončeno")
    else:
        print(f"  ⚠ Žádné holteru k zvýraznění")

    # POST-PROCESSING: Červeně zvýrazni nadlimitní hodnoty v force_distribution
    print(f"  → Zvýrazňuji nadlimitní hodnoty v tabulce force_distribution...")
    highlight_force_distribution_values(output_path, input_data, results_data)
    print(f"  ✓ Červené zvýraznění dokončeno")

    # POZNÁMKA: Podmíněný text (požadavek 2) byl vložen PŘED načtením subdokumentu (viz výše)

    print(f"[OK] Word vygenerován (PLOCHÁ STRUKTURA): {output_path}")
    print(f"  - Celkový počet klíčů: {len(context)}")
    print()
    print("V Word šabloně používej:")
    print("  {{ section1_firma.company }}")
    print("  {{ Fmax_Phk_Extenzor|czech }}  ← Číslo (auto-detekce: desetinné/celé)")
    print("  {{ phk_number_of_movements|czech }}  ← Číslo (auto-detekce: desetinné/celé)")
    print("  {{ today_date }}  ← Dnešní datum")
    print("  {% for row in table_somatometrie %}")
    print()
    print("Filtr |czech automaticky:")
    print("  - Desetinná: zaokrouhlí na 1 des. místo + čárka (8.55 → 8,6)")
    print("  - Celá: žádná des. místa (450 → 450)")
    print("  - Tisíce: mezery (2222 → 2 222, 12345.67 → 12 345,7)")


def generate_word_protocol_v3(measurement_json, results_json, template_path, output_path):
    """
    VARIANTA 3: Prefixovaná struktura
    V Word šabloně používáš: {{ m.section1_firma.company }} a {{ r.Fmax_Phk_Extenzor }}
    Krátké prefixy, jasná separace
    """
    # Získej projekt folder (parent measurement_json)
    project_folder = Path(measurement_json).parent

    # Načti oba JSONy
    with open(measurement_json, encoding='utf-8') as f:
        input_data = json.load(f)

    with open(results_json, encoding='utf-8') as f:
        results_data = json.load(f)

    # Vygeneruj podmínkové texty a přidej do input_data
    conditional_texts = generate_conditional_texts(input_data, results_data)
    input_data["section_generated_texts"] = conditional_texts

    # Konvertuj tabulky z dict na list (pro indexování v Jinja2)
    # Vytvoř 1-indexed list (s dummy elementem na indexu 0)
    # Padded až do indexu 50 pro Word template placeholders
    MAX_ROWS = 50
    for key, value in results_data.items():
        if isinstance(value, dict) and value and all(k.isdigit() for k in value.keys()):
            # Převeď slovník s číselnými klíči na list (1-indexed)
            # Padded s prázdnými daty pro template placeholders
            results_data[key] = [None] + [value.get(str(i), {}) for i in range(1, MAX_ROWS + 1)]

    # Vytvoř prefixovanou strukturu
    context = {
        "m": input_data,
        "r": results_data
    }

    # Vygeneruj Word
    doc = DocxTemplate(template_path)

    # ZAKOMENTOVÁNO: Vkládání grafů zakázáno
    # # Přidat grafy jako obrázky (z project folder) - PNG formát
    # graf1_path = project_folder / "lsz_charts" / "graf1.png"
    # graf2_path = project_folder / "lsz_charts" / "graf2.png"
    #
    # if graf1_path.exists():
    #     context["graf1"] = InlineImage(doc, str(graf1_path), width=Mm(150))
    # else:
    #     context["graf1"] = ""
    #
    # if graf2_path.exists():
    #     context["graf2"] = InlineImage(doc, str(graf2_path), width=Mm(150))
    # else:
    #     context["graf2"] = ""

    doc.render(context)
    doc.save(output_path)

    # POST-PROCESSING: Zvýrazni vybrané holteru tučně
    selected_holters = get_selected_holter_numbers(input_data)
    if selected_holters:
        print(f"  → Zvýrazňuji holteru: {selected_holters}")
        highlight_selected_holters(output_path, selected_holters)
        print(f"  ✓ Zvýraznění dokončeno")
    else:
        print(f"  ⚠ Žádné holteru k zvýraznění")

    # POST-PROCESSING: Červeně zvýrazni nadlimitní hodnoty v force_distribution
    print(f"  → Zvýrazňuji nadlimitní hodnoty v tabulce force_distribution...")
    highlight_force_distribution_values(output_path, input_data, results_data)
    print(f"  ✓ Červené zvýraznění dokončeno")

    print(f"[OK] Word vygenerován (PREFIXOVANÁ STRUKTURA): {output_path}")
    print()
    print("V Word šabloně používej:")
    print("  {{ m.section1_firma.company }}")
    print("  {{ r.Fmax_Phk_Extenzor }}")
    print("  {% for row in r.table_somatometrie %}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Generuje Word protokol ze dvou JSON souborů')
    parser.add_argument('measurement_json', help='Cesta k measurement_data JSON souboru')
    parser.add_argument('results_json', help='Cesta k lsz_results JSON souboru')
    parser.add_argument('template_path', help='Cesta k Word šabloně (.docx)')
    parser.add_argument('output_path', help='Cesta k výstupnímu Word souboru')
    parser.add_argument('--variant', choices=['v1', 'v2', 'v3'], default='v2',
                        help='Varianta generování (v1=vnořená, v2=plochá, v3=prefixovaná)')

    args = parser.parse_args()

    # Spusť vybranou variantu
    if args.variant == 'v1':
        generate_word_protocol_v1(args.measurement_json, args.results_json,
                                  args.template_path, args.output_path)
    elif args.variant == 'v2':
        generate_word_protocol_v2(args.measurement_json, args.results_json,
                                  args.template_path, args.output_path)
    elif args.variant == 'v3':
        generate_word_protocol_v3(args.measurement_json, args.results_json,
                                  args.template_path, args.output_path)
