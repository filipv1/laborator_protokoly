"""
Debug skript pro kontrolu PP tabulek v dokumentu
"""
from docx import Document
from pathlib import Path


def analyze_pp_document(docx_path):
    """Analyzuje PP dokument a vypíše detailní info o tabulkách"""
    doc = Document(docx_path)

    pp_keywords = [
        "TRUP", "Trup",
        "HLAVA_KRK", "HLAVA A KRK", "Hlava a krk",
        "PHK", "Pravá horní končetina",
        "LHK", "Levá horní končetina",
        "DOLNÍ KONČETINY", "Dolní končetiny", "DKK",
        "OSTATNÍ ČÁSTI TĚLA", "Ostatní části těla", "OST",
        "Statická", "Dynamická"
    ]

    print(f"Analyzuji dokument: {docx_path}")
    print(f"Celkem tabulek: {len(doc.tables)}")
    print("="*60)

    for table_idx, table in enumerate(doc.tables):
        print(f"\nTABULKA {table_idx}:")
        print(f"  Pocet sloupcu: {len(table.columns)}")
        print(f"  Pocet radku: {len(table.rows)}")

        if len(table.columns) == 6:
            # Zkontroluj PP keywords
            is_pp_table = False
            for row in table.rows[:5]:
                row_text = " ".join(cell.text for cell in row.cells)
                for keyword in pp_keywords:
                    if keyword in row_text:
                        is_pp_table = True
                        break
                if is_pp_table:
                    break

            if is_pp_table:
                print("  -> Je to PP tabulka!")
                print("  Obsah radku:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell_idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if cell_idx in [2, 3, 4]:  # Hodnoty sloupce
                            row_text.append(f"[{text}]")
                        else:
                            # Odstran diakritiku pro vypis
                            safe_text = text.encode('ascii', 'replace').decode('ascii')
                            row_text.append(safe_text[:30] + "..." if len(safe_text) > 30 else safe_text)
                    print(f"    Radek {row_idx}: {' | '.join(row_text)}")


def main():
    # Analyzuj aktuální PP dokument
    pp_path = Path("projects/325_Ascari_sro/PP_325_Ascari_sro_CAS_protokol.docx")
    if pp_path.exists():
        analyze_pp_document(pp_path)
    else:
        print(f"Soubor neexistuje: {pp_path}")


if __name__ == "__main__":
    main()