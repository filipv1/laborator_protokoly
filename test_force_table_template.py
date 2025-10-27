"""
Vytvorit jednoduchou Word sablonu s tabulkou force_distribution pro test cerveneho zvyrazneni
"""
from docx import Document
from docx.shared import Pt, Inches

# Vytvorit novy dokument
doc = Document()
doc.add_heading('Test Tabulky Force Distribution', 0)

# Pridej paragraf
doc.add_paragraph('Tato tabulka ukazuje rozdeleni svalovych sil:')

# Vytvorit tabulku (9 sloupcu: activity + 8 hodnot)
# Radek 1: Hlavicka
# Radek 2-5: Merene osoby a prazdne radky
# Radek 6: Casove vazeny prumer (TENTO RADEK BUDE MIT CERVENE HODNOTY)

table = doc.add_table(rows=6, cols=9)
table.style = 'Table Grid'

# Hlavicka
header_cells = table.rows[0].cells
header_cells[0].text = 'Rozpis pracovnich operaci'
header_cells[1].text = 'Extenzory PHK'
header_cells[2].text = 'Flexory PHK'
header_cells[3].text = 'Extenzory LHK'
header_cells[4].text = 'Flexory LHK'
header_cells[5].text = 'Extenzory PHK (>70%)'
header_cells[6].text = 'Flexory PHK (>70%)'
header_cells[7].text = 'Extenzory LHK (>70%)'
header_cells[8].text = 'Flexory LHK (>70%)'

# Radek 2: 1. merena osoba
row1 = table.rows[1].cells
row1[0].text = '1. merena osoba - Zakladani'
for i in range(1, 9):
    row1[i].text = '0'

# Radek 3: 2. merena osoba
row2 = table.rows[2].cells
row2[0].text = '2. merena osoba - Zakladani'
for i in range(1, 9):
    row2[i].text = '0'

# Radek 4: Prazdny
row3 = table.rows[3].cells
row3[0].text = ''
for i in range(1, 9):
    row3[i].text = ''

# Radek 5: Prazdny
row4 = table.rows[4].cells
row4[0].text = ''
for i in range(1, 9):
    row4[i].text = ''

# Radek 6: Casove vazeny prumer (REALNA DATA S PLACEHOLDERY)
row5 = table.rows[5].cells
row5[0].text = 'Casove vazeny prumer - Zakladani'
row5[1].text = '{{ results.table_force_distribution.21.force_55_70_phk_extenzory }}'
row5[2].text = '{{ results.table_force_distribution.21.force_55_70_phk_flexory }}'
row5[3].text = '{{ results.table_force_distribution.21.force_55_70_lhk_extenzory }}'
row5[4].text = '{{ results.table_force_distribution.21.force_55_70_lhk_flexory }}'
row5[5].text = '{{ results.table_force_distribution.21.force_over_70_phk_extenzory }}'
row5[6].text = '{{ results.table_force_distribution.21.force_over_70_phk_flexory }}'
row5[7].text = '{{ results.table_force_distribution.21.force_over_70_lhk_extenzory }}'
row5[8].text = '{{ results.table_force_distribution.21.force_over_70_lhk_flexory }}'

# Ulozit dokument
output_path = 'test_force_template.docx'
doc.save(output_path)
print(f'Sablona vytvorena: {output_path}')
print('')
print('Tato sablona obsahuje:')
print('  - Tabulku s placeholdery pro results.table_force_distribution.21')
print('  - Po vygenerovani bude hodnota 200 v sloupci 5 (force_over_70_phk_extenzory) CERVENA')
print('')
print('Spust generovani:')
print('  python generate_word_from_two_sources.py')
print('    "projects\\222_rentury\\measurement_data.json"')
print('    "projects\\222_rentury\\lsz_results.json"')
print('    "test_force_template.docx"')
print('    "test_force_output_CERVENE.docx"')
