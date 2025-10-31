"""
Detailni analyza textu v LSZ dokumentech
"""
from docx import Document
from pathlib import Path
import sys
import io
import difflib

# Nastaveni UTF-8 pro Windows konzoli
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

def get_all_text(doc_path):
    """Ziska vsechny texty z dokumentu"""
    doc = Document(doc_path)

    texts = {
        'paragraphs': [],
        'tables': []
    }

    # Paragrafy
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            texts['paragraphs'].append({
                'index': i,
                'text': para.text.strip()
            })

    # Tabulky
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        texts['tables'].append({
            'index': i,
            'data': table_data
        })

    return texts

def highlight_differences(text1, text2):
    """Zvyrazni rozdily mezi dvema texty"""
    if text1 == text2:
        return None

    # Pouzijeme difflib pro detekci rozdilu
    diff = list(difflib.ndiff(text1.split(), text2.split()))

    removed = []
    added = []

    for word in diff:
        if word.startswith('- '):
            removed.append(word[2:])
        elif word.startswith('+ '):
            added.append(word[2:])

    return {
        'removed': ' '.join(removed),
        'added': ' '.join(added)
    }

def main():
    base_path = Path(r"C:\Users\vaclavik\lab5\laborator_protokoly\Vzorové protokoly")

    docs = {
        'ref_2muzi': base_path / "Autorizované protokoly pro MUŽE" / "LSZ_XX_Firma_Pozice.docx",
        '2zeny': base_path / "Autorizované protokoly pro ŽENY" / "LSZ_XX_Firma_Pozice.docx",
        '1muz': base_path / "Jeden zaměstnanec" / "LSZ_jeden_MUŽ.DOCX"
    }

    print("=" * 100)
    print("DETAILNÍ TEXTOVÉ ROZDÍLY V LSZ DOKUMENTECH")
    print("=" * 100)
    print()

    # Nacteni dokumentu
    texts = {}
    for key, path in docs.items():
        if path.exists():
            texts[key] = get_all_text(path)
            print(f"Načten: {key}")

    print()

    if 'ref_2muzi' not in texts:
        print("CHYBA: Referenční dokument nenalezen!")
        return

    ref = texts['ref_2muzi']

    # ====================
    # ANALÝZA: 2 ŽENY
    # ====================
    if '2zeny' in texts:
        print("=" * 100)
        print("ROZDÍLY: LSZ PRO 2 ŽENY vs LSZ PRO 2 MUŽE (REFERENČNÍ)")
        print("=" * 100)
        print()

        comp = texts['2zeny']

        print(f"PARAGRAFŮ - REF: {len(ref['paragraphs'])}, ŽENY: {len(comp['paragraphs'])}")
        print(f"TABULEK - REF: {len(ref['tables'])}, ŽENY: {len(comp['tables'])}")
        print()

        # Porovnání paragrafů
        print("-" * 100)
        print("TEXTOVÉ ROZDÍLY V PARAGRAFECH:")
        print("-" * 100)

        diff_count = 0
        for i in range(min(len(ref['paragraphs']), len(comp['paragraphs']))):
            ref_text = ref['paragraphs'][i]['text']
            comp_text = comp['paragraphs'][i]['text']

            if ref_text != comp_text:
                diff_count += 1
                print(f"\n[PARAGRAF #{i}]")
                print(f"REF:  {ref_text}")
                print(f"ŽENY: {comp_text}")

                diff_detail = highlight_differences(ref_text, comp_text)
                if diff_detail:
                    print(f"  ❌ ODSTRANĚNO: {diff_detail['removed']}")
                    print(f"  ✅ PŘIDÁNO: {diff_detail['added']}")

        print(f"\n\nCelkem rozdílných paragrafů: {diff_count}")

        # Porovnání tabulek
        print("\n" + "-" * 100)
        print("ROZDÍLY V TABULKÁCH:")
        print("-" * 100)

        for i in range(min(len(ref['tables']), len(comp['tables']))):
            ref_table = ref['tables'][i]['data']
            comp_table = comp['tables'][i]['data']

            if ref_table != comp_table:
                print(f"\n[TABULKA #{i}]")
                print(f"  Řádků - REF: {len(ref_table)}, ŽENY: {len(comp_table)}")

                # Porovnání první řádky
                if len(ref_table) > 0 and len(comp_table) > 0:
                    if ref_table[0] != comp_table[0]:
                        print(f"  První řádek se liší:")
                        print(f"    REF:  {ref_table[0]}")
                        print(f"    ŽENY: {comp_table[0]}")

    # ====================
    # ANALÝZA: 1 MUŽ
    # ====================
    if '1muz' in texts:
        print("\n" + "=" * 100)
        print("ROZDÍLY: LSZ PRO 1 MUŽE vs LSZ PRO 2 MUŽE (REFERENČNÍ)")
        print("=" * 100)
        print()

        comp = texts['1muz']

        print(f"PARAGRAFŮ - REF: {len(ref['paragraphs'])}, 1MUŽ: {len(comp['paragraphs'])}")
        print(f"TABULEK - REF: {len(ref['tables'])}, 1MUŽ: {len(comp['tables'])}")
        print()

        # Strukturální rozdíly
        print("-" * 100)
        print("STRUKTURÁLNÍ ZMĚNY:")
        print("-" * 100)
        print(f"• Odstraněno {len(ref['paragraphs']) - len(comp['paragraphs'])} paragrafů")
        print(f"• Odstraněno {len(ref['tables']) - len(comp['tables'])} tabulek")
        print()

        # Příklady textových změn
        print("-" * 100)
        print("PŘÍKLADY TEXTOVÝCH ZMĚN (prvních 15):")
        print("-" * 100)

        diff_count = 0
        shown = 0
        max_show = 15

        for i in range(min(len(ref['paragraphs']), len(comp['paragraphs']))):
            ref_text = ref['paragraphs'][i]['text']
            comp_text = comp['paragraphs'][i]['text']

            if ref_text != comp_text:
                diff_count += 1
                if shown < max_show:
                    shown += 1
                    print(f"\n[PARAGRAF #{i}]")

                    # Zkraceni dlouhych textu
                    ref_short = ref_text if len(ref_text) < 150 else ref_text[:150] + "..."
                    comp_short = comp_text if len(comp_text) < 150 else comp_text[:150] + "..."

                    print(f"REF:  {ref_short}")
                    print(f"1MUŽ: {comp_short}")

                    diff_detail = highlight_differences(ref_text, comp_text)
                    if diff_detail and diff_detail['removed'] and diff_detail['added']:
                        # Zkraceni diff vystupu
                        removed = diff_detail['removed']
                        added = diff_detail['added']
                        if len(removed) > 100:
                            removed = removed[:100] + "..."
                        if len(added) > 100:
                            added = added[:100] + "..."
                        print(f"  ❌ ODSTRANĚNO: {removed}")
                        print(f"  ✅ PŘIDÁNO: {added}")

        print(f"\n\nCelkem rozdílných paragrafů: {diff_count} (zobrazeno prvních {shown})")

        # Rozdíly v tabulkách
        print("\n" + "-" * 100)
        print("ROZDÍLY V TABULKÁCH:")
        print("-" * 100)

        table_diffs = 0
        for i in range(min(len(ref['tables']), len(comp['tables']))):
            ref_table = ref['tables'][i]['data']
            comp_table = comp['tables'][i]['data']

            if ref_table != comp_table:
                table_diffs += 1
                print(f"\n[TABULKA #{i}]")
                print(f"  Řádků - REF: {len(ref_table)}, 1MUŽ: {len(comp_table)}")
                print(f"  Sloupců - REF: {len(ref_table[0]) if ref_table else 0}, "
                      f"1MUŽ: {len(comp_table[0]) if comp_table else 0}")

                # Ukázka prvních řádků
                if len(ref_table) > 0 and len(comp_table) > 0:
                    if ref_table[0] != comp_table[0]:
                        print(f"  První řádek:")
                        print(f"    REF:  {ref_table[0]}")
                        print(f"    1MUŽ: {comp_table[0]}")

        print(f"\n\nTabulek s rozdíly: {table_diffs}")

        # Chybějící tabulka
        if len(ref['tables']) > len(comp['tables']):
            print(f"\n❌ CHYBĚJÍCÍ TABULKA v dokumentu pro 1 muže:")
            print(f"   Poslední tabulka z referenčního dokumentu není přítomna.")

if __name__ == "__main__":
    main()
