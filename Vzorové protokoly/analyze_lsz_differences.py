"""
Analyza rozdilu mezi LSZ Word sablonami
"""
from docx import Document
from pathlib import Path
import json
import sys
import io

# Nastaveni UTF-8 pro Windows konzoli
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

def analyze_document(doc_path):
    """Extrahuje strukturu dokumentu"""
    doc = Document(doc_path)

    analysis = {
        'path': str(doc_path),
        'num_paragraphs': len(doc.paragraphs),
        'num_tables': len(doc.tables),
        'paragraphs': [],
        'tables': []
    }

    # Analýza paragrafů
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:  # Pouze neprázdné paragrafy
            analysis['paragraphs'].append({
                'index': i,
                'text': text[:200],  # První 200 znaků
                'style': para.style.name if para.style else None
            })

    # Analýza tabulek
    for i, table in enumerate(doc.tables):
        table_info = {
            'index': i,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'first_row': []
        }

        # První řádek tabulky
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                table_info['first_row'].append(cell.text.strip())

        # Ukázka obsahu (první 3 řádky)
        table_info['sample_content'] = []
        for row_idx, row in enumerate(table.rows[:3]):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_info['sample_content'].append(row_data)

        analysis['tables'].append(table_info)

    return analysis


def compare_documents(ref_analysis, comp_analysis, comp_name):
    """Porovná dva dokumenty a najde rozdíly"""
    differences = {
        'name': comp_name,
        'differences': []
    }

    # Porovnání počtu paragrafů
    if ref_analysis['num_paragraphs'] != comp_analysis['num_paragraphs']:
        differences['differences'].append({
            'type': 'paragraph_count',
            'ref': ref_analysis['num_paragraphs'],
            'comp': comp_analysis['num_paragraphs'],
            'diff': comp_analysis['num_paragraphs'] - ref_analysis['num_paragraphs']
        })

    # Porovnání počtu tabulek
    if ref_analysis['num_tables'] != comp_analysis['num_tables']:
        differences['differences'].append({
            'type': 'table_count',
            'ref': ref_analysis['num_tables'],
            'comp': comp_analysis['num_tables'],
            'diff': comp_analysis['num_tables'] - ref_analysis['num_tables']
        })

    # Porovnání obsahu paragrafů
    text_differences = []
    max_paras = min(len(ref_analysis['paragraphs']), len(comp_analysis['paragraphs']))

    for i in range(max_paras):
        ref_para = ref_analysis['paragraphs'][i]
        comp_para = comp_analysis['paragraphs'][i]

        if ref_para['text'] != comp_para['text']:
            text_differences.append({
                'paragraph_index': i,
                'ref_text': ref_para['text'],
                'comp_text': comp_para['text']
            })

    if text_differences:
        differences['differences'].append({
            'type': 'text_content',
            'count': len(text_differences),
            'examples': text_differences[:10]  # První 10 rozdílů
        })

    # Porovnání tabulek
    table_differences = []
    max_tables = min(len(ref_analysis['tables']), len(comp_analysis['tables']))

    for i in range(max_tables):
        ref_table = ref_analysis['tables'][i]
        comp_table = comp_analysis['tables'][i]

        table_diff = {'table_index': i}

        if ref_table['rows'] != comp_table['rows']:
            table_diff['row_count'] = {
                'ref': ref_table['rows'],
                'comp': comp_table['rows']
            }

        if ref_table['cols'] != comp_table['cols']:
            table_diff['col_count'] = {
                'ref': ref_table['cols'],
                'comp': comp_table['cols']
            }

        if ref_table['first_row'] != comp_table['first_row']:
            table_diff['first_row'] = {
                'ref': ref_table['first_row'],
                'comp': comp_table['first_row']
            }

        if len(table_diff) > 1:  # Pokud jsou nějaké rozdíly kromě indexu
            table_differences.append(table_diff)

    if table_differences:
        differences['differences'].append({
            'type': 'tables',
            'count': len(table_differences),
            'details': table_differences
        })

    return differences


def main():
    base_path = Path(r"C:\Users\vaclavik\lab5\laborator_protokoly\Vzorové protokoly")

    # Cesty k dokumentům
    docs = {
        'ref_2muzi': base_path / "Autorizované protokoly pro MUŽE" / "LSZ_XX_Firma_Pozice.docx",
        '2zeny': base_path / "Autorizované protokoly pro ŽENY" / "LSZ_XX_Firma_Pozice.docx",
        '1muz': base_path / "Jeden zaměstnanec" / "LSZ_jeden_MUŽ.DOCX"
    }

    print("=" * 80)
    print("ANALÝZA ROZDÍLŮ V LSZ DOKUMENTECH")
    print("=" * 80)
    print()

    # Načtení všech dokumentů
    print("Načítám dokumenty...")
    analyses = {}
    for key, path in docs.items():
        if path.exists():
            print(f"  ✓ {key}: {path.name}")
            analyses[key] = analyze_document(path)
        else:
            print(f"  ✗ {key}: NENALEZEN - {path}")

    print()

    if 'ref_2muzi' not in analyses:
        print("CHYBA: Referenční dokument (2 muži) nenalezen!")
        return

    # Základní statistiky
    print("=" * 80)
    print("ZÁKLADNÍ STATISTIKY")
    print("=" * 80)
    for key, analysis in analyses.items():
        print(f"\n{key}:")
        print(f"  Počet paragrafů: {analysis['num_paragraphs']}")
        print(f"  Počet tabulek: {analysis['num_tables']}")

    # Porovnání s referenčním dokumentem
    ref = analyses['ref_2muzi']

    print("\n" + "=" * 80)
    print("DETAILNÍ ROZDÍLY OPROTI LSZ PRO 2 MUŽE (REFERENČNÍ)")
    print("=" * 80)

    for key in ['2zeny', '1muz']:
        if key in analyses:
            print(f"\n{'=' * 80}")
            print(f"ROZDÍLY: {key.upper()}")
            print('=' * 80)

            comp_result = compare_documents(ref, analyses[key], key)

            if not comp_result['differences']:
                print("  Žádné strukturální rozdíly nenalezeny.")
            else:
                for diff in comp_result['differences']:
                    print(f"\n  [{diff['type'].upper()}]")

                    if diff['type'] == 'paragraph_count':
                        print(f"    Referenční: {diff['ref']} paragrafů")
                        print(f"    Porovnávaný: {diff['comp']} paragrafů")
                        print(f"    Rozdíl: {diff['diff']:+d} paragrafů")

                    elif diff['type'] == 'table_count':
                        print(f"    Referenční: {diff['ref']} tabulek")
                        print(f"    Porovnávaný: {diff['comp']} tabulek")
                        print(f"    Rozdíl: {diff['diff']:+d} tabulek")

                    elif diff['type'] == 'text_content':
                        print(f"    Počet rozdílných paragrafů: {diff['count']}")
                        print(f"    Příklady prvních rozdílů:")
                        for i, example in enumerate(diff['examples'][:5], 1):
                            print(f"\n      Paragraf #{example['paragraph_index']}:")
                            print(f"        REF: {example['ref_text'][:100]}...")
                            print(f"        COMP: {example['comp_text'][:100]}...")

                    elif diff['type'] == 'tables':
                        print(f"    Počet tabulek s rozdíly: {diff['count']}")
                        for table_diff in diff['details']:
                            print(f"\n      Tabulka #{table_diff['table_index']}:")
                            if 'row_count' in table_diff:
                                print(f"        Řádky - REF: {table_diff['row_count']['ref']}, "
                                      f"COMP: {table_diff['row_count']['comp']}")
                            if 'col_count' in table_diff:
                                print(f"        Sloupce - REF: {table_diff['col_count']['ref']}, "
                                      f"COMP: {table_diff['col_count']['comp']}")
                            if 'first_row' in table_diff:
                                print(f"        První řádek:")
                                print(f"          REF: {table_diff['first_row']['ref']}")
                                print(f"          COMP: {table_diff['first_row']['comp']}")

    # Uložení detailní analýzy do JSON
    output_file = base_path / "lsz_analysis_results.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(analyses, f, ensure_ascii=False, indent=2)

    print(f"\n\n{'=' * 80}")
    print(f"Detailní analýza uložena do: {output_file}")
    print('=' * 80)


if __name__ == "__main__":
    main()
