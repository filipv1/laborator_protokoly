"""
Test všech tří variant práce se dvěma JSONy
Vytvoří 3 testovací Word dokumenty
"""
import json
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_test_template():
    """Vytvoří testovací Word šablonu se všemi variantami placeholderů"""
    doc = Document()

    # Nadpis
    heading = doc.add_heading('TEST WORD ŠABLONY - TŘI VARIANTY', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ============================================
    # VARIANTA 1: Vnořená
    # ============================================
    doc.add_heading('VARIANTA 1: Vnořená struktura', 1)
    doc.add_paragraph('Placeholdery: {{ input.XXX }} a {{ results.XXX }}', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Firma: ').bold = True
    p.add_run('{{ input.section1_firma.company }}')

    p = doc.add_paragraph()
    p.add_run('IČO: ').bold = True
    p.add_run('{{ input.section1_firma.ico }}')

    p = doc.add_paragraph()
    p.add_run('Datum měření: ').bold = True
    p.add_run('{{ input.section1_firma.measurement_date }}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Fmax PHK Extenzor: ').bold = True
    p.add_run('{{ results.Fmax_Phk_Extenzor }} N')

    p = doc.add_paragraph()
    p.add_run('Fmax PHK Flexor: ').bold = True
    p.add_run('{{ results.Fmax_Phk_Flexor }} N')

    doc.add_paragraph()
    doc.add_paragraph('Somatometrická data (první řádek):')
    p = doc.add_paragraph('{% for row in results.table_somatometrie %}')
    p = doc.add_paragraph('Datum: {{ row.datum }}, Iniciály: {{ row.inicialy }}, Věk: {{ row.vek_roky }} let')
    p = doc.add_paragraph('{% endfor %}')

    doc.add_page_break()

    # ============================================
    # VARIANTA 2: Plochá
    # ============================================
    doc.add_heading('VARIANTA 2: Plochá struktura', 1)
    doc.add_paragraph('Placeholdery: {{ XXX }} přímo', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Firma: ').bold = True
    p.add_run('{{ section1_firma.company }}')

    p = doc.add_paragraph()
    p.add_run('IČO: ').bold = True
    p.add_run('{{ section1_firma.ico }}')

    p = doc.add_paragraph()
    p.add_run('Fmax PHK Extenzor: ').bold = True
    p.add_run('{{ Fmax_Phk_Extenzor }} N')

    doc.add_paragraph()
    doc.add_paragraph('Somatometrická data:')
    p = doc.add_paragraph('{% for row in table_somatometrie %}')
    p = doc.add_paragraph('Datum: {{ row.datum }}, Věk: {{ row.vek_roky }} let')
    p = doc.add_paragraph('{% endfor %}')

    doc.add_page_break()

    # ============================================
    # VARIANTA 3: Prefixovaná
    # ============================================
    doc.add_heading('VARIANTA 3: Prefixovaná struktura', 1)
    doc.add_paragraph('Placeholdery: {{ m.XXX }} a {{ r.XXX }}', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Firma: ').bold = True
    p.add_run('{{ m.section1_firma.company }}')

    p = doc.add_paragraph()
    p.add_run('IČO: ').bold = True
    p.add_run('{{ m.section1_firma.ico }}')

    p = doc.add_paragraph()
    p.add_run('Fmax PHK Extenzor: ').bold = True
    p.add_run('{{ r.Fmax_Phk_Extenzor }} N')

    doc.add_paragraph()
    doc.add_paragraph('Somatometrická data:')
    p = doc.add_paragraph('{% for row in r.table_somatometrie %}')
    p = doc.add_paragraph('Datum: {{ row.datum }}, Věk: {{ row.vek_roky }} let')
    p = doc.add_paragraph('{% endfor %}')

    # Ulož šablonu
    template_path = 'test_word_template_3variants.docx'
    doc.save(template_path)
    print(f"✓ Vytvořena testovací šablona: {template_path}")
    return template_path


def test_variant_1(template_path):
    """Test varianty 1: Vnořená struktura"""
    # Načti data
    with open('measurement_data_example.json', encoding='utf-8') as f:
        input_data = json.load(f)

    with open('lsz_results.json', encoding='utf-8') as f:
        results_data = json.load(f)

    # Vytvoř vnořený kontext
    context = {
        "input": input_data,
        "results": results_data
    }

    # Vygeneruj
    doc = DocxTemplate(template_path)
    doc.render(context)
    output = 'TEST_variant1_vnorena.docx'
    doc.save(output)
    print(f"✓ Varianta 1 (vnořená): {output}")


def test_variant_2(template_path):
    """Test varianty 2: Plochá struktura"""
    # Načti data
    with open('measurement_data_example.json', encoding='utf-8') as f:
        input_data = json.load(f)

    with open('lsz_results.json', encoding='utf-8') as f:
        results_data = json.load(f)

    # Slouč do ploché struktury
    context = {**input_data, **results_data}

    # Vygeneruj
    doc = DocxTemplate(template_path)
    doc.render(context)
    output = 'TEST_variant2_plocha.docx'
    doc.save(output)
    print(f"✓ Varianta 2 (plochá): {output}")


def test_variant_3(template_path):
    """Test varianty 3: Prefixovaná struktura"""
    # Načti data
    with open('measurement_data_example.json', encoding='utf-8') as f:
        input_data = json.load(f)

    with open('lsz_results.json', encoding='utf-8') as f:
        results_data = json.load(f)

    # Vytvoř prefixovaný kontext
    context = {
        "m": input_data,
        "r": results_data
    }

    # Vygeneruj
    doc = DocxTemplate(template_path)
    doc.render(context)
    output = 'TEST_variant3_prefix.docx'
    doc.save(output)
    print(f"✓ Varianta 3 (prefixovaná): {output}")


if __name__ == "__main__":
    print("=" * 60)
    print("TEST WORD GENEROVÁNÍ SE DVĚMA JSONy")
    print("=" * 60)
    print()

    # Vytvoř testovací šablonu
    template_path = create_test_template()
    print()

    # Otestuj všechny varianty
    print("Generuji testovací dokumenty...")
    print()

    try:
        test_variant_1(template_path)
        test_variant_2(template_path)
        test_variant_3(template_path)

        print()
        print("=" * 60)
        print("✓ HOTOVO! Zkontroluj tyto soubory:")
        print("  1. TEST_variant1_vnorena.docx")
        print("  2. TEST_variant2_plocha.docx")
        print("  3. TEST_variant3_prefix.docx")
        print()
        print("Doporučuji VARIANTU 1 (vnořená) pro jasnou separaci dat!")
        print("=" * 60)

    except Exception as e:
        print(f"❌ Chyba: {e}")
        import traceback
        traceback.print_exc()
