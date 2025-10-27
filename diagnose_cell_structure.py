"""
Diagnostikovat strukturu bunek v LSZ_protokol.docx
"""
import sys
from docx import Document

if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')

def diagnose_cell_structure(docx_path):
    """Analyzuj strukturu bunek v tabulce"""

    print("=" * 80)
    print(f"DIAGNOSTIKA: {docx_path}")
    print("=" * 80)
    print()

    doc = Document(docx_path)

    print(f"Celkem tabulek v dokumentu: {len(doc.tables)}")
    print()

    # Vypis VSECHNY tabulky
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            print(f"Tabulka #{table_idx}: prilis malo radku ({len(table.rows)})")
            continue

        header_row = table.rows[0]
        header_text = " ".join([cell.text for cell in header_row.cells])

        print(f"Tabulka #{table_idx}:")
        print(f"  Hlavicka: {header_text[:100]}...")
        print(f"  Pocet radku: {len(table.rows)}")
        print(f"  Pocet sloupcu: {len(header_row.cells)}")
        print()

    # Najdi tabulku s force distribution
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue

        header_row = table.rows[0]
        header_text = " ".join([cell.text for cell in header_row.cells]).lower()

        if "extenzor" in header_text or "flexor" in header_text:
            if any("innost" in cell.text.lower() or "activity" in cell.text.lower() or "rozpis" in cell.text.lower() for cell in header_row.cells):
                print("=" * 80)
                print(f"FORCE DISTRIBUTION TABULKA: #{table_idx}")
                print("=" * 80)
                print(f"Pocet radku: {len(table.rows)}")
                print()

                # Analyzuj bunky s hodnotami
                for row_idx, row in enumerate(table.rows):
                    activity = row.cells[0].text.strip()

                    if not activity or row_idx == 0:
                        continue

                    print(f"Radek {row_idx}: {activity}")

                    # Zkontroluj sloupce 5-8 (force_over_70)
                    for col_idx in [5, 6, 7, 8]:
                        if col_idx >= len(row.cells):
                            continue

                        cell = row.cells[col_idx]
                        cell_value = cell.text.strip()

                        if not cell_value or cell_value == '0':
                            continue

                        print(f"  Sloupec {col_idx}: hodnota = '{cell_value}'")
                        print(f"    Pocet paragraphs: {len(cell.paragraphs)}")

                        for para_idx, para in enumerate(cell.paragraphs):
                            print(f"      Paragraph {para_idx}: text = '{para.text}'")
                            print(f"        Pocet runs: {len(para.runs)}")

                            for run_idx, run in enumerate(para.runs):
                                print(f"          Run {run_idx}:")
                                print(f"            Text: '{run.text}'")
                                print(f"            Bold: {run.bold}")
                                print(f"            Color: {run.font.color.rgb if run.font.color.rgb else 'None (auto)'}")

                    print()

                break

    print("=" * 80)


if __name__ == "__main__":
    diagnose_cell_structure("projects/222_rentury/LSZ_protokol_NEW.docx")
